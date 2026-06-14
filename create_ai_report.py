# -*- coding: utf-8 -*-
"""
Grup 4 — Yaşca Diş Kliniği Yönetim Sistemi (Lisans Bitirme Projesi)
İSTÜN Mühendislik ve Doğa Bilimleri Fakültesi Bitirme Projesi Yazım Kuralları (EK-3) uyumlu.
- A4, Times New Roman 12pt, 1.5 satır, 2.5 cm kenar boşlukları, 1.25 cm ilk satır girintisi
- Başlık 14pt BÜYÜK, alt başlık 12pt; çizelge başlığı üstte, şekil başlığı altta
- Sayfa numaraları: ön sayfalar i,ii,iii (Romen); Giriş'ten itibaren 1,2,3 (Arap); alt-orta
- Onay, Beyan, Önsöz, Abstract sayfaları; APA kaynaklar
"""
import os, sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ───────────────────────── Yollar ─────────────────────────
ROOT = os.path.dirname(os.path.abspath(__file__))
DIAGRAMS = os.path.join(ROOT, "docs", "diagrams")
SHOTS = os.path.join(ROOT, "docs", "screenshots")
SWAGGER = os.path.join(ROOT, "swagger_screenshot.png")
OUT = os.path.join(ROOT, "Grup4_Yasca_Proje_Raporu.docx")

FONT = "Times New Roman"
FS = Pt(12)
LS = 1.5
INDENT = Cm(1.25)
A4_W, A4_H = Cm(21.0), Cm(29.7)

doc = Document()

# ───────────────────────── Font yardımcıları ─────────────────────────
def set_run_font(run, name=FONT):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{name}" w:cs="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(a), name)

def set_style_font(style_obj, name=FONT, size=None):
    style_obj.font.name = name
    if size:
        style_obj.font.size = size
    el = style_obj.element
    rPr = el.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>'); el.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:eastAsia="{name}" w:cs="{name}"/>')
        rPr.insert(0, rFonts)
    else:
        for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(a), name)

dd = doc.styles.element.find(qn("w:docDefaults"))
if dd is None:
    dd = parse_xml(f'<w:docDefaults {nsdecls("w")}/>'); doc.styles.element.insert(0, dd)
rpd = dd.find(qn("w:rPrDefault"))
if rpd is None:
    rpd = parse_xml(f'<w:rPrDefault {nsdecls("w")}/>'); dd.append(rpd)
rpd.append(parse_xml(
    f'<w:rPr {nsdecls("w")}><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:eastAsia="{FONT}" w:cs="{FONT}"/>'
    f'<w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr>'))

st = doc.styles["Normal"]
set_style_font(st, FONT, FS)
st.paragraph_format.line_spacing = LS
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Başlık stilleri: H1 = 14pt BÜYÜK (18nk sonra), H2 = 12pt (12nk sonra)
h1s = doc.styles["Heading 1"]
h1s.font.color.rgb = RGBColor(0, 0, 0); h1s.font.bold = True
h1s.paragraph_format.line_spacing = LS
h1s.paragraph_format.space_before = Pt(0); h1s.paragraph_format.space_after = Pt(18)
set_style_font(h1s, FONT, Pt(14))
for lvl in (2, 3):
    hs = doc.styles[f"Heading {lvl}"]
    hs.font.color.rgb = RGBColor(0, 0, 0); hs.font.bold = True
    hs.paragraph_format.line_spacing = LS
    hs.paragraph_format.space_before = Pt(6); hs.paragraph_format.space_after = Pt(12)
    set_style_font(hs, FONT, Pt(12))

try:
    lb = doc.styles["List Bullet"]; set_style_font(lb, FONT, FS); lb.paragraph_format.line_spacing = LS
except Exception:
    pass

def enable_update_fields():
    uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true")
    doc.settings.element.append(uf)

# ───────────────────────── İçerik yardımcıları ─────────────────────────
def tr_upper(s):
    return s.replace("ı", "I").replace("i", "İ").upper()

def h1(text):
    doc.add_heading(tr_upper(text), level=1)

def h2(text):
    doc.add_heading(text, level=2)

def p(text="", bold=False, italic=False, size=None, align=None, after=Pt(6), indent=True):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.line_spacing = LS
    para.paragraph_format.space_after = after
    if indent and align is None:
        para.paragraph_format.first_line_indent = INDENT
    r = para.add_run(text)
    set_run_font(r); r.font.size = size or FS
    r.bold = bold; r.italic = italic
    return para

def bullet(text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.line_spacing = LS
    para.paragraph_format.left_indent = Cm(1.0 + level * 0.6)
    para.clear()
    r = para.add_run(text); set_run_font(r); r.font.size = FS
    return para

def code_block(code):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_before = Pt(4); para.paragraph_format.space_after = Pt(4)
    para._p.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>'))
    r = para.add_run(code); set_run_font(r, "Consolas"); r.font.size = Pt(9)
    return para

def table(headers, rows, widths=None, fs=10):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, hh in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        pp = c.paragraphs[0]; pp.paragraph_format.line_spacing = 1.0; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(hh); r.bold = True; set_run_font(r); r.font.size = Pt(fs); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F497D"/>'))
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            pp = c.paragraphs[0]; pp.paragraph_format.line_spacing = 1.0
            r = pp.add_run(str(val)); set_run_font(r); r.font.size = Pt(fs)
            if ri % 2 == 0:
                c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EDF3"/>'))
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def add_image(path, width=Inches(5.6)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    print("  [UYARI] görsel yok:", path)
    return False

def figure(path, number, caption, width=Inches(5.6)):
    add_image(path, width)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_before = Pt(12)  # 2 tam aralık üstten
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(f"Şekil {number}: {caption}"); set_run_font(r); r.font.size = Pt(10); r.bold = True

def caption_table(number, caption):
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_before = Pt(12); cap.paragraph_format.space_after = Pt(12)  # 2 tam aralık alta
    r = cap.add_run(f"Çizelge {number}: {caption}"); set_run_font(r); r.font.size = Pt(10); r.bold = True

def page_break():
    doc.add_page_break()

def add_toc_field():
    para = doc.add_paragraph(); para.add_run()
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    rr = OxmlElement("w:r"); tt = OxmlElement("w:t")
    tt.text = "İçindekiler tablosu için Word'de bu alanı seçip F9 (Alanları Güncelle) tuşuna basın."
    rr.append(tt); fld.append(rr); para._p.append(fld)

def _page_field_run(paragraph):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(i); run._r.append(e)
    set_run_font(run); run.font.size = Pt(11)

def setup_section(section, number=True, fmt=None, start=None):
    section.page_width, section.page_height = A4_W, A4_H
    section.top_margin = section.bottom_margin = Cm(2.5)
    section.left_margin = section.right_margin = Cm(2.5)
    sectPr = section._sectPr
    for el in sectPr.findall(qn("w:pgNumType")):
        sectPr.remove(el)
    if fmt or start is not None:
        pg = OxmlElement("w:pgNumType")
        if fmt:
            pg.set(qn("w:fmt"), fmt)
        if start is not None:
            pg.set(qn("w:start"), str(start))
        # pgNumType, şema sırasına göre lnNumType'tan sonra / cols'tan önce gelmeli
        anchor = sectPr.find(qn("w:cols"))
        if anchor is None:
            anchor = sectPr.find(qn("w:docGrid"))
        if anchor is not None:
            anchor.addprevious(pg)
        else:
            pgMar = sectPr.find(qn("w:pgMar"))
            if pgMar is not None:
                pgMar.addnext(pg)
            else:
                sectPr.append(pg)
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(fp.runs):
        r._r.getparent().remove(r._r)
    if number:
        _page_field_run(fp)

D = lambda f: os.path.join(DIAGRAMS, f)
S = lambda f: os.path.join(SHOTS, f)

# ═══════════════════════════════════════════════════════════
# Şekil / Çizelge listeleri
# ═══════════════════════════════════════════════════════════
FIG_LIST = [
    ("6.1",  "Swagger (OpenAPI) Canlı API Dokümantasyonu"),
    ("12.1", "Kullanıcı Girişi Sekans Diyagramı (SD-01)"),
    ("12.2", "Kullanıcı Girişi ve Oturum Doğrulama – Detaylı Sekans Diyagramı"),
    ("12.3", "Hasta Kaydı ve Anamnez Güncelleme Sekans Diyagramı (SD-02)"),
    ("12.4", "Randevu Oluşturma ve Çakışma Kontrolü Sekans Diyagramı (SD-03)"),
    ("12.5", "Günlük Dashboard Özet Akışı Sekans Diyagramı (SD-04)"),
    ("12.6", "Tedavi ve Ödeme Girişi Sekans Diyagramı (SD-05)"),
    ("12.7", "Klinik Ayarları Güncelleme Sekans Diyagramı (SD-06)"),
    ("12.8", "Domain (Varlık) Katmanı Sınıf Diyagramı"),
    ("12.9", "Uygulama Katmanı (Boundary/Control) Sınıf Diyagramı"),
    ("17.1", "Genel Tanıtım (SaaS Landing) Sayfası"),
    ("17.2", "Klinik Giriş (Login) Ekranı"),
    ("17.3", "Ana Panel (Dashboard)"),
    ("17.4", "Randevu Takvimi (Haftalık Görünüm)"),
    ("17.5", "Hasta Yönetimi / Hasta Listesi"),
    ("17.6", "Hasta Profili Ekranı"),
    ("17.7", "Dijital Diş Şeması (Odontogram)"),
    ("17.8", "Hasta Tedavi ve Randevu Geçmişi"),
    ("17.9", "Tedavi Türleri Yönetimi"),
    ("17.10","Klinik Ayarları"),
]
TBL_LIST = [
    ("1.1",  "Kullanıcı Rolleri ve Yetkileri"),
    ("3.1",  "Teknoloji Yığını Özeti"),
    ("3.2",  "Backend Çatı Karar Matrisi"),
    ("3.3",  "Frontend Çatı Karar Matrisi"),
    ("3.4",  "Veritabanı Karar Matrisi"),
    ("4.1",  "Ekip Yapısı ve Görev Dağılımı"),
    ("8.1",  "PEP 8 / İsimlendirme Standartları"),
    ("8.2",  "REST API Uç Noktaları Özeti"),
    ("10.1", "Katman Bazlı Test Metrikleri"),
    ("10.2", "Birim Test Durum Tablosu (Temsilî)"),
    ("13.1", "Veri Modeli Varlıkları ve İlişkileri"),
    ("14.1", "GitHub Actions CI/CD İş Akışları"),
    ("18.1", "Use-Case (Kullanım Durumu) Özeti"),
    ("21.1", "Takım Üyeleri Proje Değerlendirmesi"),
    ("22.1", "Yapay Zeka Araçlarının Kullanım Alanları"),
]

# ╔══════════════════════════════════════════════════════════╗
# ║  BÖLÜM A — KAPAK / ONAY / BEYAN / ÖNSÖZ / İÇİNDEKİLER     ║
# ║  (Sayfa numarası YOK)                                     ║
# ╚══════════════════════════════════════════════════════════╝
def C(text, bold=False, italic=False, size=None, after=Pt(4)):
    p(text, bold=bold, italic=italic, size=size, align=WD_ALIGN_PARAGRAPH.CENTER, after=after, indent=False)

# --- KAPAK ---
C("T.C.", bold=True, size=Pt(14))
C("İSTANBUL SAĞLIK VE TEKNOLOJİ ÜNİVERSİTESİ", bold=True, size=Pt(14))
C("MÜHENDİSLİK VE DOĞA BİLİMLERİ FAKÜLTESİ", bold=True, size=Pt(14))
for _ in range(3):
    doc.add_paragraph()
C("LİSANS BİTİRME PROJESİ", bold=True, size=Pt(14))
for _ in range(2):
    doc.add_paragraph()
C("YAŞCA DİŞ KLİNİĞİ YÖNETİM SİSTEMİ", bold=True, size=Pt(18))
C("Çok-Kiracılı (Multi-Tenant) SaaS Tabanlı Klinik Yönetim Uygulaması", italic=True, size=Pt(13))
for _ in range(3):
    doc.add_paragraph()
C("Hazırlayanlar", bold=True, size=Pt(12))
for m in ["Yaman Halloum  –  Öğrenci No: ………………",
          "Ali Üre  –  Öğrenci No: ………………",
          "Cihan Kurtbey  –  Öğrenci No: ………………",
          "Şükrü Yeşilmen  –  Öğrenci No: ………………"]:
    C(m, size=Pt(12), after=Pt(2))
for _ in range(3):
    doc.add_paragraph()
C("Danışman: Dr. Öğr. Üyesi ………………………", size=Pt(12))
C("Yazılım Mühendisliği Anabilim Dalı", size=Pt(12))
for _ in range(3):
    doc.add_paragraph()
C("İSTANBUL – 2026", bold=True, size=Pt(12))
page_break()

# --- ONAY SAYFASI ---
C("T.C.", bold=True, size=Pt(13))
C("İSTANBUL SAĞLIK VE TEKNOLOJİ ÜNİVERSİTESİ", bold=True, size=Pt(13))
C("MÜHENDİSLİK VE DOĞA BİLİMLERİ FAKÜLTESİ", bold=True, size=Pt(13))
C("BİTİRME PROJESİ ONAYI", bold=True, size=Pt(13))
doc.add_paragraph()
p("Hazırlayan: Yaman Halloum, Ali Üre, Cihan Kurtbey, Şükrü Yeşilmen (Grup 4)", indent=False)
p("Proje Başlığı: Yaşca Diş Kliniği Yönetim Sistemi", indent=False)
p("Sınav Tarihi: …… / …… / 2026", indent=False)
doc.add_paragraph()
p("Danışman ve bitirme projesi sınav jürisi değerlendirmesi sonucu yukarıda bilgileri verilen "
  "öğrenciler tarafından hazırlanan proje …………………… bulunmuştur.", indent=False)
doc.add_paragraph(); doc.add_paragraph()
p("Danışman: Dr. Öğr. Üyesi ………………………………                         İmza: ……………………", indent=False)
doc.add_paragraph()
p("Jüri Üyesi: ………………………………                                            İmza: ……………………", indent=False)
p("Jüri Üyesi: ………………………………                                            İmza: ……………………", indent=False)
p("Jüri Üyesi: ………………………………                                            İmza: ……………………", indent=False)
page_break()

# --- BEYAN ---
h1("Beyan")
p("Bu projenin bana ait olduğunu, tüm aşamalarında etik dışı davranışımın olmadığını, içinde yer "
  "alan bütün bilgileri akademik ve etik kurallar içinde elde ettiğimi, kullanmış olduğum bütün "
  "bilgilere kaynak gösterdiğimi ve bu kaynakları da kaynaklar listesine aldığımı, yine bu "
  "projenin tüm aşamalarında patent ve telif haklarını ihlal edici bir davranışımın olmadığını "
  "bildiririm.")
doc.add_paragraph(); doc.add_paragraph()
p("Yaman Halloum (Ad, Soyad; İmza): ……………………", indent=False)
p("Ali Üre (Ad, Soyad; İmza): ……………………", indent=False)
p("Cihan Kurtbey (Ad, Soyad; İmza): ……………………", indent=False)
p("Şükrü Yeşilmen (Ad, Soyad; İmza): ……………………", indent=False)
page_break()

# --- ÖNSÖZ ---
h1("Önsöz")
p("Bu çalışma, diş hekimliği kliniklerinin günlük operasyonel süreçlerini dijitalleştiren, "
  "çok-kiracılı (multi-tenant) bir SaaS uygulaması olan Yaşca Diş Kliniği Yönetim Sistemi’nin "
  "tasarım, gerçekleme, test ve dağıtım aşamalarını kapsamaktadır. Proje; yazılım mühendisliğinin "
  "analizden bakıma uzanan yaşam döngüsünü uçtan uca deneyimleme fırsatı sunmuştur.")
p("Projenin her aşamasında bilgi ve yönlendirmeleriyle bize destek olan danışmanımız Dr. Öğr. "
  "Üyesi ……………………’ye, çalışma boyunca özveriyle katkı sunan tüm ekip arkadaşlarımıza ve "
  "bizleri her zaman destekleyen ailelerimize teşekkür ederiz.")
page_break()

# --- İÇİNDEKİLER ---
h1("İçindekiler")
add_toc_field()
p("(Not: Word'de Ctrl+A ile tümünü seçip F9 tuşuna basarak içindekiler tablosunu ve sayfa "
  "numaralarını otomatik güncelleyebilirsiniz.)", italic=True, size=Pt(10), indent=False)

# ╔══════════════════════════════════════════════════════════╗
# ║  BÖLÜM B — ÖN SAYFALAR (Romen: i, ii, iii ...)            ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_section(WD_SECTION.NEW_PAGE)

# --- KISALTMALAR ---
h1("Kısaltmalar")
table(["Kısaltma", "Açıklama"],
      [("API", "Application Programming Interface (Uygulama Programlama Arayüzü)"),
       ("CI/CD", "Continuous Integration / Continuous Deployment"),
       ("CRUD", "Create-Read-Update-Delete"),
       ("DRF", "Django REST Framework"),
       ("E2E", "End-to-End (Uçtan Uca) Test"),
       ("ER", "Entity-Relationship (Varlık-İlişki) Diyagramı"),
       ("FDI", "Fédération Dentaire Internationale (Diş Numaralandırma Sistemi)"),
       ("JWT", "JSON Web Token"),
       ("KVKK", "Kişisel Verilerin Korunması Kanunu"),
       ("ORM", "Object-Relational Mapping"),
       ("PEP", "Python Enhancement Proposal"),
       ("RBAC", "Role-Based Access Control (Rol Bazlı Erişim Kontrolü)"),
       ("REST", "Representational State Transfer"),
       ("SaaS", "Software as a Service (Hizmet Olarak Yazılım)"),
       ("UML", "Unified Modeling Language"),
       ("WCAG", "Web Content Accessibility Guidelines")],
      [3, 13])
page_break()

# --- ŞEKİL LİSTESİ ---
h1("Şekil Listesi")
for num, cap in FIG_LIST:
    para = doc.add_paragraph(); para.paragraph_format.line_spacing = 1.0
    r = para.add_run(f"Şekil {num}: {cap}"); set_run_font(r); r.font.size = Pt(11)
page_break()

# --- ÇİZELGE LİSTESİ ---
h1("Çizelge Listesi")
for num, cap in TBL_LIST:
    para = doc.add_paragraph(); para.paragraph_format.line_spacing = 1.0
    r = para.add_run(f"Çizelge {num}: {cap}"); set_run_font(r); r.font.size = Pt(11)
page_break()

# --- ÖZET ---
h1("Özet")
C("YAŞCA DİŞ KLİNİĞİ YÖNETİM SİSTEMİ", bold=True)
C("Grup 4 – Yazılım Mühendisliği Bölümü, Lisans Bitirme Projesi")
doc.add_paragraph()
p("Yaşca; diş hekimliği kliniklerinin hasta kaydı, anamnez, çakışma kontrollü randevu, tedavi "
  "takibi, ödeme yönetimi ve dijital diş şeması (odontogram) süreçlerini tek bir güvenli "
  "platformda toplayan, çok-kiracılı (multi-tenant) bir SaaS web uygulamasıdır. Arka uç Python "
  "3.12, Django REST Framework ve django-tenants ile; ön uç React (TypeScript), Vite ve Tailwind "
  "CSS ile; veritabanı katmanı ise şema bazlı izolasyon sağlayan PostgreSQL ile geliştirilmiştir. "
  "Her klinik, kendi izole PostgreSQL şemasına sahip bir kiracı olarak modellenmiş; kiracı "
  "çözümlemesi X-Tenant başlığı ve özel bir ara katman (middleware) ile yapılmıştır. Kimlik "
  "doğrulama JWT, yetkilendirme ise rol bazlı erişim kontrolü (RBAC) ile sağlanmaktadır. Kalite "
  "güvencesi için GitHub Actions üzerinde çalışan, toplam 403 otomatik test (151 ön yüz birim, "
  "169 arka uç, 83 uçtan uca) içeren çok katmanlı bir test altyapısı ve WCAG 2.1 AA "
  "erişilebilirlik geçidi kurulmuştur. Uygulama; ön yüzü Vercel, arka ucu Render üzerinde olmak "
  "üzere canlı ortamda yayındadır. Bu rapor, projenin mimarisini, teknolojilerini, geliştirme ve "
  "test süreçlerini, kalite pratiklerini ve yapay zeka destekli araçların geliştirme sürecindeki "
  "kullanımını belgelemektedir.")
doc.add_paragraph()
p("Anahtar Kelimeler: Klinik Yönetim Sistemi, Çok-Kiracılı SaaS, Django REST Framework, React, "
  "PostgreSQL", bold=True)
page_break()

# --- ABSTRACT ---
h1("Abstract")
C("YAŞCA DENTAL CLINIC MANAGEMENT SYSTEM", bold=True)
C("Group 4 – Department of Software Engineering, Graduation Project")
doc.add_paragraph()
p("Yaşca is a multi-tenant Software-as-a-Service (SaaS) web application that consolidates the "
  "patient registration, anamnesis, conflict-checked appointment scheduling, treatment tracking, "
  "payment management and digital dental chart (odontogram) processes of dental clinics into a "
  "single secure platform. The backend is built with Python 3.12, Django REST Framework and "
  "django-tenants; the frontend with React (TypeScript), Vite and Tailwind CSS; and the data "
  "layer with PostgreSQL, which provides schema-based isolation. Each clinic is modeled as a "
  "tenant owning its own isolated PostgreSQL schema, and tenant resolution is performed via an "
  "X-Tenant header and a dedicated middleware. Authentication is handled with JWT and "
  "authorization with role-based access control (RBAC). For quality assurance, a multi-layered "
  "test infrastructure of 403 automated tests (151 frontend unit, 169 backend, 83 end-to-end) "
  "running on GitHub Actions, together with a WCAG 2.1 AA accessibility gate, has been "
  "established. The application is deployed live, with the frontend on Vercel and the backend on "
  "Render. This report documents the system's architecture, technologies, development and testing "
  "processes, quality practices, and the use of AI-assisted tools throughout development.")
doc.add_paragraph()
p("Keywords: Clinic Management System, Multi-Tenant SaaS, Django REST Framework, React, "
  "PostgreSQL", bold=True)

# ╔══════════════════════════════════════════════════════════╗
# ║  BÖLÜM C — ANA METİN (Arap: 1, 2, 3 ...)                  ║
# ╚══════════════════════════════════════════════════════════╝
doc.add_section(WD_SECTION.NEW_PAGE)

# ───── 1. GİRİŞ ─────
h1("1. Giriş ve Tanım")
p("Sağlık sektöründe klinik operasyonlarının büyük bölümü hâlâ kâğıt tabanlı veya birbirinden "
  "kopuk yazılımlarla yürütülmektedir. Bu durum; veri kaybı, randevu çakışması, finansal takip "
  "zorlukları ve kişisel sağlık verilerinin güvenliği açısından ciddi riskler oluşturur. Yaşca "
  "Diş Kliniği Yönetim Sistemi, diş kliniklerinin günlük operasyonel iş akışlarını "
  "dijitalleştirmek, veri güvenliğini sağlamak ve klinik verimliliğini artırmak amacıyla "
  "geliştirilmiş, açık kaynaklı (FOSS) ve modern bir web uygulamasıdır.")
p("Sistem; hasta kaydı ve anamnez (tıbbi özgeçmiş) yönetimi, çakışma kontrollü randevu takvimi, "
  "tedavi ve finansal takip, dijital diş şeması (odontogram), tedavi türü tanımlama, klinik "
  "ayarları, kullanıcı/personel yönetimi ve işlem geçmişi (audit log) modüllerini tek bir çatı "
  "altında sunar. Her kliniğin verileri, PostgreSQL şema tabanlı çok-kiracılı mimari ile "
  "fiziksel olarak izole edilir.")
p("Projenin Amacı ve Kapsamı", bold=True)
p("Projenin amacı, klinik personelinin (Yönetici, Hekim, Asistan) günlük işlerini hızlandıran, "
  "güvenli ve ölçeklenebilir bir SaaS platformu ortaya koymaktır. Sistemin kapsamı klinik "
  "personeli ile sınırlandırılmıştır; hastalar sistemin kullanıcısı değil, sistemde verisi "
  "tutulan birer varlıktır. Mevcut sürüm gerçek ödeme altyapısı (sanal POS) ve devlet sağlık "
  "sistemleri (MEDULA vb.) entegrasyonu içermemektedir. Çizelge 1.1’de tanımlı kullanıcı "
  "rolleri ve yetkileri özetlenmiştir.")
caption_table("1.1", "Kullanıcı Rolleri ve Yetkileri")
table(["Rol", "Açıklama", "Erişim ve Yetkiler"],
      [("Yönetici (Admin)", "Başhekim / klinik sahibi",
        "Klinik ayarları, personel yönetimi, tüm verilere tam erişim, kullanıcı ekleme ve şifre sıfırlama, işlem geçmişi"),
       ("Hekim (Doktor)", "Sistemin ana kullanıcısı",
        "Kendi hastalarının randevuları, tedavi girişi, odontogram, dashboard"),
       ("Asistan / Sekreter", "İdari süreçler",
        "Hasta kaydı, randevu oluşturma, ödeme girişi")],
      [3.5, 4, 9])
page_break()

# ───── 2. PROJE YAKLAŞIMI ─────
h1("2. Proje Yaklaşımı ve Mimari Bağlam")
p("Yaşca, istemci–sunucu (client–server) ayrımına dayalı, API odaklı (API-first) bir web "
  "uygulamasıdır. Ön yüz (React) ile arka uç (Django REST Framework) tamamen JSON tabanlı bir "
  "REST API üzerinden haberleşir. Sistem, çok-kiracılı (multi-tenant) bir SaaS olarak "
  "tasarlanmış; her klinik kendi izole PostgreSQL şemasına sahip bir kiracı (tenant) olarak "
  "modellenmiştir.")
h2("2.1 Mimari Yaklaşımı (Monorepo ve Katmanlı Mimari)")
p("Proje, tek bir kod deposunda hem ön yüz hem de arka uç kodunu barındıran monorepo yapısıyla "
  "yönetilmiştir (frontend/ + backend/ + docs/). Bu yaklaşım; ortak sürüm yönetimi, atomik "
  "commit’ler ve API sözleşmelerinin tek yerde tutulması açısından küçük ekipler için avantaj "
  "sağlamıştır.")
p("Sistem, Boundary–Control–Entity (Sınır–Denetim–Varlık) ilkesine dayalı katmanlı bir "
  "mimariyle tasarlanmıştır:")
bullet("Boundary (Arayüz): React bileşenleri (LoginPage, Dashboard, AppointmentCalendar, "
       "PatientProfile, DentalChart vb.).")
bullet("Control (Denetim/İş Mantığı): Ön yüzde merkezî api.ts servis katmanı; arka uçta DRF "
       "ViewSet ve Serializer sınıfları.")
bullet("Entity (Varlık): Django ORM modelleri (Patient, Appointment, Treatment, Payment vb.).")
p("Mimarinin temel kuralı, Boundary katmanının Entity ile doğrudan konuşmaması; her erişimin "
  "Control katmanı üzerinden geçmesidir. Bu kural, Django’nun MVT mimarisiyle örtüşür ve test "
  "edilebilirliği güvence altına alır.")
page_break()

# ───── 3. TEKNOLOJİLER ─────
h1("3. Kullanılan Ana Teknolojiler")
p("Teknoloji yığını; performans, ekip yetkinliği, güvenlik ve nesne tabanlı tasarıma uygunluk "
  "kriterleriyle seçilmiştir. Çizelge 3.1 yığının tamamını özetler; sonrasında ana bileşenler "
  "ayrıntılandırılmıştır.")
caption_table("3.1", "Teknoloji Yığını Özeti")
table(["Katman", "Teknolojiler"],
      [("Ön Yüz (Frontend)", "React 18, TypeScript, Vite, Tailwind CSS v4, Radix UI / MUI, "
        "lucide-react, react-router, react-hook-form, recharts, react-dnd, i18next (TR/EN)"),
       ("Arka Uç (Backend)", "Python 3.12, Django 5.2, Django REST Framework, simplejwt (JWT), "
        "django-tenants, drf-spectacular (Swagger), django-filter, django-anymail[resend], gunicorn"),
       ("Veritabanı", "PostgreSQL (schema-per-tenant izolasyon)"),
       ("Test", "Vitest, Testing Library, MSW, Playwright, vitest-axe, Stryker, pytest, "
        "factory-boy + Faker, Pact"),
       ("DevOps / Dağıtım", "Docker & docker-compose, GitHub Actions, Vercel (ön yüz), "
        "Render (arka uç + PostgreSQL), Codecov")],
      [3.5, 12.5], fs=9)
h2("3.1 Django")
p("Backend, Python tabanlı Django ve Django REST Framework (DRF) ile geliştirilmiştir (Django "
  "Software Foundation, 2026). Django “batteries-included” felsefesiyle ORM, kimlik doğrulama, "
  "yönetim paneli ve güvenlik korumalarını (CSRF, SQL injection vb.) hazır sunar. DRF ise "
  "serileştirme (serializer), görünüm kümeleri (ViewSet) ve izin (permission) sınıflarıyla REST "
  "API geliştirmeyi standartlaştırır (Django REST Framework, 2026). JWT tabanlı kimlik doğrulama "
  "djangorestframework-simplejwt; canlı API dokümantasyonu drf-spectacular ile sağlanmıştır. "
  "Karar matrisinde (Çizelge 3.2) Django REST en yüksek puanı almıştır.")
caption_table("3.2", "Backend Çatı Karar Matrisi")
table(["Kriter", "Django REST", "Express.js", "Spring Boot"],
      [("Geliştirme hızı / öğrenme eğrisi", "5", "4", "2"),
       ("Dahili güvenlik (Auth/RBAC)", "5", "2", "4"),
       ("Ekip yetkinliği", "5", "3", "2"),
       ("OOP / sınıf modeli uyumu", "4", "3", "5"),
       ("TOPLAM", "19", "12", "13")], [6, 3.3, 3.3, 3.3], fs=9)
h2("3.2 React")
p("Ön yüz, React (TypeScript) ile geliştirilmiştir (Meta, 2026). React’in bileşen tabanlı "
  "(component-based) ve sanal DOM (Virtual DOM) yaklaşımı, özellikle odontogram gibi sık "
  "güncellenen, çok sayıda etkileşimli öğe içeren ekranlarda performans avantajı sağlar. "
  "TypeScript ile statik tip denetimi sayesinde hata oranı düşürülmüş (Microsoft, 2026); Vite "
  "ile hızlı geliştirme, Tailwind CSS ile tutarlı bir tasarım dili elde edilmiştir. Karar "
  "matrisi Çizelge 3.3’te sunulmuştur.")
caption_table("3.3", "Frontend Çatı Karar Matrisi")
table(["Kriter", "React", "Vue", "Angular"],
      [("Topluluk / kaynak", "5", "4", "3"),
       ("Odontogram için DOM performansı", "5", "4", "3"),
       ("Ekip yetkinliği", "5", "2", "1"),
       ("Bileşen uyumu", "5", "5", "5"),
       ("TOPLAM", "20", "15", "12")], [6, 3.3, 3.3, 3.3], fs=9)
h2("3.3 Veritabanı")
p("Veritabanı olarak PostgreSQL seçilmiştir (The PostgreSQL Global Development Group, 2026). "
  "ACID uyumluluğu, güçlü ilişkisel bütünlük garantileri ve şema (schema) desteği, çok-kiracılı "
  "mimarinin doğal temelini oluşturur. Her kliniğin verisi ayrı bir PostgreSQL şemasında "
  "tutularak kiracılar arası veri sızıntısı yapısal olarak engellenir. JSON alan desteği "
  "sayesinde esnek yapılar da ilişkisel model içinde saklanabilir. Karar matrisi Çizelge 3.4’te "
  "verilmiştir.")
caption_table("3.4", "Veritabanı Karar Matrisi")
table(["Kriter", "PostgreSQL", "MongoDB", "SQLite"],
      [("Veri bütünlüğü (ACID)", "5", "2", "4"),
       ("İlişki modelleme", "5", "3", "5"),
       ("Ölçeklenebilirlik", "5", "5", "1"),
       ("Medikal veri güvenliği", "5", "3", "2"),
       ("TOPLAM", "20", "13", "12")], [6, 3.3, 3.3, 3.3], fs=9)
page_break()

# ───── 4. EKİP ─────
h1("4. Ekip Yapısı ve Görev Dağılımı")
p("Proje, Grup 4 tarafından dört kişilik bir ekiple yürütülmüştür. Görev dağılımı; üyelerin "
  "yetkinlikleri ve sürüm geçmişindeki katkıları dikkate alınarak yapılmıştır (Çizelge 4.1).")
caption_table("4.1", "Ekip Yapısı ve Görev Dağılımı")
table(["Üye", "Başlıca Sorumluluk Alanı"],
      [("Yaman Halloum", "Backend (Django/DRF), multi-tenant altyapı, model ve API geliştirme"),
       ("Ali Üre", "Tam yığın (full-stack) geliştirme, mimari kurulum, dağıtım (Vercel/Render)"),
       ("Cihan Kurtbey", "Frontend, kimlik doğrulama akışı, şifre sıfırlama, e-posta entegrasyonu"),
       ("Şükrü Yeşilmen", "Test, dokümantasyon ve kalite güvencesi katkıları")],
      [4, 12])
p("Sürüm yönetimi sisteminden (Git) elde edilen verilere göre proje deposunda toplam 119 commit "
  "ve birden çok geliştirme dalı (Ali, Cihan, Yaman, main) bulunmaktadır.")
page_break()

# ───── 5. GELİŞTİRME SÜRECİ ─────
h1("5. Yazılım Geliştirme Süreci")
p("Geliştirme, üniversite ders takvimine paralel olarak artımlı (iteratif-artımlı) bir süreçle "
  "yürütülmüştür. Her adım, bir önceki adımın çıktısını temel alarak ürünü olgunlaştırmıştır:")
bullet("Adım 1 – Analiz ve Tasarım: İster analizi (22 fonksiyonel + 10 fonksiyonel olmayan "
       "gereksinim), sekans ve sınıf diyagramları (UML).")
bullet("Adım 2 – Teknoloji Karar Analizi: Karar matrisleriyle teknoloji yığınının seçimi.")
bullet("Adım 3 – Gerçekleme: İç/dış arayüzler, REST API, multi-tenant altyapı ve ön yüz.")
bullet("Adım 4 – Test: Birim testler, test piramidi ve CI/CD altyapısının kurulması.")
bullet("Adım 5 – Dağıtım ve İyileştirme: Vercel/Render dağıtımı, erişilebilirlik ve performans "
       "iyileştirmeleri, dokümantasyon.")
p("Süreç boyunca her özellik; tasarım → gerçekleme → test → kod gözden geçirme → birleştirme "
  "döngüsünden geçmiş; kalite geçitleri (CI) otomatik olarak uygulanmıştır.")
page_break()

# ───── 6. DOKÜMANTASYON ─────
h1("6. Teknik Dokümantasyon Yönetimi")
p("Proje, kod ile birlikte yaşayan (docs-as-code) bir dokümantasyon stratejisi izlemiştir. "
  "Dokümanlar depo içinde versiyonlanmış ve kod değişiklikleriyle eşzamanlı güncellenmiştir:")
bullet("README.md: kurulum, çalıştırma, demo verisi ve teknoloji özeti.")
bullet("docs/CONTRIBUTING.md: PR süreci ve kod gözden geçirme kontrol listesi.")
bullet("docs/adr/: Mimari Karar Kayıtları (ADR) — test stratejisi, multi-tenant izolasyon, "
       "a11y sıfır tolerans gibi kararların gerekçeleri.")
bullet("docs/TESTING.md, TEST_PYRAMID.md, TEST_METRICS.md: test rehberi ve otomatik üretilen metrikler.")
bullet("docs/uml/: PlantUML kaynaklı sekans ve sınıf diyagramları.")
bullet("Canlı API dokümantasyonu: drf-spectacular ile üretilen OpenAPI şeması ve Swagger arayüzü "
       "(Şekil 6.1).")
figure(SWAGGER, "6.1", "Swagger (OpenAPI) Canlı API Dokümantasyonu")
page_break()

# ───── 7. ÜRÜN YAPILARI ─────
h1("7. Teknoloji ve Ürün Yapıları")
p("Bu bölüm, ürünün ön yüz ve arka uç bileşenlerinin iç yapısını ve güvenlik tasarımını ele alır.")
h2("7.1 Frontend Yapısı")
p("Ön yüz, src/app altında modüler bir yapıya sahiptir: components (ekran ve diyalog "
  "bileşenleri), contexts (AuthContext ile JWT oturum yönetimi), hooks (örn. useClinicNavigate), "
  "services (api.ts) ve locales (i18next çeviri dosyaları, TR/EN). Uygulama yönlendirmesi "
  "react-router ile yapılır; canlı ortamda klinik paneline /app/:slug yol kalıbıyla, lokal "
  "geliştirmede ise alt alan adı (subdomain) ile erişilir. Tüm API istekleri tek bir api.ts "
  "katmanından geçer; bu katman JWT erişim jetonunu Authorization başlığına, kiracı bilgisini "
  "ise X-Tenant başlığına otomatik ekler ve 401 durumunda jetonu sessizce yeniler (refresh).")
code_block(
    "// frontend/src/app/services/api.ts (özet)\n"
    "const getAuthHeaders = () => ({\n"
    "  'Content-Type': 'application/json',\n"
    "  ...(token ? { Authorization: `Bearer ${token}` } : {}),\n"
    "  ...(TENANT_SUBDOMAIN ? { 'X-Tenant': TENANT_SUBDOMAIN } : {}),\n"
    "});")
h2("7.2 Backend, Veritabanı ve Güvenlik Yapısı")
p("Arka uç; api (klinik alan modelleri ve iş mantığı) ve customers (kiracı kaydı/yönetimi) "
  "uygulamalarından oluşur. Güvenlik katmanı çok bileşenlidir:")
bullet("Kimlik Doğrulama: JWT (access + refresh) tabanlı, durumsuz (stateless) oturum.")
bullet("Yetkilendirme: Rol Bazlı Erişim Kontrolü (RBAC) — hekim yalnızca kendi randevu ve "
       "tedavilerini görür; klinik ayarları ve personel yönetimi yalnızca yöneticiye açıktır.")
bullet("Kiracı İzolasyonu: HeaderTenantMiddleware her istekte X-Tenant başlığını (yoksa Host "
       "başlığını) okuyarak isteği ilgili PostgreSQL şemasına yönlendirir.")
bullet("Denetim İzi (Audit Log): AuditLog modeli CREATE/UPDATE/DELETE işlemlerini kullanıcı, IP "
       "ve değişiklik ayrıntılarıyla birlikte kaydeder (KVKK/güvenlik).")
code_block(
    "class HeaderTenantMiddleware:\n"
    "    def __call__(self, request):\n"
    "        tenant = None\n"
    "        header = request.META.get('HTTP_X_TENANT', '').strip()   # 1) X-Tenant\n"
    "        if header:\n"
    "            tenant = TenantModel.objects.filter(schema_name=header).first()\n"
    "        if not tenant:                                            # 2) Host fallback\n"
    "            host = request.get_host().split(':')[0]\n"
    "            d = DomainModel.objects.filter(domain=host).first()\n"
    "            tenant = d.tenant if d else None\n"
    "        if not tenant:                                            # 3) public\n"
    "            tenant = TenantModel.objects.get(schema_name='public')\n"
    "        request.tenant = tenant\n"
    "        connection.set_tenant(tenant)\n"
    "        return self.get_response(request)")
page_break()

# ───── 8. İSİMLENDİRME ─────
h1("8. İsimlendirme Standartları (PEP 8 / PEP 257)")
p("Python tarafında PEP 8 (kod stili) ve PEP 257 (docstring sözleşmeleri) esas alınmıştır (van "
  "Rossum, Warsaw ve Coghlan, 2001; Goodger ve van Rossum, 2001). Tüm model ve servis sınıfları, "
  "amacını açıklayan Türkçe docstring’ler içerir (örn. Appointment modelindeki “Randevu kaydı. "
  "F-006, F-007, F-008, F-009, F-019.”). Ön yüzde ise yerleşik TypeScript/React isimlendirme "
  "kuralları uygulanmıştır (Çizelge 8.1).")
caption_table("8.1", "PEP 8 / İsimlendirme Standartları")
table(["Öğe", "Kural", "Örnek"],
      [("Modül / paket", "snake_case", "api, customers, middleware.py"),
       ("Sınıf", "PascalCase", "AppointmentViewSet, CustomUser"),
       ("Fonksiyon / değişken", "snake_case", "get_queryset, work_start_time"),
       ("Sabit", "UPPER_CASE", "API_BASE, DEFAULT_COUNTRY"),
       ("React bileşeni", "PascalCase.tsx", "PatientProfile.tsx, DentalChart.tsx"),
       ("Değişken/fonksiyon (TS)", "camelCase", "fetchCurrentUser, isAuthenticated"),
       ("Docstring", "PEP 257 (üçlü tırnak)", '"""Hasta kaydı. F-003 ..."""')],
      [4, 5, 7.5], fs=9)
h2("8.1 API İsimlendirme Standartları")
p("REST API uç noktaları; kaynak odaklı, çoğul ve küçük harfli isimlerle tasarlanmıştır. HTTP "
  "fiilleri (GET/POST/PUT/PATCH/DELETE) eylemi belirtir; eylem isimleri URL’ye gömülmez "
  "(Çizelge 8.2).")
caption_table("8.2", "REST API Uç Noktaları Özeti")
table(["Uç Nokta", "Metot", "Açıklama"],
      [("/api/auth/token/", "POST", "JWT ile giriş (access + refresh)"),
       ("/api/auth/token/refresh/", "POST", "Erişim jetonu yenileme"),
       ("/api/auth/me/", "GET", "Oturum açan kullanıcı profili"),
       ("/api/patients/", "GET / POST", "Hasta listeleme / oluşturma"),
       ("/api/patients/{id}/", "GET / PUT / DELETE", "Hasta detay / güncelle / pasifleştir"),
       ("/api/appointments/", "GET / POST", "Randevu (çakışma kontrollü)"),
       ("/api/treatments/", "GET / POST", "Tedavi kaydı"),
       ("/api/payments/", "GET / POST", "Ödeme kaydı"),
       ("/api/treatment-types/", "GET / POST", "Tedavi türü yönetimi"),
       ("/api/dashboard/today/", "GET", "Günlük özet (randevu/hasta sayıları)"),
       ("/api/schema/ , /api/docs/", "GET", "OpenAPI şeması ve Swagger arayüzü")],
      [5.5, 3.5, 7], fs=9)
page_break()

# ───── 9. TEST DATASI ─────
h1("9. Test Datası Üretim Süreci")
p("Testlerin gerçekçi ve tekrarlanabilir olması için sahte (sentetik) veri üretimi "
  "otomatikleştirilmiştir. Arka uçta factory-boy ve Faker kütüphaneleriyle model fabrikaları "
  "tanımlanmış; böylece her test kendi izole verisini üretebilmektedir. Ayrıca seed_demo_data "
  "yönetim komutu, demo amacıyla public kiracıyı, ayrı klinik kiracılarını ve her klinik için "
  "personel, rastgele hasta, tedavi ve ödeme kayıtlarını tek komutla oluşturur. Birim testlerde "
  "dış bağımlılıklar (django-tenants, PostgreSQL) unittest.mock / monkeypatch ile taklit "
  "edilerek metotlar gerçekten izole edilmiştir.")
code_block(
    "# Demo verisini tek komutla üretme (docker-compose)\n"
    "docker-compose run --rm backend sh -c \\\n"
    "  \"python manage.py flush --no-input && python manage.py seed_demo_data\"")
page_break()

# ───── 10. ÖLÇÜMLEME ─────
h1("10. Ölçümleme ve Süreç Metrikleri")
p("Projenin test altyapısı, katmanlara dağılmış toplam 403 otomatik test içermektedir. "
  "Çizelge 10.1, test-metrics aracıyla otomatik üretilen katman bazlı dağılımı gösterir.")
caption_table("10.1", "Katman Bazlı Test Metrikleri")
table(["Katman", "Test Sayısı", "Dağılım (%)"],
      [("Ön Yüz Birim (Vitest)", "151", "37,5"),
       ("Arka Uç (pytest)", "169", "41,9"),
       ("Uçtan Uca (Playwright)", "83", "20,6"),
       ("TOPLAM", "403", "100")],
      [6, 4, 4])
p("Bunlara ek olarak; mutasyon testi (Stryker), sözleşme testi (Pact), görsel regresyon, "
  "Lighthouse performans denetimi ve yük testi ayrı CI iş akışlarıyla yürütülmektedir. Kod "
  "kapsamı Codecov ile izlenmekte; arka uçta %78 asgari kapsam (coverage) geçidi "
  "(--cov-fail-under=78) uygulanmaktadır.")
p("Birim Test Sonuçları", bold=True)
p("Sınıf diyagramından seçilen kritik sınıfların çekirdek metotları, pozitif ve negatif "
  "senaryolarla birim test edilmiş; ilgili test dosyalarının tamamı 54/54 PASSED ile "
  "sonuçlanmıştır. Çizelge 10.2’de temsilî bir alt küme sunulmuştur.")
caption_table("10.2", "Birim Test Durum Tablosu (Temsilî)")
table(["Test ID", "Hedef Sınıf", "Tür", "Senaryo", "Durum"],
      [("UT-001", "AppointmentCreateSerializer", "Negatif", "Aynı hekime aynı saatte ikinci randevu", "PASSED"),
       ("UT-004", "AppointmentCreateSerializer", "Pozitif", "İptal edilmiş randevunun slotu bloke etmemesi", "PASSED"),
       ("UT-006", "HeaderTenantMiddleware", "Pozitif", "X-Tenant başlığı ile doğru kiracı çözümleme", "PASSED"),
       ("UT-010", "HeaderTenantMiddleware", "Negatif", "Boş X-Tenant → public kiracıya düşme", "PASSED"),
       ("UT-014", "RegisterClinicView", "Negatif", "Zorunlu alan eksik → HTTP 400", "PASSED"),
       ("UT-020", "PatientSerializer", "Pozitif", "Nested anamnez ile hasta oluşturma", "PASSED"),
       ("UT-025", "TreatmentSerializer", "Negatif", "Aynı gün/diş/tür mükerrer → ValidationError", "PASSED")],
      [2.2, 4.5, 1.8, 5.5, 2], fs=9)
page_break()

# ───── 11. İŞ BİRLİĞİ ─────
h1("11. İş Birliğine Dayalı Geliştirme")
p("Geliştirme; özellik dalları (feature branch) ve çekme istekleri (Pull Request) üzerinden "
  "yürütülen, iş birliğine dayalı bir akışla yönetilmiştir (Conventional Commits, 2023). Süreç "
  "kuralları şunlardır:")
bullet("Branch isimlendirme: feat/…, fix/…, test/…, docs/… ön ekleri.")
bullet("Commit mesajları: Conventional Commits standardı (feat:, fix:, test:, docs:, refactor:, chore:).")
bullet("Her PR’da: en az 1 onay (code review), CI’nin yeşil olması (SQLite + PostgreSQL matrisi, "
       "lint, a11y) ve PR şablonunun doldurulması.")
bullet("Birleştirme (merge): Squash & merge; final commit Conventional Commits biçiminde.")
bullet("Yerel kalite geçidi: husky + lint-staged ile commit öncesi (pre-commit) otomatik lint.")
p("Kod gözden geçirme için ayrıntılı bir kontrol listesi (test, kod kalitesi, güvenlik, "
  "erişilebilirlik) tanımlanmış; --no-verify ile geçit atlama, coverage eşiğini düşürme ve "
  "gerekçesiz test atlama gibi pratikler açıkça yasaklanmıştır.")
page_break()

# ───── 12. UML ─────
h1("12. Mimari Çizim ve Katmanlı Yapı")
p("Sistemin dinamik davranışı sekans diyagramlarıyla, statik yapısı ise sınıf diyagramlarıyla "
  "modellenmiştir. Sekans diyagramlarında arayüzden gönderilen mesajlar (login(), "
  "createAppointment(), fetchDashboardToday() vb.) doğrudan api.ts servis katmanının "
  "metotlarıyla eşleşecek biçimde, “sınıf–sekans tutarlılığı” ilkesiyle hazırlanmıştır.")
p("12.1 Sekans Diyagramları (Dinamik Model)", bold=True)
figure(D("WhatsApp Image 2026-04-25 at 17.03.46.jpeg"), "12.1", "Kullanıcı Girişi Sekans Diyagramı (SD-01)")
figure(D("WhatsApp Image 2026-04-25 at 17.02.41.jpeg"), "12.2", "Kullanıcı Girişi ve Oturum Doğrulama – Detaylı Sekans Diyagramı")
figure(D("WhatsApp Image 2026-04-25 at 17.24.32.jpeg"), "12.3", "Hasta Kaydı ve Anamnez Güncelleme Sekans Diyagramı (SD-02)")
figure(D("WhatsApp Image 2026-04-25 at 17.06.03.jpeg"), "12.4", "Randevu Oluşturma ve Çakışma Kontrolü Sekans Diyagramı (SD-03)")
figure(D("WhatsApp Image 2026-04-25 at 17.23.21.jpeg"), "12.5", "Günlük Dashboard Özet Akışı Sekans Diyagramı (SD-04)")
figure(D("WhatsApp Image 2026-04-25 at 17.00.59.jpeg"), "12.6", "Tedavi ve Ödeme Girişi Sekans Diyagramı (SD-05)")
figure(D("WhatsApp Image 2026-04-25 at 17.25.28.jpeg"), "12.7", "Klinik Ayarları Güncelleme Sekans Diyagramı (SD-06)")
p("12.2 Sınıf Diyagramları (Statik Model)", bold=True)
p("Sekans diyagramlarındaki tüm arayüz, servis ve varlık nesneleri iki ana sınıf diyagramında "
  "toplanmıştır. Domain katmanı veritabanı varlıklarını ve ilişkilerini; uygulama katmanı ise "
  "Boundary ve Control bileşenlerini gösterir.")
figure(D("WhatsApp Image 2026-04-25 at 17.27.36.jpeg"), "12.8", "Domain (Varlık) Katmanı Sınıf Diyagramı")
figure(D("WhatsApp Image 2026-04-25 at 17.28.41.jpeg"), "12.9", "Uygulama Katmanı (Boundary/Control) Sınıf Diyagramı")
page_break()

# ───── 13. VERİ MODELİ ─────
h1("13. Veri Modeli")
p("Veri modeli, public (paylaşılan) şema ve kiracı (tenant) şemaları olmak üzere iki düzeyde "
  "tasarlanmıştır. Public şema kiracı yönetimini (client/tenant ve domain tabloları); her "
  "kiracı şeması ise kliniğe ait alan verilerini barındırır. Çizelge 13.1’de kiracı şemasındaki "
  "temel varlıklar ve ilişkileri özetlenmiştir.")
caption_table("13.1", "Veri Modeli Varlıkları ve İlişkileri")
table(["Varlık (Model)", "Açıklama", "Temel İlişkiler"],
      [("CustomUser", "Klinik personeli (admin/doctor/assistant)", "1:N Appointment, Treatment, TreatmentType"),
       ("Patient", "Hasta kaydı (ad, soyad, telefon zorunlu)", "1:1 Anamnesis; 1:N Appointment/Treatment/Payment/Document"),
       ("Anamnesis", "Tıbbi özgeçmiş", "1:1 Patient (CASCADE)"),
       ("TreatmentType", "Tedavi türü ve varsayılan fiyat", "1:N Treatment"),
       ("Appointment", "Çakışma kontrollü randevu", "N:1 Patient, Doctor; opsiyonel Treatment"),
       ("Treatment", "Uygulanan/planlanan tedavi (diş no, FDI)", "N:1 Patient, Doctor, TreatmentType; 1:N Payment"),
       ("Payment", "Ödeme kaydı (tutar, tarih)", "N:1 Patient; opsiyonel Treatment"),
       ("Document", "Hasta dosyaları (röntgen vb.)", "N:1 Patient, uploaded_by"),
       ("AuditLog", "İşlem geçmişi (CREATE/UPDATE/DELETE)", "N:1 User; GenericForeignKey ile herhangi bir varlık"),
       ("ClinicSettings", "Çalışma saatleri ve klinik ayarları", "Tekil yapılandırma kaydı")],
      [3.5, 6.5, 6], fs=9)
p("Randevu çakışması, aynı hekime aynı tarih ve saatte yalnızca “planlanmış” durumdaki bir "
  "randevunun verilebilmesi kuralıyla serializer doğrulama (validate) katmanında engellenir.")
page_break()

# ───── 14. TEST VE DEPLOYMENT ─────
h1("14. Test ve Deployment Süreçleri")
h2("14.1 Test Süreci ve Uygulamaları")
p("Test stratejisi, test piramidi yaklaşımına dayanır: tabanda hızlı ve çok sayıda birim test, "
  "ortada entegrasyon testleri, tepede ise az sayıda uçtan uca (E2E) test. Testler GitHub "
  "Actions üzerinde, hem hızlı geri bildirim için SQLite hem de gerçek kiracı izolasyonunu "
  "doğrulamak için PostgreSQL olmak üzere çift modlu (matrix) çalıştırılır. Ek olarak görsel "
  "regresyon, sözleşme (Pact), mutasyon (Stryker), Lighthouse ve yük testleri ayrı iş "
  "akışlarında koşulur (Çizelge 14.1).")
caption_table("14.1", "GitHub Actions CI/CD İş Akışları")
table(["İş Akışı", "Amaç"],
      [("ci.yml", "Backend (SQLite+PostgreSQL matrisi), frontend (tsc+eslint+vitest), a11y geçidi"),
       ("e2e-tests.yml", "Playwright ile uçtan uca testler"),
       ("contract-tests.yml", "Pact ile tüketici-sağlayıcı sözleşme testleri"),
       ("visual-regression.yml", "Görsel regresyon (snapshot) testleri"),
       ("mutation.yml", "Stryker ile mutasyon testi (test etkinliği ölçümü)"),
       ("lighthouse.yml", "Performans/erişilebilirlik denetimi (Lighthouse CI)"),
       ("load-test.yml", "Yük testi"),
       ("test-metrics.yml", "Test metriklerinin otomatik üretimi"),
       ("flaky-report.yml", "Kararsız (flaky) testlerin raporlanması")],
      [4.5, 11.5], fs=9)
h2("14.2 Deployment (Dağıtım) Süreci")
p("Uygulama, tek komutla ayağa kalkabilen Docker/docker-compose yapısıyla paketlenmiştir. Canlı "
  "ortamda ön yüz Vercel’de (https://yasca-dental-clinic.vercel.app), arka uç ise Render’da "
  "(gunicorn ile) barındırılmaktadır. Veritabanı PostgreSQL’dir. Şifre sıfırlama gibi "
  "işlemlerde e-posta gönderimi, bulut barındırma sağlayıcısının SMTP port kısıtlamalarını "
  "aşmak için Resend HTTP API’si (django-anymail) üzerinden yapılır. Canlı ortamda kiracı "
  "bilgisi, ön yüzden arka uca X-Tenant başlığı ile taşınır.")
page_break()

# ───── 15. UYGULAMA PLANI ─────
h1("15. Uygulama Planı")
p("Mevcut sürümün ardından planlanan sürüm yol haritası aşağıdaki gibidir:")
table(["Versiyon", "Planlanan Özellikler", "Tahmini Tarih"],
      [("v1.1", "SMS/WhatsApp entegrasyonu, takvimde sürükle-bırak (drag & drop)", "2026 Q3"),
       ("v1.2", "Raporlama modülü, PDF çıktı desteği", "2026 Q4"),
       ("v2.0", "Mobil uygulama (React Native), gelişmiş çok dilli destek", "2027 Q1"),
       ("v2.1", "Röntgen/görüntüleme entegrasyonu, gelişmiş analitik dashboard", "2027 Q2")],
      [3, 10, 4])
page_break()

# ───── 16. KALİTE ─────
h1("16. Kalite Pratikleri")
p("Kalite, hem otomatik geçitler hem de süreç kurallarıyla güvence altına alınmıştır.")
h2("16.1 Versiyon Yönetimi")
p("Kaynak kod Git ile yönetilmiş; GitHub üzerinde dal (branch) tabanlı iş akışı, PR gözden "
  "geçirmesi ve Conventional Commits standardı uygulanmıştır. Birleştirmeler squash & merge ile "
  "yapılmış, böylece ana dal (main) geçmişi temiz tutulmuştur.")
h2("16.2 Geliştirme Ortamı ve IDE")
bullet("Editör: VS Code / Cursor (yapay zeka entegrasyonlu).")
bullet("Çalışma zamanı: Python 3.12, Node.js 18+; konteynerizasyon için Docker Desktop.")
bullet("Yardımcılar: PlantUML eklentisi (UML), ESLint + Prettier, husky (git hook).")
h2("16.3 Erişilebilirlik ve Statik Analiz")
p("WCAG 2.1 AA için “sıfır tolerans” politikası benimsenmiş (W3C, 2018); vitest-axe ve "
  "@axe-core/playwright ile otomatik erişilebilirlik testleri CI’de zorunlu (required) geçit "
  "olarak çalıştırılmıştır. Ön yüzde eslint --max-warnings 0 ve tsc --noEmit ile sıfır "
  "uyarı/sıfır tip hatası geçidi uygulanır; any kullanımı yasaktır.")
page_break()

# ───── 17. ÖNYÜZLER ─────
h1("17. Uygulama Önyüzleri ve Kullanım Kılavuzu")
p("Aşağıdaki ekran görüntüleri, sistemin canlı ortamından (https://yasca-dental-clinic.vercel.app) "
  "bir klinik kiracısı ile giriş yapılarak alınmıştır ve temel kullanım kılavuzu işlevi görür.")
shots = [
    ("00_public_landing.png", "17.1", "Genel Tanıtım (SaaS Landing) Sayfası",
     "Ziyaretçilerin karşılandığı, klinik kayıt ve giriş yönlendirmelerinin yapıldığı tanıtım sayfası."),
    ("01_login.png", "17.2", "Klinik Giriş (Login) Ekranı",
     "Kliniğe özel giriş ekranı; e-posta/kullanıcı adı ve şifre ile JWT tabanlı kimlik doğrulama, dil seçimi ve şifre sıfırlama."),
    ("02_dashboard.png", "17.3", "Ana Panel (Dashboard)",
     "Bugünkü randevular, bekleyen hastalar ve toplam hasta gibi özet kartlar ile günün randevu listesi."),
    ("03_randevular.png", "17.4", "Randevu Takvimi (Haftalık Görünüm)",
     "Günlük/haftalık görünüm; boş slota tıklayarak çakışma kontrollü randevu oluşturma."),
    ("04_hastalar.png", "17.5", "Hasta Yönetimi / Hasta Listesi",
     "Ad, soyad veya telefon ile arama; hasta ekleme ve detaya gitme."),
    ("05_hasta_profili.png", "17.6", "Hasta Profili Ekranı",
     "Profil bilgileri, anamnez, tedavi geçmişi, ödemeler, dokümanlar ve diş şeması sekmeleri."),
    ("06_odontogram.png", "17.7", "Dijital Diş Şeması (Odontogram)",
     "FDI numaralandırmasına göre üst/alt çene; işlem türlerinin renk kodlu gösterimi (dolgu, kanal, kron, çekim vb.)."),
    ("09_tedavi_gecmisi.png", "17.8", "Hasta Tedavi ve Randevu Geçmişi",
     "Hastaya ait tedavi ve randevu kayıtlarının kronolojik dökümü."),
    ("07_tedavi_turleri.png", "17.9", "Tedavi Türleri Yönetimi",
     "Tedavi türleri ve varsayılan fiyatlarının tanımlanması, düzenlenmesi ve silinmesi."),
    ("08_ayarlar.png", "17.10", "Klinik Ayarları",
     "Çalışma saatleri, çalışma günleri ve numara/ülke ayarlarının yönetimi."),
]
for fname, num, cap, desc in shots:
    p(desc, size=Pt(11))
    figure(S(fname), num, cap, width=Inches(5.9))
page_break()

# ───── 18. USE-CASE / UML ─────
h1("18. Use-Case'ler ve UML Dokümantasyonu")
p("Sistemin temel kullanım durumları (use-case) Çizelge 18.1’de özetlenmiştir. Bu kullanım "
  "durumlarının dinamik ve statik UML modelleri (sekans ve sınıf diyagramları) Bölüm 12’de "
  "sunulmuştur; PlantUML kaynak dosyaları docs/uml/ dizininde versiyonlanmaktadır.")
caption_table("18.1", "Use-Case (Kullanım Durumu) Özeti")
table(["Kullanım Durumu", "Aktör", "Açıklama"],
      [("Giriş yap", "Tüm roller", "E-posta/şifre ile JWT tabanlı oturum açma"),
       ("Hasta kaydı oluştur", "Asistan, Yönetici", "Ad/soyad/telefon ile hızlı kayıt"),
       ("Randevu oluştur", "Asistan, Hekim", "Çakışma kontrollü randevu"),
       ("Tedavi gir", "Hekim", "Odontogram üzerinden işlem kaydı"),
       ("Ödeme al", "Asistan, Yönetici", "Tutar ve açıklama ile ödeme"),
       ("Klinik ayarlarını düzenle", "Yönetici", "Çalışma saatleri ve günleri"),
       ("Personel yönet", "Yönetici", "Hekim/asistan ekleme, şifre sıfırlama")],
      [4.5, 4, 7.5], fs=9)
page_break()

# ───── 19. PROJE YÖNETİMİ ─────
h1("19. Proje Yönetim Pratikleri")
p("Proje yönetimi; şeffaflık, izlenebilirlik ve sürekli entegrasyon ilkeleri üzerine "
  "kurulmuştur:")
bullet("İş takibi: Görevler dal (branch) ve PR’lar üzerinden izlenmiş; her iş bir özellik dalına "
       "ve gözden geçirilen bir çekme isteğine bağlanmıştır.")
bullet("Kilometre taşları (milestone): Geliştirme, ders adımlarına (Adım 1–5) göre "
       "kilometretaşlarına bölünmüştür.")
bullet("Sürekli entegrasyon: Her push/PR’da CI iş akışları otomatik koşturularak erken geri "
       "bildirim sağlanmıştır.")
bullet("Karar kayıtları: Önemli teknik kararlar docs/adr/ altında ADR formatında belgelenmiştir.")
bullet("İletişim ve gözden geçirme: Kod gözden geçirme (en az 1 onay) zorunlu tutularak bilgi "
       "paylaşımı ve ortak sahiplik sağlanmıştır.")
page_break()

# ───── 20. NELER ÖĞRENDİK ─────
h1("20. Neler Öğrendik (Başarılı ve Başarısız Yönler) ve Öneriler")
p("Başarılı Yönler", bold=True)
bullet("Çok katmanlı test altyapısı (403 test) ve CI/CD geçitleri, değişikliklerin güvenle "
       "birleştirilebilmesini sağladı.")
bullet("Schema-per-tenant izolasyon, kiracılar arası veri güvenliğini yapısal olarak güvence "
       "altına aldı.")
bullet("Merkezî api.ts servis katmanı ve net BCE ayrımı, ön yüz–arka uç entegrasyonunu "
       "sadeleştirdi.")
p("Başarısız / Zorlanılan Yönler", bold=True)
bullet("Multi-tenant kurulum ve django-tenants şema yönetimi ilk aşamada karmaşık bulundu; ara "
       "katman (middleware) ile çözüldü.")
bullet("Çakışma kontrolünde iptal edilmiş randevuların da çakışma sayılması hatası tespit "
       "edilip status filtresiyle giderildi.")
bullet("Test piramidinde birim test oranı (%37,5) hedefin (%70) altında kaldı.")
p("Öneriler", bold=True)
bullet("İş mantığını ViewSet/Serializer’dan bağımsız bir Servis Katmanına (Service Layer) taşımak.")
bullet("Birim test oranını artırarak test piramidini dengelemek.")
bullet("API hatalarını RFC 7807 (Problem Details) standardında döndüren merkezî bir hata yönetimi kurmak.")
page_break()

# ───── 21. TAKIM DEĞERLENDİRME ─────
h1("21. Takım Üyelerinin Proje Değerlendirme Raporu")
caption_table("21.1", "Takım Üyeleri Proje Değerlendirmesi")
table(["Üye", "Değerlendirme"],
      [("Yaman Halloum", "Backend ve multi-tenant altyapıda kritik katkı; Django/DRF derinliği arttı."),
       ("Ali Üre", "Tam yığın geliştirme ve dağıtım sorumluluğu; en yüksek katkı hacmi."),
       ("Cihan Kurtbey", "Kimlik doğrulama, şifre sıfırlama ve e-posta entegrasyonunda etkin rol."),
       ("Şükrü Yeşilmen", "Test, kalite ve dokümantasyon süreçlerine katkı.")],
      [4, 12])
page_break()

# ───── 22. YAPAY ZEKA ─────
h1("22. Yapay Zeka Kullanım Pratikleri")
p("Proje boyunca yapay zeka destekli araçlar (GitHub Copilot, Anthropic Claude / Claude Code, "
  "ChatGPT vb.) bir “eş-programcı (pair programmer)” olarak, yazılım yaşam döngüsünün hemen her "
  "aşamasında kullanılmıştır. Temel ilke, üretilen hiçbir çıktının test edilmeden ve insan "
  "denetiminden (human-in-the-loop) geçmeden birleştirilmemesidir (Çizelge 22.1).")
caption_table("22.1", "Yapay Zeka Araçlarının Kullanım Alanları")
table(["Aşama", "Yapay Zeka Kullanımı", "İnsan Denetimi"],
      [("Tasarım / Analiz", "Gereksinim netleştirme, UML ve veri modeli taslakları",
        "Kararların ekipçe onaylanması"),
       ("Kodlama", "Model/serializer/viewset ve React bileşeni iskeletleri, boilerplate üretimi",
        "Kod gözden geçirme, refactoring"),
       ("Test", "Pozitif/negatif test senaryoları ve test verisi fabrikaları",
        "Testlerin çalıştırılıp doğrulanması"),
       ("Hata Ayıklama", "Hata mesajı yorumlama, olası kök neden ve çözüm önerileri",
        "Çözümün uygulanıp test edilmesi"),
       ("Dokümantasyon", "README/ADR/docstring taslakları, bu raporun yapılandırılması",
        "İçeriğin doğrulanması ve düzenlenmesi"),
       ("Otomasyon", "Playwright ile ekran görüntüsü ve rapor üretim betikleri",
        "Çıktıların gözden geçirilmesi")],
      [3.2, 7.3, 5.5], fs=9)
p("Somut Kullanım Örnekleri", bold=True)
bullet("Kod Üretimi: Tekrar eden CRUD serializer/viewset yapıları ve Radix UI tabanlı diyalog "
       "bileşenleri, kısa komutlarla hızlıca iskelet hâline getirilmiş, ardından elle "
       "iyileştirilmiştir.")
bullet("Test Üretimi: Randevu çakışma kontrolü gibi kritik kurallar için hem pozitif hem negatif "
       "senaryolar türetilmiş; ekip bu testleri çalıştırarak doğrulamıştır.")
bullet("Hata Ayıklama: Multi-tenant kurulumdaki hatalarda olası nedenler sıralanmış; çözümler "
       "ekipçe uygulanmıştır.")
bullet("Bu Rapor: Kod tabanı analiz edilip, canlı siteden Playwright ile otomatik ekran "
       "görüntüleri alınarak ve python-docx ile biçimlendirilerek üretilmiş; içerik ekip "
       "tarafından doğrulanmıştır.")
p("Komut (Prompt) Pratikleri", bold=True)
bullet("Bağlam verme: İlgili dosya/standartların açıkça belirtilmesi (örn. “PEP 8’e uy”).")
bullet("Küçük ve doğrulanabilir adımlar: Büyük görevlerin parçalanması ve her adımın testle doğrulanması.")
bullet("Güvenlik: Hassas verilerin (şifre, TCKN) komutlara gömülmemesi; kalite geçitlerinin atlatılmaması.")
p("Genel Değerlendirme: Yapay zeka destekli geliştirme bir “otomatik pilot” değil, bir “güç "
  "çarpanı (force multiplier)” olarak değerlendirilmiştir. Doğru komut pratikleri, güçlü test "
  "geçitleri ve sürekli insan denetimi ile birlikte kullanıldığında verimliliği belirgin "
  "biçimde artırmış; ancak üretilen çıktının doğruluğu daima test ve gözden geçirme ile "
  "doğrulanmıştır.")
page_break()

# ───── KAYNAKLAR (APA) ─────
h1("Kaynaklar")
refs = [
    "Conventional Commits. (2023). Conventional Commits 1.0.0. https://www.conventionalcommits.org/",
    "Django REST Framework. (2026). Django REST framework documentation. https://www.django-rest-framework.org/",
    "Django Software Foundation. (2026). Django documentation. https://docs.djangoproject.com/",
    "django-tenants. (2026). Tenant schemas for Django. https://django-tenants.readthedocs.io/",
    "Goodger, D., & van Rossum, G. (2001). PEP 257 – Docstring conventions. https://peps.python.org/pep-0257/",
    "Meta. (2026). React documentation. https://react.dev/",
    "Microsoft. (2026). Playwright documentation. https://playwright.dev/",
    "Microsoft. (2026). TypeScript documentation. https://www.typescriptlang.org/docs/",
    "OpenAPI Initiative. (2026). OpenAPI specification. https://spec.openapis.org/",
    "The PostgreSQL Global Development Group. (2026). PostgreSQL documentation. https://www.postgresql.org/docs/",
    "van Rossum, G., Warsaw, B., & Coghlan, N. (2001). PEP 8 – Style guide for Python code. https://peps.python.org/pep-0008/",
    "W3C. (2018). Web Content Accessibility Guidelines (WCAG) 2.1. https://www.w3.org/TR/WCAG21/",
]
for rr in refs:
    para = doc.add_paragraph(); para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.left_indent = Cm(1.25); para.paragraph_format.first_line_indent = Cm(-1.25)
    para.paragraph_format.space_after = Pt(6)
    r = para.add_run(rr); set_run_font(r); r.font.size = Pt(12)
page_break()

# ───── EKLER ─────
h1("Ekler")
p("EK A – Gerçek Kod Örnekleri", bold=True)
p("A.1 Randevu Çakışma Algoritması (Serializer doğrulama):", bold=True)
code_block(
    "class AppointmentCreateSerializer(serializers.ModelSerializer):\n"
    "    def validate(self, data):\n"
    "        doctor, date, time = data.get('doctor'), data.get('date'), data.get('time')\n"
    "        if doctor and date and time:\n"
    "            qs = Appointment.objects.filter(\n"
    "                doctor=doctor, date=date, time=time,\n"
    "                status=Appointment.Status.SCHEDULED, is_active=True)\n"
    "            if self.instance:\n"
    "                qs = qs.exclude(pk=self.instance.pk)\n"
    "            if qs.exists():\n"
    "                raise serializers.ValidationError(\n"
    "                    'Bu hekime bu saatte zaten randevu kayıtlı.')\n"
    "        return data")
p("A.2 RBAC ile Sorgu Filtreleme (ViewSet):", bold=True)
code_block(
    "class AppointmentViewSet(AuditLogMixin, viewsets.ModelViewSet):\n"
    "    permission_classes = [IsAuthenticated]\n"
    "    def get_queryset(self):\n"
    "        user = self.request.user\n"
    "        qs = Appointment.objects.filter(is_active=True)\\\n"
    "                                .select_related('patient', 'doctor')\n"
    "        if user.role == CustomUser.Role.DOCTOR and not user.is_superuser:\n"
    "            qs = qs.filter(doctor=user)\n"
    "        return qs.order_by('date', 'time')")
doc.add_paragraph()
p("EK B – Çalıştırma ve Demo Bilgileri", bold=True)
bullet("Canlı ortam (ön yüz): https://yasca-dental-clinic.vercel.app")
bullet("Lokal kurulum: docker-compose up --build -d && npm run dev")
bullet("Demo verisi: docker-compose run --rm backend sh -c \"python manage.py seed_demo_data\"")
bullet("Kod deposu yapısı: monorepo (frontend/ + backend/ + docs/).")
page_break()

# ───── ÖZGEÇMİŞ ─────
h1("Özgeçmiş")
p("Bu proje, İstanbul Sağlık ve Teknoloji Üniversitesi Yazılım Mühendisliği Bölümü öğrencileri "
  "Yaman Halloum, Ali Üre, Cihan Kurtbey ve Şükrü Yeşilmen’den oluşan Grup 4 tarafından "
  "hazırlanmıştır. Ekip; arka uç ve ön uç geliştirme, çok-kiracılı mimari, test otomasyonu, "
  "DevOps/dağıtım ve teknik dokümantasyon alanlarında ortak çalışma yürütmüştür.")

# ───────────────────────── Bölümleri sonlandır ─────────────────────────
setup_section(doc.sections[0], number=False)                       # Kapak/Onay/Beyan/Önsöz/İçindekiler
setup_section(doc.sections[1], number=True, fmt="lowerRoman", start=1)  # Ön sayfalar i,ii,iii
setup_section(doc.sections[2], number=True, fmt="decimal", start=1)     # Giriş'ten itibaren 1,2,3

enable_update_fields()
doc.save(OUT)
print("Rapor olusturuldu:", OUT)
