from docx import Document
from docx.shared import Pt
import sys

def format_docx(filename):
    doc = Document(filename)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            if not run.font.size:
                run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.5
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        if not run.font.size:
                            run.font.size = Pt(12)

    doc.save(filename)

if __name__ == "__main__":
    format_docx('Grup_4_Proje_Adım3.docx')
