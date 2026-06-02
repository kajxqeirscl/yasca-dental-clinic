# -*- coding: utf-8 -*-
"""
Yaşca Diş Kliniği Yönetim Sistemi - Final Proje Raporu Oluşturucu
Yazılım Gerçekleme, Test ve Bakım Dersi
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style Helpers ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # Dark blue
    if level == 1:
        hs.font.size = Pt(18)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark blue background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F497D"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cells[col_idx].text = ''
            p = cells[col_idx].paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Alternating row colors
            if row_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EDF3"/>')
                cells[col_idx]._tc.get_or_add_tcPr().append(shading)
    
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)
    
    return table

def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    return p

def add_code_block(doc, code, language="python"):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    # Add shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    pPr.append(shading)
    
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

# ═══════════════════════════════════════════════════════════════════════════
# KAPAK SAYFASI
# ═══════════════════════════════════════════════════════════════════════════

# University name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
run = p.add_run('İSTANBUL SAĞLIK VE TEKNOLOJİ ÜNİVERSİTESİ')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('MÜHENDİSLİK VE DOĞA BİLİMLERİ FAKÜLTESİ')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('YAZILIM MÜHENDİSLİĞİ BÖLÜMÜ')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Separator line
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 50)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Course name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
run = p.add_run('YAZ402 – Yazılım Gerçekleme, Test ve Bakım')
run.bold = True
run.font.size = Pt(14)

# Project title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(36)
run = p.add_run('FINAL PROJE RAPORU')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Yaşca Diş Kliniği Yönetim Sistemi')
run.font.size = Pt(16)
run.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Çok-Kiracılı (Multi-Tenant) SaaS Mimarisi')
run.font.size = Pt(12)

# Separator
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('━' * 50)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Sistem Analizinden Test Sonuçlarına Tam Dokümantasyon')
run.font.size = Pt(12)
run.italic = True

# Group info
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Grup 4')
run.bold = True
run.font.size = Pt(14)

members = [
    'Yaman Halloum',
    'Ali Üre',
    'Cihan Kurtbey',
    'Berkay Aydın'
]
for member in members:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(member)
    run.font.size = Pt(12)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2024 – 2025 Bahar Dönemi')
run.font.size = Pt(11)

# Page break after cover
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# İÇİNDEKİLER
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('İÇİNDEKİLER', level=1)

toc_items = [
    ('1.', 'Giriş ve Proje Kapsamı'),
    ('2.', 'Yazılım Gereksinimleri'),
    ('  2.1.', 'Fonksiyonel İsterler'),
    ('  2.2.', 'Fonksiyonel Olmayan İsterler'),
    ('3.', 'Detaylı UML Tasarımı'),
    ('  3.1.', 'Sekans Diyagramları'),
    ('  3.2.', 'Sınıf Diyagramları'),
    ('4.', 'Teknoloji Karar Analizi (DAR)'),
    ('  4.1.', 'Backend Dil ve Framework Seçimi'),
    ('  4.2.', 'Frontend Framework Seçimi'),
    ('  4.3.', 'Veri Tabanı Seçimi'),
    ('  4.4.', 'Karar Matrisleri'),
    ('5.', 'İç ve Dış Arayüzler'),
    ('  5.1.', 'API Endpoint Kontratları'),
    ('  5.2.', 'Dış Servis Entegrasyonları'),
    ('  5.3.', 'Multi-Tenant Mimarisi'),
    ('6.', 'Birim Test Sonuçları'),
    ('  6.1.', 'Test Ortamı ve Araçlar'),
    ('  6.2.', 'Test Edilen Kritik Sınıflar'),
    ('  6.3.', 'Birim Test Sonuç Matrisi'),
    ('7.', 'Bakım ve Gelecek Önerileri'),
    ('8.', 'Sonuç'),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    if num.startswith('  '):
        p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. GİRİŞ VE PROJE KAPSAMI
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('1. Giriş ve Proje Kapsamı', level=1)

doc.add_heading('Projenin Amacı', level=2)
doc.add_paragraph(
    'Bu proje, diş kliniklerinin günlük operasyonlarını dijitalleştirmek ve '
    'merkezi bir platform üzerinden yönetmelerini sağlamak amacıyla geliştirilmiş '
    'çok-kiracılı (multi-tenant) bir SaaS (Software as a Service) uygulamasıdır. '
    '"Yaşca Diş Kliniği Yönetim Sistemi" adı verilen proje; hasta kaydı, randevu '
    'yönetimi, tedavi ve finansal takip, dijital odontogram (diş şeması) ve klinik '
    'yönetim modüllerini tek bir çatı altında sunan kapsamlı bir web uygulamasıdır.'
)

doc.add_paragraph(
    'Sağlık sektöründe hâlâ kağıt tabanlı veya birbirinden kopuk yazılımlarla '
    'yürütülen klinik süreçleri, veri kaybı, randevu çakışması ve güvenlik '
    'açıkları gibi ciddi sorunlara yol açmaktadır. Bu proje, söz konusu problemleri '
    'çözmek üzere tasarlanmış olup, her bir kliniğin verilerini fiziksel olarak '
    'izole eden PostgreSQL şema tabanlı multi-tenant mimarisi ile güvenliği en üst '
    'düzeyde tutmaktadır.'
)

doc.add_heading('Ana Hedefler', level=2)
hedefler = [
    'Hasta kayıt ve anamnez (tıbbi özgeçmiş) yönetimini dijitalleştirmek',
    'Çakışma kontrollü randevu sistemi ile operasyonel verimliliği artırmak',
    'Tedavi ve ödeme takibini entegre bir yapıda sunmak',
    'Dijital diş şeması (Odontogram) ile klinik görselleştirme sağlamak',
    'Güvenli JWT tabanlı kimlik doğrulama ve rol bazlı yetkilendirme uygulamak',
    'Multi-tenant (çok-kiracılı) SaaS mimarisi ile birden fazla kliniğe hizmet vermek'
]
for h in hedefler:
    add_bullet(doc, h)

doc.add_heading('Hedef Kullanıcılar', level=2)
doc.add_paragraph(
    'Sistem; Klinik Yöneticisi (Admin), Hekim (Doktor) ve Asistan olmak üzere '
    'üç temel kullanıcı rolüne sahiptir. Her rol, kendi yetkilerine uygun '
    'modüllere erişebilmektedir.'
)

kullanicilar = [
    ('Klinik Yöneticisi (Admin)', 'Klinik ayarları, çalışma saatleri, personel yönetimi, tüm verilere tam erişim'),
    ('Hekim (Doktor)', 'Kendi hastalarının randevuları, tedavi girişi, odontogram, dashboard'),
    ('Asistan', 'Hasta kaydı, randevu oluşturma, ödeme girişi')
]

tbl = add_styled_table(doc, 
    ['Rol', 'Erişim ve Yetkiler'],
    kullanicilar,
    [5, 12]
)

doc.add_heading('Kritik Başarı Kriterleri', level=2)
kriterler = [
    'Randevu çakışması %0 oranında gerçekleşmelidir',
    'Kiracılar arası veri sızıntısı kesinlikle önlenmelidir',
    'API yanıt süreleri 2 saniyenin altında olmalıdır',
    'Sistem, en az 54 birim testin tamamını başarıyla geçmelidir',
    'Tüm API endpoint\'leri Swagger (OpenAPI) ile dokümante edilmelidir'
]
for k in kriterler:
    add_bullet(doc, k)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. YAZILIM GEREKSİNİMLERİ
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('2. Yazılım Gereksinimleri', level=1)

doc.add_heading('2.1. Fonksiyonel İsterler', level=2)
doc.add_paragraph(
    'Sistemin ne yapması gerektiğini tanımlayan, Proje Adımı 1\'de oluşturulan '
    'fonksiyonel gereksinim listesi aşağıda özetlenmiştir:'
)

fr_table = [
    ('FR-01', 'Kimlik Doğrulama', 'Klinik personelinin JWT tabanlı güvenli girişi. Kullanıcı adı ve parola ile sisteme giriş yapılması, access ve refresh token çifti döndürülmesi.'),
    ('FR-02', 'Hasta Yönetimi', 'Yeni hasta kaydı oluşturulması ve tıbbi özgeçmiş (Anamnez) eklenmesi. Hasta bilgilerinin güncellenmesi ve listelenmesi.'),
    ('FR-03', 'Randevu Yönetimi', 'Çakışma kontrollü randevu oluşturulması. Aynı hekime aynı tarih ve saatte birden fazla randevu girilmesinin engellenmesi.'),
    ('FR-04', 'Klinik Operasyon (Dashboard)', 'Günlük özet (Dashboard) verilerinin görüntülenmesi: bugüne ait randevular, tamamlanan işlem sayıları ve aktif hasta sayısı.'),
    ('FR-05', 'Tedavi ve Finans', 'Yapılan tedavilerin sisteme girilmesi, ödeme kayıtlarının oluşturulması ve tedavi-finans entegrasyonunun sağlanması.'),
    ('FR-06', 'Sistem Yönetimi', 'Yöneticinin klinik çalışma saatleri ve ayarlarını güncellemesi. Yetki kontrolü (403 Forbidden) ile yalnızca Admin rolünün erişimi.'),
]

add_styled_table(doc,
    ['ID', 'Gereksinim', 'Açıklama'],
    fr_table,
    [2, 3.5, 11.5]
)

doc.add_heading('2.2. Fonksiyonel Olmayan İsterler', level=2)
doc.add_paragraph(
    'Performans, güvenlik ve ölçeklenebilirlik gibi kalite nitelikleri '
    'aşağıda listelenmiştir:'
)

nfr_table = [
    ('NFR-01', 'Performans', 'API sorgu süresi 2 saniyeden az olmalıdır.'),
    ('NFR-02', 'Güvenlik', 'JWT tabanlı kimlik doğrulama, RBAC (Role-Based Access Control) ile yetkilendirme uygulanmalıdır.'),
    ('NFR-03', 'Veri İzolasyonu', 'Multi-tenant mimaride her klinik verisi PostgreSQL şema seviyesinde fiziksel olarak izole edilmelidir.'),
    ('NFR-04', 'Ölçeklenebilirlik', 'Sistem, birden fazla kliniğin eşzamanlı kullanımını desteklemelidir.'),
    ('NFR-05', 'Bakım Yapılabilirlik', 'Modüler mimari ve yüksek test kapsamı ile bakım maliyeti düşürülmelidir.'),
    ('NFR-06', 'Kullanılabilirlik', 'Kullanıcı arayüzü sezgisel ve Türkçe olmalıdır.'),
]

add_styled_table(doc,
    ['ID', 'Kategori', 'Açıklama'],
    nfr_table,
    [2, 3.5, 11.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. DETAYLI UML TASARIMI
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('3. Detaylı UML Tasarımı', level=1)

doc.add_heading('Sınıf ve Sekans Analizi', level=2)
doc.add_paragraph(
    'Proje Adımı 1 kapsamında, sistemin nesne tabanlı tasarımı standart UML '
    'notasyonlarına uygun olarak modellenmiştir. Bu bölümde, sekans ve sınıf '
    'diyagramlarının projeye katkısı ve diyagramlar arası tutarlılık açıklanmaktadır.'
)

doc.add_heading('3.1. Sekans Diyagramları', level=2)
doc.add_paragraph(
    'Projenin temel kullanım durumlarını (use-case) temsil eden dinamik süreçler, '
    'Boundary (Arayüz), Control (Servis/API) ve Entity (Varlık) nesneleri kullanılarak '
    'modellenmiştir. Aşağıda her bir fonksiyonel gereksinim için tasarlanan sekans '
    'diyagramlarının özeti verilmiştir:'
)

sekans_aciklamalari = [
    ('SD-01: Kullanıcı Girişi ve Oturum Doğrulaması (FR-01)',
     'Klinik personelinin sisteme giriş yaptığı ve yetkilendirme jetonunun (JWT) '
     'localStorage üzerinde saklanarak kullanıcı rolünün doğrulandığı süreçtir. '
     'LoginPage → ApiService → AuthController → CustomUser etkileşim akışı modellenmiştir.'),
    ('SD-02: Hasta Kaydı ve Anamnez Güncelleme (FR-02)',
     'Asistan veya Hekimin yeni bir hasta profili oluşturduğu ve tıbbi özgeçmiş '
     '(Anamnez) kayıtlarının veritabanına yazıldığı iş akışıdır. PatientDialog → '
     'ApiService → PatientViewSet → Patient + Anamnesis etkileşimi kurgulanmıştır.'),
    ('SD-03: Randevu Oluşturma ve Çakışma Kontrolü (FR-03)',
     'Asistanın yeni randevu talebi oluşturduğu süreçtir. Sistem, belirtilen saat '
     'diliminde ilgili hekimin uygunluğunu (Exists kontrolü) denetler. Çakışma durumunda '
     'hata döndürülürken, uygunluk durumunda randevu "Scheduled" statüsünde kaydedilir.'),
    ('SD-04: Günlük Dashboard Özet Akışı (FR-04)',
     'Klinik personelinin güne başlarken açtığı ana panelin veri çekme sürecidir. '
     'Sistem, bugüne ait randevuları, tamamlanan işlem sayılarını ve aktif hasta '
     'sayısını filtreleyerek arayüze sunar.'),
    ('SD-05: Tedavi ve Ödeme Girişi (FR-05)',
     'Hekimin veya asistanın bir hastaya ait tedavi türlerini çektiği, işlemi '
     'kaydettiği ve eşzamanlı olarak finansal ödemeyi sisteme yansıttığı uçtan uca '
     'klinik operasyon sürecidir.'),
    ('SD-06: Klinik Ayarları Güncelleme (FR-06)',
     'Yalnızca Admin yetkisine sahip yöneticinin erişebildiği, yetki kontrolü '
     '(Authorization 403 Forbidden) aşamalarını içeren klinik çalışma saatleri '
     'düzenleme sürecidir.'),
]

for title, desc in sekans_aciklamalari:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

p = doc.add_paragraph()
run = p.add_run('Not: Sekans diyagramlarının görsel temsilleri Proje Adımı 1 raporunda (Grup_4_Proje_Adım1.pdf) yer almaktadır.')
run.italic = True
run.font.size = Pt(10)

doc.add_page_break()

doc.add_heading('3.2. Sınıf Diyagramları', level=2)

doc.add_heading('3.2.1. Domain Katmanı (Veri Varlıkları) Sınıf Diyagramı', level=3)
doc.add_paragraph(
    'Veritabanı tablolarının nesne tabanlı karşılıklarını (Entities) ve ilişkilerini '
    '(1..* vb. çokluk değerleri) gösteren modeldir. Temel varlıklar şunlardır:'
)

domain_classes = [
    ('Clinic', 'Kiracı (Tenant) bilgileri, çalışma saatleri, adres'),
    ('CustomUser', 'Kullanıcılar (Hekim, Asistan, Admin) – rol ve yetki bilgileri'),
    ('Patient', 'Hasta bilgileri – ad, soyad, telefon, TCKN'),
    ('Anamnesis', 'Tıbbi özgeçmiş – alerji, kronik hastalık bilgileri'),
    ('Appointment', 'Randevu – tarih, saat, durum (Scheduled/Completed/Cancelled)'),
    ('Treatment', 'Tedavi kaydı – tedavi türü, diş numarası, fiyat'),
    ('TreatmentType', 'Tedavi türleri – dolgu, kanal, çekim vb.'),
    ('Payment', 'Ödeme kaydı – tutar, ödeme yöntemi, tarih'),
]

add_styled_table(doc,
    ['Sınıf', 'Açıklama'],
    domain_classes,
    [4, 13]
)

doc.add_paragraph()
doc.add_paragraph(
    'Sınıflar arası ilişki türleri:'
)
iliskiler = [
    'Clinic → Patient: 1..* Association (Bir kliniğe birden fazla hasta kayıtlıdır)',
    'Clinic → CustomUser: 1..* Association (Bir klinikte birden fazla personel bulunur)',
    'Patient → Appointment: 1..* Composition (Hasta silindiğinde randevuları da silinir)',
    'Patient → Anamnesis: 1..1 Composition (Her hastanın en fazla bir anamez kaydı vardır)',
    'Appointment → Treatment: 1..* Aggregation (Bir randevuda birden fazla tedavi yapılabilir)',
    'Treatment → Payment: 1..1 Association (Her tedaviye bir ödeme bağlanır)',
]
for i in iliskiler:
    add_bullet(doc, i)

doc.add_heading('3.2.2. Uygulama Katmanı (Sınır ve Kontrol Sınıfları)', level=3)
doc.add_paragraph(
    'Frontend arayüz bileşenlerinin (Boundary), API servisleriyle (ApiService) nasıl '
    'haberleştiğini ve Backend Controller (ViewSet/View) yapılarını gösteren mimari '
    'tasarımdır. Bu katmandaki temel sınıflar:'
)

uygulama_classes = [
    ('LoginPage (Boundary)', 'Kullanıcı giriş arayüzü'),
    ('DashboardUI (Boundary)', 'Ana panel / günlük özet arayüzü'),
    ('PatientDialog (Boundary)', 'Hasta kayıt ve düzenleme formu'),
    ('AppointmentDialog (Boundary)', 'Randevu oluşturma formu'),
    ('ApiService (Control)', 'Frontend-Backend arası HTTP iletişim servisi'),
    ('AuthController (Control)', 'JWT tabanlı kimlik doğrulama denetleyicisi'),
    ('PatientViewSet (Control)', 'Hasta CRUD işlemleri denetleyicisi'),
    ('AppointmentViewSet (Control)', 'Randevu CRUD ve çakışma kontrolü denetleyicisi'),
    ('DashboardView (Control)', 'Günlük özet veri toplama denetleyicisi'),
]

add_styled_table(doc,
    ['Sınıf (Stereotip)', 'Sorumluluk'],
    uygulama_classes,
    [6, 11]
)

doc.add_heading('3.3. Tasarım Kararları ve Tutarlılık Analizi', level=2)
doc.add_paragraph(
    'Tasarım süreci, "Sınıf ve Sekans diyagramları arası mutlak tutarlılık" prensibiyle '
    'iteratif olarak gerçekleştirilmiştir:'
)

tutarlilik = [
    ('Metod ve Mesaj Tutarlılığı', 
     'Sekans diyagramlarında arayüzden gönderilen fetchDashboardToday(), createAppointment(payload) '
     'veya login(username, password) gibi tüm mesajlar, Uygulama Katmanı Sınıf Diyagramındaki '
     'ApiService sınıfının metodları (operasyonları) olarak birebir tanımlanmıştır.'),
    ('Sorumlulukların Ayrılığı (Separation of Concerns)',
     'Arayüzler (Boundary) doğrudan veritabanı (Entity) ile konuşturulmamış, araya güvenliği ve '
     'iş mantığını sağlamak amacıyla Kontrolcüler (ViewSet/Serializer) yerleştirilerek MVC/MVT '
     'mimarisine uyum sağlanmıştır.'),
    ('Çokluk (Multiplicity) Doğrulaması',
     'Sınıf diyagramında Clinic ile Patient veya CustomUser (Hekim) arasında kurulan 1..* '
     '(Bire-Çok) ilişkiler, sistemin gerçek dünya işleyişini doğru bir şekilde yansıtmaktadır.'),
]

for title, desc in tutarlilik:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

p = doc.add_paragraph()
run = p.add_run('Not: Sınıf diyagramlarının görsel temsilleri Proje Adımı 1 raporunda (Grup_4_Proje_Adım1.pdf) yer almaktadır.')
run.italic = True
run.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. TEKNOLOJİ KARAR ANALİZİ (DAR)
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('4. Teknoloji Karar Analizi (DAR)', level=1)

doc.add_paragraph(
    'Proje Adımı 2 kapsamında, Adım 1\'deki sistem analizi ve nesne tabanlı tasarım '
    '(Sınıf ve Sekans diyagramları) tamamlanan projenin kodlama (gerçekleme) aşamasında '
    'kullanılacak teknoloji yığınının (tech stack) belirlenmesi amacıyla karar analiz '
    'raporu (DAR) hazırlanmıştır. Geliştirme platformları, mimari frameworkler ve veri '
    'yönetim sistemleri; performans, ekip yetkinliği, güvenlik ve maliyet kriterleri '
    'doğrultusunda alternatifleriyle karşılaştırılarak analiz edilmiştir.'
)

doc.add_heading('4.1. Backend Dil ve Framework Seçimi', level=2)

doc.add_heading('Backend Dil Seçimi: Python', level=3)
doc.add_paragraph(
    'Proje Adımı 1\'deki sınıf diyagramında tasarlanan nesneler (User, Patient, Appointment) '
    've aralarındaki kalıtım/ilişki yapıları, nesne yönelimli bir dil gerektirmektedir. '
    'Java güçlü bir alternatif olsa da, Python\'un sunduğu hızlı prototipleme (rapid prototyping) '
    'yeteneği ve sağlık verisi şifreleme/hashing kütüphanelerinin zenginliği sebebiyle Python '
    'tercih edilmiştir.'
)

doc.add_heading('Backend Framework: Django REST Framework (DRF)', level=3)
doc.add_paragraph(
    'Adım 1\'deki sekans diyagramlarında yer alan kontrolcü nesnelerin (AuthController, '
    'AppointmentController) iş mantığı, Django\'nun MVC (Model-View-Template/Controller) '
    'tabanlı mimarisiyle birebir örtüşmektedir. AuthController senaryosu, DRF\'in yerleşik '
    'JWT (JSON Web Token) kütüphanesiyle doğrudan gerçeklenebilmektedir. Express.js\'te bu '
    'güvenlik katmanlarının sıfırdan yazılması gerekeceği için DRF seçilmiştir.'
)

doc.add_heading('4.2. Frontend Framework Seçimi: React.js', level=2)
doc.add_paragraph(
    'Sınıf diyagramındaki arayüz bileşenleri (PatientDialog, DashboardUI), React.js\'in '
    '"Component-Based" (Bileşen tabanlı) mimarisiyle doğrudan eşleşir. Özellikle projenin en '
    'kritik özelliği olan interaktif diş şemasının (Dijital Odontogram) DOM manipülasyonu, '
    'React\'ın Virtual DOM yapısı sayesinde performans kaybı yaşanmadan gerçekleştirilebilir. '
    'Frontend dil olarak TypeScript/JavaScript tercih edilmiştir; sekans diyagramlarında sınır '
    'sınıfları (Boundary/UI) üzerinden taşınan veri paketlerinin (Payload/DTO) tür güvenliğini '
    '(type-safety) sağlamak amacıyla JavaScript\'in süperseti olan TypeScript kullanılmıştır.'
)

doc.add_heading('4.3. Veri Tabanı Seçimi: PostgreSQL', level=2)
doc.add_paragraph(
    'Adım 1\'deki bütünsel sınıf diyagramında Clinic, Patient ve Appointment sınıfları arasında '
    'kesin yapısal ilişkiler (1..* Association ve Composition) bulunmaktadır. Bu kesin ilişki '
    'ağları ve sağlık/finansal verilerin gerektirdiği ACID (Atomicity, Consistency, Isolation, '
    'Durability) uyumluluğu nedeniyle ilişkisel veri tabanı (RDBMS) olan PostgreSQL seçilmiştir. '
    'MongoDB gibi şemasız (schemaless) sistemler veri bütünlüğü riskleri taşıdığı için elenmiştir.'
)

doc.add_paragraph(
    'ORM Aracı olarak Django ORM kullanılmıştır. Sınıf diyagramındaki nesnelerin veritabanı '
    'tablolarına eşlenmesi (Mapping) süreci Django ORM ile sağlanmakta; sınıflar arası '
    '"Bire-Çok" ilişkiler ForeignKey ile koda dökülmektedir.'
)

doc.add_heading('4.4. Karar Matrisleri', level=2)
doc.add_paragraph(
    'Aşağıdaki tablolarda teknoloji alternatifleri proje ihtiyaçlarına göre 1 (En Düşük) ile 5 '
    '(En Yüksek) arasında puanlanarak değerlendirilmiştir.'
)

doc.add_heading('4.4.1. Backend Framework Karar Matrisi', level=3)
add_styled_table(doc,
    ['Kriter', 'Django REST (Python)', 'Express.js (Node)', 'Spring Boot (Java)'],
    [
        ('Geliştirme Hızı / Öğrenme Eğrisi', '5', '4', '2'),
        ('Dahili Güvenlik (Auth/RBAC)', '5', '2', '4'),
        ('Ekip Yetkinliği', '5', '3', '2'),
        ('Sınıf Diyagramı (OOP) Uyumluluğu', '4', '3', '5'),
        ('Toplam Puan', '19', '12', '13'),
    ],
    [5.5, 3.5, 3.5, 3.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Sonuç: ')
run.bold = True
p.add_run('Güvenlik altyapısı ve ekibin hızı göz önüne alınarak Django REST seçilmiştir.')

doc.add_heading('4.4.2. Veri Tabanı Sistemleri Karar Matrisi', level=3)
add_styled_table(doc,
    ['Kriter', 'PostgreSQL (Relational)', 'MongoDB (NoSQL)', 'SQLite (Relational)'],
    [
        ('Veri Bütünlüğü (ACID Uyumu)', '5', '2', '4'),
        ('Sınıf/İlişki Modelleme Kolaylığı', '5', '3', '5'),
        ('Üretim (Production) Ölçeklenebilirliği', '5', '5', '1'),
        ('Medikal Veri Güvenliği Standartları', '5', '3', '2'),
        ('Toplam Puan', '20', '13', '12'),
    ],
    [5.5, 3.5, 3.5, 3.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Sonuç: ')
run.bold = True
p.add_run('İlişkisel yapının zorunluluğu ve ölçeklenebilirlik sebebiyle PostgreSQL seçilmiştir.')

doc.add_heading('4.4.3. Frontend Framework Karar Matrisi', level=3)
add_styled_table(doc,
    ['Kriter', 'React.js', 'Vue.js', 'Angular'],
    [
        ('Topluluk Desteği ve Kaynak', '5', '4', '3'),
        ('Odontogram için DOM Performansı', '5', '4', '3'),
        ('Ekip Yetkinliği', '5', '2', '1'),
        ('Sınır Sınıfları (Component) Uyumu', '5', '5', '5'),
        ('Toplam Puan', '20', '15', '12'),
    ],
    [5.5, 3.5, 3.5, 3.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Sonuç: ')
run.bold = True
p.add_run('DOM manipülasyon gücü ve ekip tecrübesi nedeniyle React.js seçilmiştir.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 5. İÇ VE DIŞ ARAYÜZLER
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('5. İç ve Dış Arayüzler', level=1)

doc.add_paragraph(
    'Bu bölüm, Proje Adımı 3 kapsamında gerçeklenen Opsiyon A (İç ve Dış Arayüzlerin '
    'Gerçeklenmesi ve Dokümantasyonu) çalışmasını içermektedir. "Yaşca Diş Kliniği Yönetim '
    'Sistemi", React.js (Frontend) ve Django REST Framework (Backend) kullanılarak geliştirilen '
    'API odaklı (API-first) bir SaaS projesidir. Sistemde yoğun bir arayüz/entegrasyon '
    'mimarisi bulunduğu için iç ve dış arayüzlerin detaylı kontratlarının tanımlanması ve '
    'OpenAPI (Swagger) ile dokümante edilmesi seçilmiştir.'
)

doc.add_heading('5.1. İç Arayüzler ve API Endpoint Kontratları', level=2)
doc.add_paragraph(
    'Projedeki katmanlar (Model, ViewSet, Serializer) birbirleriyle JSON paketleri üzerinden '
    'haberleşmektedir. Endüstri standardı olan Swagger (drf-spectacular) kullanılarak tüm API '
    'uç noktaları canlı olarak dokümante edilmiştir.'
)

# API endpoints
doc.add_heading('A. Kimlik Doğrulama: POST /api/auth/token/', level=3)
api_details = [
    ('Açıklama', 'Sisteme giriş yapılarak yetki jetonu (JWT) alınması.'),
    ('Request Body', '{"username": "admin_user", "password": "securepassword123"}'),
    ('Response (200 OK)', '{"access": "eyJhbGciOi...", "refresh": "eyJhbGciOi..."}'),
    ('Hata Kodları', '401 Unauthorized (Hatalı giriş), 400 Bad Request (Eksik parametre)'),
]
add_styled_table(doc, ['Parametre', 'Değer'], api_details, [4, 13])

doc.add_paragraph()

doc.add_heading('B. Hasta Yönetimi: GET & POST /api/patients/', level=3)
doc.add_paragraph('Sisteme yeni hasta kaydı yapılması ve listelenmesi (FR-02).')
doc.add_paragraph('POST Request Body:')
add_code_block(doc, '''{
  "first_name": "Ahmet",
  "last_name": "Yılmaz",
  "phone": "05551234567"
}''')

doc.add_paragraph('GET Response Örneği (200 OK):')
add_code_block(doc, '''[
  {
    "id": 1,
    "first_name": "Ahmet",
    "last_name": "Yılmaz",
    "full_name": "Ahmet Yılmaz",
    "phone": "05551234567",
    "tckn": null,
    "last_visit": "2026-05-20"
  }
]''')
doc.add_paragraph('Hata Kodları: 400 Bad Request (Zorunlu alan eksik).')

doc.add_heading('C. Randevu Yönetimi: POST /api/appointments/', level=3)
api_randevu = [
    ('Açıklama', 'Çakışma kontrollü yeni randevu oluşturulması (FR-03).'),
    ('Request Body', '{"patient": 15, "doctor": 3, "date": "2026-06-15", "time": "14:30:00"}'),
    ('Response (201 Created)', 'Başarıyla oluşturulan randevu nesnesi (id: 42).'),
    ('Hata Kodları', '400 Bad Request (Çakışma - Conflict), 403 Forbidden (Yetkisiz)'),
]
add_styled_table(doc, ['Parametre', 'Değer'], api_randevu, [4, 13])

doc.add_paragraph()

doc.add_heading('D. Günlük Özet (Dashboard): GET /api/dashboard/today/', level=3)
doc.add_paragraph(
    'Ana ekranda gösterilecek günlük aktif randevular, tamamlanan işlemler ve toplam '
    'hasta sayısı verilerini tek bir uç noktadan döner (FR-04).'
)
doc.add_paragraph('Response Body (200 OK):')
add_code_block(doc, '''{
  "today_appointments": [...],
  "today_total": 12,
  "today_completed": 4,
  "total_patients": 150
}''')

doc.add_heading('5.2. Gerçek Kod Örnekleri', level=2)

doc.add_heading('5.2.1. Model Katmanı (Entity) – Appointment Modeli', level=3)
add_code_block(doc, '''class Appointment(models.Model):
    class Status(models.TextChoices):
        SCHEDULED = "scheduled", "Planlandı"
        COMPLETED = "completed", "Tamamlandı"

    patient = models.ForeignKey(Patient, on_delete=models.CASCADE,
                                related_name="appointments")
    doctor = models.ForeignKey(CustomUser, on_delete=models.PROTECT)
    date = models.DateField("Tarih")
    time = models.TimeField("Saat")
    status = models.CharField(max_length=20, choices=Status.choices,
                              default=Status.SCHEDULED)''')

doc.add_heading('5.2.2. Serializer Katmanı – Çakışma Algoritması (Control)', level=3)
doc.add_paragraph('Hekimin aynı saatte başka bir randevusu varsa algoritma 400 hatası fırlatır:')
add_code_block(doc, '''class AppointmentSerializer(serializers.ModelSerializer):
    def validate(self, data):
        # FR-03: Aynı hekime aynı saatte randevu çakışması kontrolü
        doctor, date, time = data.get("doctor"), data.get("date"), data.get("time")
        existing = Appointment.objects.filter(
            doctor=doctor, date=date, time=time,
            status=Appointment.Status.SCHEDULED
        )
        if existing.exists():
            raise serializers.ValidationError(
                "Bu hekime bu saatte zaten randevu kayıtlı."
            )
        return data''')

doc.add_heading('5.2.3. ViewSet Katmanı (Controller)', level=3)
doc.add_paragraph('Gelen isteklerin karşılandığı, filtreleme ve iş mantığının yönetildiği sınıf:')
add_code_block(doc, '''class AppointmentViewSet(AuditLogMixin, viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]

    def get_serializer_class(self):
        if self.action in ("create", "update"):
            return AppointmentCreateSerializer
        return AppointmentSerializer

    def get_queryset(self):
        user = self.request.user
        qs = Appointment.objects.filter(is_active=True) \\
                                .select_related("patient", "doctor")
        if user.role == CustomUser.Role.DOCTOR and not user.is_superuser:
            qs = qs.filter(doctor=user)
        return qs.order_by("date", "time")''')

doc.add_page_break()

doc.add_heading('5.3. Dış Arayüzler (External Interfaces)', level=2)

doc.add_heading('5.3.1. Harici SMS Sağlayıcı Entegrasyonu', level=3)
doc.add_paragraph(
    'Randevu oluşturulduğunda harici SMS API\'sine istek atılarak hastaya bildirim gönderilir.'
)
sms_details = [
    ('Senaryo', 'Randevu oluşturulduğunda API\'ye istek atar.'),
    ('Dış Kontrat (Payload)', '{"api_key": "YOUR_SECRET_KEY", "phone_number": "+905551234567", "message": "Randevunuz onaylanmıştır."}'),
    ('Hata Yönetimi', 'Dış API yanıt vermezse (Timeout), işlem çökmez. Hata log sistemine kaydedilir ve randevu asenkron olarak tamamlanır.'),
]
add_styled_table(doc, ['Özellik', 'Detay'], sms_details, [4, 13])

doc.add_heading('5.3.2. Frontend (React) Entegrasyonu', level=3)
doc.add_paragraph(
    'Frontend uygulaması, sunucu ile axios kullanarak haberleşmektedir. JWT (Bearer Token) '
    'otomatik olarak başlığa (Header) eklenir.'
)
add_code_block(doc, '''// frontend/src/services/api.js
import axios from 'axios';

const api = axios.create({ baseURL: 'http://127.0.0.1:8000/api/' });

api.interceptors.request.use((config) => {
    const token = localStorage.getItem('accessToken');
    if (token) config.headers.Authorization = `Bearer ${token}`;
    return config;
});

export const fetchDashboardData = async () => {
    const response = await api.get('/dashboard/today/');
    return response.data;
};''')

doc.add_heading('5.3.3. Multi-Tenant (SaaS) Mimarisi', level=3)
doc.add_paragraph(
    'Uygulama django-tenants kütüphanesi kullanılarak çoklu kiracı yapısında kodlanmıştır. '
    'HeaderTenantMiddleware, HTTP Request içerisindeki domain/tenant kimliğini okur ve '
    'PostgreSQL şemasını o kliniğe yönlendirir. Böylece veri güvenliği fiziksel izolasyonla '
    'sağlanır.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 6. BİRİM TEST SONUÇLARI
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('6. Birim Test Sonuçları', level=1)

doc.add_paragraph(
    'Bu bölüm, Proje Adımı 4 kapsamında gerçekleştirilen birim test çalışmalarının '
    'sonuçlarını sunmaktadır. Sistemin çekirdek iş mantığını oluşturan sınıflara ait '
    'metotların doğruluğu birim (unit) test seviyesinde doğrulanmıştır.'
)

doc.add_heading('6.1. Test Ortamı ve Kullanılan Araçlar', level=2)

test_araclari = [
    ('pytest', 'Test koşum motoru', 'Python için endüstri standardı; az boilerplate, sade assert sözdizimi.'),
    ('pytest-django', 'Django entegrasyonu', '@pytest.mark.django_db ile her teste izole, otomatik geri alınan (rollback) veritabanı sağlar.'),
    ('factory-boy + Faker', 'Sahte test verisi üretimi', 'Test nesnelerini (Hasta, Doktor, Tedavi) tek satırda ve Türkçe veriyle üretir.'),
    ('unittest.mock / monkeypatch', 'Bağımlılık taklidi (mocking)', 'Veritabanı/dış bağımlılıkları taklit ederek metodu gerçekten izole eder.'),
]

add_styled_table(doc,
    ['Araç', 'Görevi', 'Neden Seçildi?'],
    test_araclari,
    [4, 3.5, 9.5]
)

doc.add_paragraph()
doc.add_paragraph(
    'Mocking (Taklit Etme) Yaklaşımı: Bağımlılıkların testi engellememesi için taklit '
    '(mocking) teknikleri kullanılmıştır. HeaderTenantMiddleware ve RegisterClinicView '
    'testleri, gerçek bir PostgreSQL/django-tenants kurulumuna ihtiyaç duymadan, kiracı '
    'modeli ve şema işlemleri monkeypatch ile taklit edilerek izole şekilde çalıştırılmıştır.'
)

doc.add_heading('6.2. Test Edilen Kritik Sınıflar', level=2)

doc.add_paragraph(
    'Sınıf diyagramından, zengin iş mantığı ve hata yönetimi içeren 5 kritik sınıf seçildi:'
)

kritik_siniflar = [
    ('1', 'AppointmentCreateSerializer', 'validate()', 'Randevu çakışması (aynı hekim/saat) engellenmeli.', 'Hayır'),
    ('2', 'HeaderTenantMiddleware', '__call__()', 'Her isteğin doğru kiracıya (tenant) yönlendirilmesi; izolasyonun temeli.', 'Evet'),
    ('3', 'RegisterClinicView / CheckDomainView', 'post() / get()', 'Yeni klinik kayıt akışı; hatalı/eksik girdilerin reddi.', 'Evet'),
    ('4', 'PatientSerializer', 'create(), update()', 'Hasta + anamnez kaydı; zorunlu alan doğrulaması.', 'Hayır'),
    ('5', 'TreatmentSerializer', 'validate()', 'Aynı gün/aynı dişe mükerrer tedavi girişinin engellenmesi.', 'Hayır'),
]

add_styled_table(doc,
    ['#', 'Sınıf', 'Hedef Metot', 'Neden Kritik?', 'Mocking'],
    kritik_siniflar,
    [1, 4.5, 2.5, 7, 2]
)

doc.add_heading('6.3. Test Tasarım Metodolojisi', level=2)
doc.add_paragraph('Her metot için iki senaryo türü tasarlanmıştır:')

add_bullet(doc, 'Pozitif (Başarılı) senaryo: Girdiler beklenen aralıkta olduğunda metodun doğru sonucu üretmesi (ör. boş randevu saatinin kabul edilmesi).')
add_bullet(doc, 'Negatif (Başarısız) senaryo: Hatalı/geçersiz girdilerle sistemin doğru şekilde hata fırlatması veya işlemi reddetmesi (ör. çakışan randevunun ValidationError ile reddi).')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Önemli: ')
run.bold = True
p.add_run(
    'Negatif senaryolarda, kodun doğru şekilde hata fırlatması/işlemi reddetmesi testin '
    'başarıyla sonuçlandığı anlamına gelir ve durum PASSED olarak işaretlenir.'
)

doc.add_page_break()

doc.add_heading('6.4. Birim Test Sonuç Matrisi', level=2)
doc.add_paragraph(
    'Koşum: pytest api/tests/test_serializers.py api/tests/test_middleware.py '
    'customers/tests/test_register.py → 54 passed. Aşağıda her sınıf için temsili '
    'pozitif/negatif satırlar listelenmiştir.'
)

# Test results table - split into groups for readability
test_results = [
    ('UT-001', 'AppointmentCreateSerializer', 'validate', 'Negatif', 'Aynı hekime, aynı tarih ve saatte ikinci randevu eklenmesi', 'PASSED'),
    ('UT-002', 'AppointmentSerializer', 'validate', 'Negatif', 'Çakışma kontrolünün güncelleme serializer\'ında da reddi', 'PASSED'),
    ('UT-003', 'AppointmentCreateSerializer', 'validate', 'Pozitif', 'Aynı hekime farklı saatte randevu eklenmesi', 'PASSED'),
    ('UT-004', 'AppointmentCreateSerializer', 'validate', 'Pozitif', 'İPTAL edilmiş randevunun aynı slotu bloke etmemesi', 'PASSED'),
    ('UT-005', 'AppointmentSerializer', 'validate', 'Pozitif', 'Mevcut randevu güncellenirken kendi slotunu çakışma saymaması', 'PASSED'),
    ('UT-006', 'HeaderTenantMiddleware', '__call__', 'Pozitif', 'X-Tenant header\'ı ile doğru kiracının çözülmesi (mock\'lu)', 'PASSED'),
    ('UT-007', 'HeaderTenantMiddleware', '__call__', 'Pozitif', 'Header yokken Host adından kiracı çözülmesi', 'PASSED'),
    ('UT-008', 'HeaderTenantMiddleware', '__call__', 'Pozitif', 'Header ve Host birlikte verildiğinde Header\'ın önceliği', 'PASSED'),
    ('UT-009', 'HeaderTenantMiddleware', '__call__', 'Negatif', 'Geçersiz/bilinmeyen X-Tenant değeri verilmesi', 'PASSED'),
    ('UT-010', 'HeaderTenantMiddleware', '__call__', 'Negatif', 'Boş/whitespace X-Tenant header gönderilmesi', 'PASSED'),
    ('UT-011', 'HeaderTenantMiddleware', '__call__', 'Negatif', 'Public tenant tanımsızken hiçbir eşleşme olmaması', 'PASSED'),
    ('UT-012', 'RegisterClinicView', 'post', 'Pozitif', 'Geçerli verilerle yeni klinik kaydı (mock\'lu şema)', 'PASSED'),
    ('UT-013', 'RegisterClinicView', 'post', 'Pozitif', 'Kayıt sonrası admin (superuser) kullanıcı oluşması', 'PASSED'),
    ('UT-014', 'RegisterClinicView', 'post', 'Negatif', 'Zorunlu alan (ör. admin_password) eksik gönderilmesi', 'PASSED'),
    ('UT-015', 'RegisterClinicView', 'post', 'Negatif', 'Zaten kayıtlı subdomain ile tekrar kayıt denenmesi', 'PASSED'),
    ('UT-016', 'RegisterClinicView', 'post', 'Negatif', 'Kayıt sırasında beklenmedik bir istisna oluşması (mock\'lu)', 'PASSED'),
    ('UT-017', 'CheckDomainView', 'get', 'Pozitif', 'Var olan subdomain için uygunluk sorgusu', 'PASSED'),
    ('UT-018', 'CheckDomainView', 'get', 'Negatif', 'Boş subdomain parametresi ile sorgu', 'PASSED'),
    ('UT-019', 'PatientSerializer', 'create', 'Pozitif', 'Anamnez olmadan yalnızca hasta kaydı oluşturma', 'PASSED'),
    ('UT-020', 'PatientSerializer', 'create', 'Pozitif', 'İç içe (nested) anamnez verisiyle hasta oluşturma', 'PASSED'),
    ('UT-021', 'PatientSerializer', 'update', 'Pozitif', 'Hastada anamnez yoksa güncelleme de oluşturulması', 'PASSED'),
    ('UT-022', 'PatientSerializer', 'create', 'Negatif', 'Telefon (zorunlu alan) gönderilmeden kayıt denenmesi', 'PASSED'),
    ('UT-023', 'PatientSerializer', 'create', 'Negatif', 'Boş ad (first_name) ile kayıt denenmesi', 'PASSED'),
    ('UT-024', 'TreatmentSerializer', 'validate', 'Pozitif', 'Aynı gün, aynı hasta, FARKLI diş numarasına tedavi', 'PASSED'),
    ('UT-025', 'TreatmentSerializer', 'validate', 'Negatif', 'Aynı gün, aynı diş, aynı tedavi türünün mükerrer eklenmesi', 'PASSED'),
]

add_styled_table(doc,
    ['Test ID', 'Hedef Sınıf', 'Metot', 'Tür', 'Senaryo Açıklaması', 'Durum'],
    test_results,
    [1.5, 3.5, 1.5, 1.5, 7, 1.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Özet: ')
run.bold = True
p.add_run(
    '25 temsili senaryo — 13 Pozitif + 12 Negatif — Tümü PASSED. '
    'İlgili 3 test dosyasının tam koşumu: 54 passed.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 7. BAKIM VE GELECEK ÖNERİLERİ
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('7. Bakım ve Gelecek Önerileri', level=1)

doc.add_paragraph(
    'Dersin "Bakım" ayağına uygun olarak sistemin gelecekteki olası güncelleme '
    'ihtiyaçları ve teknik borçları bu bölümde analiz edilir.'
)

doc.add_heading('7.1. Refactoring İhtiyaçları', level=2)
refactoring = [
    ('Servis Katmanı Ayrımı', 
     'Mevcut durumda iş mantığının bir kısmı ViewSet ve Serializer katmanlarında yer almaktadır. '
     'İlerleyen aşamalarda, karmaşık iş kurallarını (randevu çakışması, ödeme doğrulaması) '
     'bağımsız bir Service Layer\'a taşımak bakım maliyetini düşürecektir.'),
    ('Test Kapsamının Genişletilmesi',
     'Mevcut birim testler 5 kritik sınıfı kapsamaktadır. Entegrasyon testleri (API seviyesinde) '
     've uçtan uca (E2E) testlerin eklenmesi ile test kapsamı artırılmalıdır.'),
    ('Hata Yönetimi Merkezi Yapısı',
     'API hatalarının standart bir formatta (RFC 7807 Problem Details) döndürülmesi için '
     'merkezi bir exception handler yapısına geçiş planlanmaktadır.'),
]

for title, desc in refactoring:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_heading('7.2. Yeni Özellik Önerileri', level=2)
yeni_ozellikler = [
    'Gerçek SMS/WhatsApp entegrasyonu ile otomatik randevu hatırlatmaları',
    'Randevu takviminde sürükle-bırak (drag & drop) desteği',
    'Dijital röntgen ve görüntüleme dosyalarının hasta profiline eklenmesi',
    'Raporlama modülü: Aylık/yıllık gelir, hasta istatistikleri, tedavi dağılımları',
    'Mobil uygulama (React Native) ile hekimlerin uzaktan erişimi',
    'Çok dilli destek (i18n) ile farklı ülkelere açılma imkânı',
]
for o in yeni_ozellikler:
    add_bullet(doc, o)

doc.add_heading('7.3. Versiyon Güncelleme Planı', level=2)

versiyon_plani = [
    ('v1.1', 'SMS entegrasyonu, takvim drag & drop', '2025 Q3'),
    ('v1.2', 'Raporlama modülü, PDF çıktı', '2025 Q4'),
    ('v2.0', 'Mobil uygulama, çok dilli destek', '2026 Q1'),
    ('v2.1', 'Röntgen entegrasyonu, gelişmiş analitik', '2026 Q2'),
]

add_styled_table(doc,
    ['Versiyon', 'Planlanan Özellikler', 'Tahmini Tarih'],
    versiyon_plani,
    [3, 10, 4]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 8. SONUÇ
# ═══════════════════════════════════════════════════════════════════════════

doc.add_heading('8. Sonuç', level=1)

doc.add_paragraph(
    'Bu proje süreci boyunca, "Yaşca Diş Kliniği Yönetim Sistemi" adlı çok-kiracılı '
    '(multi-tenant) SaaS uygulaması, yazılım mühendisliğinin temel disiplinleri olan '
    'analiz, tasarım, gerçekleme, test ve bakım aşamalarının tamamını kapsayacak şekilde '
    'geliştirilmiştir.'
)

doc.add_paragraph(
    'Proje Adımı 1\'de gerçekleştirilen sistem analizi ve UML modelleme çalışmasıyla, '
    '6 fonksiyonel gereksinim (FR-01 ila FR-06) detaylı sekans ve sınıf diyagramları ile '
    'modellenmiştir. Boundary-Control-Entity ayrımına dayanan bu modelleme, ilerleyen '
    'aşamalarda kod mimarisinin temelini oluşturmuştur.'
)

doc.add_paragraph(
    'Proje Adımı 2\'de hazırlanan Teknoloji Karar Analizi (DAR) raporu ile Django REST '
    'Framework (Python), React.js (TypeScript) ve PostgreSQL\'den oluşan teknoloji yığını, '
    'karar matrisleri kullanılarak objektif kriterlerle seçilmiştir.'
)

doc.add_paragraph(
    'Proje Adımı 3\'te iç ve dış arayüzlerin gerçeklenmesi ve dokümantasyonu (Opsiyon A) '
    'kapsamında, RESTful API kontratları Swagger (drf-spectacular) ile canlı olarak '
    'dokümante edilmiş; frontend-backend entegrasyonu ve multi-tenant altyapısı '
    'kodlanmıştır.'
)

doc.add_paragraph(
    'Proje Adımı 4\'te gerçekleştirilen birim test çalışmasında, sınıf diyagramından '
    'seçilen 5 kritik sınıfın çekirdek metotları, hem pozitif hem negatif senaryolarla '
    'test edilmiştir. Bağımlılıklar (veritabanı, kiracı modeli, şema işlemleri) '
    'monkeypatch/MagicMock ile taklit edilerek metotlar gerçekten izole edilmiştir. '
    'İlgili test dosyalarının tamamı 54/54 PASSED ile sonuçlanmıştır.'
)

doc.add_heading('Temel Çıktılar', level=2)
ciktilar = [
    '✓ Uçtan uca test edilmiş modüller',
    '✓ Standartlara uygun UML dokümantasyonu',
    '✓ Teknik karar verme yetkinliği',
    '✓ Multi-tenant SaaS mimarisi ile ölçeklenebilir altyapı',
    '✓ Swagger ile canlı API dokümantasyonu',
    '✓ 54/54 birim testin başarıyla geçmesi',
]
for c in ciktilar:
    add_bullet(doc, c)

doc.add_paragraph()
doc.add_paragraph(
    'Bu proje, yazılım geliştirme yaşam döngüsünün tüm aşamalarını kapsayan '
    'kapsamlı bir mühendislik çalışması olarak tamamlanmıştır. Edinilen deneyimler '
    've oluşturulan altyapı, sistemin gelecekte genişletilmesi için sağlam bir temel '
    'oluşturmaktadır.'
)

# ═══════════════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════════════

output_path = os.path.join(r"c:\Users\Ali\yasca-dental-clinic", "Grup4_Final_Proje_Raporu.docx")
doc.save(output_path)
print(f"Rapor basariyla olusturuldu: {output_path}")
