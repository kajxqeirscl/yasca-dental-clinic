# -*- coding: utf-8 -*-
"""
Yasca Bitirme Projesi - ISTUN Tez Yazim Sablonuna Uygun Rapor
5 ANA BOLUM: Giris | Kuramsal Temeller | Materyal ve Yontem | Bulgular | Sonuc ve Oneriler
34 konu bu 5 bolumun alt basliklari olarak yerlesti.
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = r"c:\Users\Ali\yasca-dental-clinic"
SS = r"C:\Users\Ali\.gemini\antigravity-ide\brain\574ceb4e-bc73-4dce-b3e0-c86fd8757753"
DIAG = os.path.join(BASE, "docs", "diagrams")
OUT = os.path.join(BASE, "Grup4_Yasca_Proje_Raporu.docx")

C = WD_ALIGN_PARAGRAPH.CENTER
L = WD_ALIGN_PARAGRAPH.LEFT
R = WD_ALIGN_PARAGRAPH.RIGHT

def shd(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def p(doc, text, sz=12, b=False, al=None, sa=6, sb=0, ind=None, color=None, it=False):
    pa = doc.add_paragraph(); r = pa.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(sz)
    if b: r.bold = True
    if it: r.italic = True
    if color: r.font.color.rgb = RGBColor(*color)
    if al is not None: pa.alignment = al
    pf = pa.paragraph_format; pf.space_after = Pt(sa); pf.space_before = Pt(sb)
    if ind: pf.first_line_indent = Cm(ind)
    return pa

def h1(doc, text):
    hd = doc.add_heading(text, level=1)
    for r in hd.runs: r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    hd.paragraph_format.space_before=Pt(18); hd.paragraph_format.space_after=Pt(10)
    return hd

def h2(doc, text):
    hd = doc.add_heading(text, level=2)
    for r in hd.runs: r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    hd.paragraph_format.space_before=Pt(12); hd.paragraph_format.space_after=Pt(6)
    return hd

def h3(doc, text):
    hd = doc.add_heading(text, level=3)
    for r in hd.runs: r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    hd.paragraph_format.space_before=Pt(8); hd.paragraph_format.space_after=Pt(4)
    return hd

def tbl(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = hdr
        for pp in c.paragraphs:
            pp.alignment = C
            for rr in pp.runs: rr.bold=True; rr.font.name='Times New Roman'; rr.font.size=Pt(10)
        shd(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(ct)
            for pp in c.paragraphs:
                for rr in pp.runs: rr.font.name='Times New Roman'; rr.font.size=Pt(10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows: row.cells[i].width = Cm(w)
    return t

def img(doc, path, w=5.5, cap=None):
    if not os.path.exists(path):
        p(doc, f"[Resim bulunamadi: {os.path.basename(path)}]", sz=10, color=(255,0,0)); return
    try:
        doc.add_picture(path, width=Inches(w)); doc.paragraphs[-1].alignment = C
        if cap: p(doc, cap, sz=10, al=C, sa=12, b=True)
    except Exception as e:
        p(doc, f"[Resim hatasi: {e}]", sz=10, color=(255,0,0))

def code(doc, text):
    pa = doc.add_paragraph(); r = pa.add_run(text); r.font.name='Consolas'; r.font.size=Pt(9)
    pa.paragraph_format.space_before=Pt(6); pa.paragraph_format.space_after=Pt(6)
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), 'F5F5F5'); s.set(qn('w:val'), 'clear')
    pa._element.get_or_add_pPr().append(s)

def bul(doc, text, sz=11):
    p(doc, text, sz=sz, sa=3)

# ============================================================================
doc = Document()
for sec in doc.sections:
    sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
style = doc.styles['Normal']; style.font.name='Times New Roman'; style.font.size=Pt(12)

# ==================== KAPAK ====================
p(doc, "T.C.", sz=14, b=True, al=C, sb=36)
p(doc, "ISTANBUL SAGLIK VE TEKNOLOJI UNIVERSITESI", sz=14, b=True, al=C)
p(doc, "MUHENDISLIK VE DOGA BILIMLERI FAKULTESI", sz=14, b=True, al=C, sa=24)
p(doc, "LISANS BITIRME PROJESI", sz=16, b=True, al=C, sa=36)
p(doc, "YASCA: COKLU KIRACILI DIS KLINIGI", sz=16, b=True, al=C, sa=3)
p(doc, "YONETIM SaaS PLATFORMU", sz=16, b=True, al=C, sa=12)
p(doc, "", sz=12, sa=6)
p(doc, "Ders: Python Projelerinde Yapay Zeka Kullanimi", sz=13, b=True, al=C, sa=48)
for nm in ["Yaman Halloum", "Ali Ure", "Cihan Kurtbey", "Sukru Yesilmen"]:
    p(doc, nm, sz=12, al=C, sa=3)
p(doc, "", sz=12, sa=24)
p(doc, "Danismani: Ogr. Gor. Oguz Oztoprak", sz=12, al=C)
p(doc, "Yazilim Muhendisligi Bolumu", sz=12, al=C, sa=48)
p(doc, "ISTANBUL - 2025", sz=14, b=True, al=C)
doc.add_page_break()

# ==================== ONAY ====================
p(doc, "T.C.", sz=12, b=True, al=C, sb=24)
p(doc, "ISTANBUL SAGLIK VE TEKNOLOJI UNIVERSITESI", sz=12, b=True, al=C)
p(doc, "MUHENDISLIK VE DOGA BILIMLERI FAKULTESI", sz=12, b=True, al=C, sa=18)
p(doc, "BITIRME PROJESI ONAYI", sz=14, b=True, al=C, sa=18)
p(doc, "Hazirlayan: Yaman Halloum, Ali Ure, Cihan Kurtbey, Sukru Yesilmen", sz=12)
p(doc, "Proje Basligi: Yasca - Coklu Kiracili Dis Klinigi Yonetim SaaS Platformu", sz=12)
p(doc, "Sinav Tarihi: ......./......./2025", sz=12, sa=24)
p(doc, "Danismani ve Bitirme projesi sinav jurisi degerlendirmesi sonucu ogrenci ... bulunmustur.", sz=12, sa=18)
t = doc.add_table(rows=4, cols=3); t.style='Table Grid'
for i, hdr in enumerate(["Unvan, Ad Soyad", "Gorev", "Imza"]):
    t.rows[0].cells[i].text = hdr
    for pp in t.rows[0].cells[i].paragraphs:
        for rr in pp.runs: rr.bold=True; rr.font.name='Times New Roman'; rr.font.size=Pt(11)
    shd(t.rows[0].cells[i], 'D9E2F3')
for ri, (n, g) in enumerate([("Ogr. Gor. Oguz Oztoprak","Danismani"),("","Juri Uyesi"),("","Juri Uyesi")]):
    t.rows[ri+1].cells[0].text=n; t.rows[ri+1].cells[1].text=g
doc.add_page_break()

# ==================== BEYAN ====================
p(doc, "BEYAN", sz=14, b=True, al=C, sb=24, sa=18)
p(doc, "Bu projenin bize ait oldugunu, tum asamalarinda etik disi davranisimizin olmadigini, icinde yer alan butun bilgileri akademik ve etik kurallar icinde elde ettigimizi, kullanmis oldugumuz butun bilgilere kaynak gosterdigimizi, bu projenin Istanbul Saglik ve Teknoloji Universitesi Muhendislik ve Doga Bilimleri Fakultesi Bitirme Projesi Yonergesine uygun olarak hazirlandigini beyan ederiz.", sz=12, ind=1.25, sa=36)
for nm in ["Yaman Halloum", "Ali Ure", "Cihan Kurtbey", "Sukru Yesilmen"]:
    p(doc, f"{nm} .......................", sz=12, sa=6)
doc.add_page_break()

# ==================== ONSOZ ====================
p(doc, "ONSOZ", sz=14, b=True, al=C, sb=24, sa=18)
p(doc, "Bu bitirme projesi kapsaminda, dis klinikleri icin modern ve olceklenebilir bir SaaS yonetim platformu olan Yasca'yi gelistirdik. Proje suresince yazilim muhendisliginin temel prensiplerini uygulama firsati bulduk ve yapay zeka araclarini gelistirme sureclerimize entegre ederek verimlilik ve kalite konusunda onemli kazanimlar elde ettik.", sz=12, ind=1.25)
p(doc, "Projemizin her asamasinda bize rehberlik eden danismanmiz Ogr. Gor. Oguz Oztoprak'a, bilgi ve deneyimlerini paylasan hocalarimiza ve destek veren ailelerimize tesekkur ederiz.", sz=12, ind=1.25)
p(doc, "Haziran 2025, Istanbul", sz=12, al=R, sb=18)
p(doc, "Grup 4", sz=12, al=R)
doc.add_page_break()

# ==================== KISALTMALAR ====================
p(doc, "KISALTMALAR", sz=14, b=True, al=C, sb=24, sa=18)
tbl(doc, ["Kisaltma", "Aciklama"], [
    ["SaaS","Software as a Service"], ["API","Application Programming Interface"],
    ["REST","Representational State Transfer"], ["DRF","Django REST Framework"],
    ["JWT","JSON Web Token"], ["RBAC","Role-Based Access Control"],
    ["CRUD","Create, Read, Update, Delete"], ["CI/CD","Continuous Integration / Continuous Delivery"],
    ["E2E","End-to-End Test"], ["A11y","Accessibility (Erisilebilirlik)"],
    ["WCAG","Web Content Accessibility Guidelines"], ["FDI","Uluslararasi Dis Numaralandirma Sistemi"],
    ["KVKK","Kisisel Verilerin Korunmasi Kanunu"], ["ORM","Object-Relational Mapping"],
    ["PR","Pull Request"], ["ADR","Architecture Decision Record"],
], widths=[3, 13])
doc.add_page_break()

# ==================== SEKIL LISTESI ====================
p(doc, "SEKIL LISTESI", sz=14, b=True, al=C, sb=24, sa=24)
p(doc, "[Bu bolum Word'de otomatik olusturulacaktir]", sz=11, al=C, color=(128,128,128))
doc.add_page_break()

# ==================== CIZELGE LISTESI ====================
p(doc, "CIZELGE LISTESI", sz=14, b=True, al=C, sb=24, sa=24)
p(doc, "[Bu bolum Word'de otomatik olusturulacaktir]", sz=11, al=C, color=(128,128,128))
doc.add_page_break()

# ==================== OZET ====================
p(doc, "OZET", sz=14, b=True, al=C, sb=24, sa=12)
p(doc, "YASCA: COKLU KIRACILI DIS KLINIGI YONETIM SaaS PLATFORMU", sz=12, b=True, al=C, sa=12)
p(doc, "Bu proje, dis kliniklerinin operasyonel is akislarini dijitallestirmek, veri guvenligini saglamak ve klinik verimliligini artirmak amaciyla gelistirilen acik kaynakli, coklu kiracili (multi-tenant) bir SaaS uygulamasidir. Yasca, Django REST Framework tabanli backend ve React.js tabanli frontend ile gelistirilmis olup, django-tenants kutuphanesi kullanilarak her klinik icin izole PostgreSQL sema yapisi sunmaktadir. Sistem; hasta, randevu, tedavi, odeme yonetimi, dental sema (FDI), rol bazli erisim kontrolu (RBAC), islem gecmisi (Audit Log) ve coklu dil destegi gibi kapsamli ozellikler icermektedir. Proje suresince yapay zeka araclari aktif olarak kullanilmis; kod uretimi, test yazimi, hata ayiklama ve dokumantasyon sureclerinde onemli verimlilik artislari saglanmistir. 403 adet test ile %78 backend ve %57 frontend coverage degerlerine ulasilmistir.", sz=12, ind=1.25)
p(doc, "Anahtar Kelimeler: Multi-tenant SaaS, Django REST Framework, React.js, Yapay Zeka, Dis Klinigi Yonetimi", sz=12, sb=12, b=True)
doc.add_page_break()

# ==================== ABSTRACT ====================
p(doc, "ABSTRACT", sz=14, b=True, al=C, sb=24, sa=12)
p(doc, "YASCA: MULTI-TENANT DENTAL CLINIC MANAGEMENT SaaS PLATFORM", sz=12, b=True, al=C, sa=12)
p(doc, "This project is an open-source, multi-tenant SaaS application developed to digitize the operational workflows of dental clinics, ensure data security, and increase clinical efficiency. Yasca is built with a Django REST Framework backend and React.js frontend, utilizing the django-tenants library to provide isolated PostgreSQL schema structures for each clinic. Throughout the development process, AI tools (GitHub Copilot, ChatGPT, Cursor AI, Claude Code) were actively employed, yielding significant productivity gains. A total of 403 tests were created with 78% backend and 57% frontend code coverage.", sz=12, ind=1.25)
p(doc, "Keywords: Multi-tenant SaaS, Django REST Framework, React.js, Artificial Intelligence, Dental Clinic Management", sz=12, sb=12, b=True)
doc.add_page_break()

# ==================== ICINDEKILER ====================
p(doc, "I C I N D E K I L E R", sz=14, b=True, al=C, sb=24, sa=24)
p(doc, "[Bu bolum Word'de otomatik olusturulacaktir - Referanslar > Icindekiler Tablosu]", sz=11, al=C, color=(128,128,128))
doc.add_page_break()

# ====================================================================
# BOLUM 1: GIRIS  (Sablon: Bolum 1)
# ====================================================================
h1(doc, "1. GIRIS")

# 1.1 Giris-Tanim (Konu #2)
h2(doc, "1.1 Giris ve Tanim")
p(doc, "Bu proje, 'Python Projelerinde Yapay Zeka Kullanimi' dersi kapsaminda hazirlanan bir bitirme projesidir. Proje, dis kliniklerinin gunluk operasyonlarini dijital ortama tasiyarak klinik verimliligini artiran, veri guvenligini saglayan ve olceklenebilir bir yapida hizmet sunan modern bir SaaS (Software as a Service) platformu gelistirmeyi amaclamaktadir.", sz=12, ind=1.25)
p(doc, "Yasca adi verilen bu platform, coklu kiracili (multi-tenant) mimarisi sayesinde tek bir yazilim kurulumu uzerinden birden fazla dis klinigine es zamanli olarak hizmet verebilmektedir. Her klinik icin fiziksel duzey veri izolasyonu saglayan PostgreSQL schema-based yaklasim kullanilmaktadir.", sz=12, ind=1.25)
p(doc, "Projenin kapsami; hasta kaydi ve anamnez yonetimi, randevu planlama ve cakisma kontrolu, tedavi ve odeme takibi, dental sema (FDI dis numaralandirma) sistemi, dokuman yonetimi, rol bazli erisim kontrolu (RBAC), islem gecmisi (Audit Log) kayitlari, klinik ayarlari yonetimi ve cok dilli destek (Turkce/Ingilizce) gibi temel is sureclerini kapsamaktadir.", sz=12, ind=1.25)

# 1.2 Projenin Onemi
h2(doc, "1.2 Projenin Anlami ve Onemi")
p(doc, "Turkiye'de bircok dis klinigi hala manuel kayit sistemleri veya yetersiz yazilim cozumleri kullanmaktadir. Bu durum, veri kaybi, randevu karmasasi, odeme takip sorunlari ve KVKK uyumsuzlugu gibi ciddi problemlere yol acmaktadir. Yasca, bu sorunlari cozmek amaciyla tasarlanmis olup, saglik yazilimlarinda beklenen yuksek guvenlik standartlarini (veri izolasyonu, Audit Trail, sifreleme) karsilamaktadir.", sz=12, ind=1.25)
p(doc, "Ek olarak, proje gelistirme surecinde yapay zeka araclarinin (GitHub Copilot, ChatGPT, Cursor AI, Claude Code, Gemini) etkin kullaniminin yazilim gelistirme sureclerine etkisi de calismada detayli olarak ele alinmistir.", sz=12, ind=1.25)
doc.add_page_break()

# ====================================================================
# BOLUM 2: KURAMSAL TEMELLER VE LITERATUR TARAMASI  (Sablon: Bolum 2)
# Alt basliklar: Konu #3-8
# ====================================================================
h1(doc, "2. KURAMSAL TEMELLER VE LITERATUR TARAMASI")

# 2.1 Proje Yaklasimi ve Mimari Baglam (Konu #3)
h2(doc, "2.1 Proje Yaklasimi ve Mimari Baglam")
p(doc, "Yasca projesi, API-First monolitik mimari yaklasimini benimsemistir. Backend tarafinda Django monoliti tek bir servis olarak calismakta, frontend ise ayri bir React SPA (Single Page Application) olarak backend API'larini tuketmektedir. Bu yaklasim, kucuk-orta olcekli takimlar icin hizli gelistirme dongusu saglamakta ve operasyonel karmasikligi azaltmaktadir.", sz=12, ind=1.25)

# 2.1.1 Mimari Yaklasimi (Konu #4)
h3(doc, "2.1.1 Mimari Yaklasimi (Monorepo)")
p(doc, "Proje monolitik bir yaklasim benimsemistir. Mikro-servis mimarisi yerine monolitin tercih edilme sebepleri:", sz=12, ind=1.25)
for item in ["4 kisilik kucuk takim icin operasyonel karmasikligin dusuk tutulmasi",
    "Tek deployment birimi ile bakim kolayligi", "Django-tenants kutuphanesinin monolitik yapiyla uyumu",
    "Gelistirme hizi: Tek repo (monorepo) ile hizli iterasyon", "API-First yaklasim: Frontend ve backend arasinda acik REST kontrati"]:
    bul(doc, f"- {item}")
p(doc, "Multi-tenant yapi, django-tenants kutuphanesi araciligiyla PostgreSQL'in schema-based isolation mekanizmasi uzerine insa edilmistir. Her yeni klinik kaydedildiginde otomatik olarak izole bir veritabani semasi olusturulmaktadir.", sz=12, ind=1.25, sb=6)

# 2.2 Kullanilacak Ana Teknolojiler (Konu #5)
h2(doc, "2.2 Kullanilacak Ana Teknolojiler")
p(doc, "Projede kullanilan teknoloji yigini asagidaki tabloda ozetlenmistir.", sz=12, ind=1.25)
p(doc, "Cizelge 2.1: Projede Kullanilan Teknoloji Yigini", sz=10, b=True, al=C, sb=12)
tbl(doc, ["Katman", "Teknoloji", "Surum", "Kullanim Amaci"], [
    ["Backend","Python","3.12","Ana programlama dili"], ["Backend","Django","5.2+","Web framework"],
    ["Backend","DRF","3.14+","REST API"], ["Backend","django-tenants","3.7+","Multi-tenant izolasyon"],
    ["Backend","SimpleJWT","5.3+","JWT kimlik dogrulama"], ["Backend","drf-spectacular","0.27+","OpenAPI/Swagger"],
    ["Frontend","React","18.3","UI framework"], ["Frontend","TypeScript","6.0+","Tip guvenligi"],
    ["Frontend","Vite","6.3","Build araci"], ["Frontend","Tailwind CSS","4.1","CSS framework"],
    ["Frontend","Radix UI","Cesitli","UI primitifleri"], ["Veritabani","PostgreSQL","15","Iliskisel veritabani"],
    ["Test","pytest / Vitest / Playwright","-","Test motorlari"], ["DevOps","Docker + GitHub Actions","-","CI/CD"],
    ["Deploy","Vercel + Render","-","Hosting"],
], widths=[2.5, 3.5, 2, 5.5])

# 2.2.1 Django (Konu #6)
h3(doc, "2.2.1 Django")
p(doc, "Django 5.2, Python tabanli yuksek seviyeli bir web framework'udur. ORM katmani, veritabani islemlerini Python siniflari uzerinden yonetmeyi saglar. Django REST Framework (DRF) ise RESTful API gelistirmeyi kolaylastiran, serializer, viewset ve permission siniflari sunan guclu bir eklentidir. Projede ek olarak: django-tenants (multi-tenant), djangorestframework-simplejwt (JWT), drf-spectacular (OpenAPI/Swagger), django-cors-headers (CORS), django-anymail (e-posta) kutuphaneleri kullanilmistir.", sz=12, ind=1.25)

# 2.2.2 React (Konu #7)
h3(doc, "2.2.2 React")
p(doc, "Frontend tarafinda React 18, TypeScript ile birlikte kullanilmistir. Build araci olarak Vite, stilizasyon icin Tailwind CSS v4, UI bilesen kutuphanesi olarak Radix UI primitifleri, ikon seti olarak Lucide React, animasyonlar icin Motion, form yonetimi icin React Hook Form, grafik gorsellestirmesi icin Recharts kullanilmistir. Cok dilli destek i18next ile saglanmistir.", sz=12, ind=1.25)

# 2.2.3 Veritabani (Konu #8)
h3(doc, "2.2.3 Veritabani")
p(doc, "Veritabani olarak PostgreSQL 15 kullanilmistir. django-tenants kutuphanesi, PostgreSQL'in CREATE SCHEMA ozelligini kullanarak her kiraciya izole bir sema olusturur. Paylasilan veriler (Client, Domain) 'public' semasinda, kiraciya ozel veriler (Patient, Appointment vb.) ilgili kiracinin semasinda tutulur. Bu yaklasim, veri izolasyonunu fiziksel duzeyde garanti eder.", sz=12, ind=1.25)
doc.add_page_break()

# ====================================================================
# BOLUM 3: MATERYAL VE YONTEM  (Sablon: Bolum 3)
# Alt basliklar: Konu #9-17, 27, 28
# ====================================================================
h1(doc, "3. MATERYAL VE YONTEM")

# 3.1 Ekip Yapisi (Konu #9)
h2(doc, "3.1 Ekip Yapisi ve Gorev Dagilimi")
p(doc, "Cizelge 3.1: Ekip Uyeleri ve Sorumluluk Alanlari", sz=10, b=True, al=C, sb=12)
tbl(doc, ["Uye", "Ana Sorumluluk", "Detay"], [
    ["Yaman Halloum","Backend Gelistirici","Django/DRF, multi-tenant altyapi, model ve API gelistirme"],
    ["Ali Ure","Full-Stack Gelistirici","Mimari kurulum, tam yigin (full-stack) gelistirme, dagitim (Vercel/Render)"],
    ["Cihan Kurtbey","Frontend Gelistirici","Kimlik dogrulama akisi, sifre sifirlama, e-posta entegrasyonu"],
    ["Sukru Yesilmen","Kalite Guvence","Test, dokumantasyon ve kalite guvencesi katkilari"],
], widths=[3.5, 4, 8])

# 3.2 Yazilim Gelistirme Sureci (Konu #10)
h2(doc, "3.2 Yazilim Gelistirme Sureci")
p(doc, "Proje, Scrum'dan esinlenen cevik (Agile) bir gelistirme sureci ile yurutulmustur. Gelistirme dongusu: Planlama > Gelistirme > Test > Code Review > Merge > Deploy seklinde yapilandirilmistir.", sz=12, ind=1.25)
for item in ["Haftalik Sprint Planning: Her hafta basinda gorev belirleme ve onceliklendirme",
    "GitHub Issues: Gorev takibi, bug raporlama ve ozellik istekleri",
    "GitHub Projects (Kanban): Gorsel is akisi takibi (To Do > In Progress > Review > Done)",
    "Pull Request (PR) bazli code review: Her PR en az 1 onay gerektirir",
    "Conventional Commits formati: feat:, fix:, test:, docs:, refactor:, chore:",
    "Pre-commit hook'lari (Husky): ESLint, Prettier, TypeScript type-check otomatik kontrol"]:
    bul(doc, f"- {item}")

# 3.3 Teknik Dokumantasyon Yonetimi (Konu #11)
h2(doc, "3.3 Teknik Dokumantasyon Yonetimi")
p(doc, "Proje dokumantasyonu, kodla birlikte yasayan (living documentation) yaklasimi ile yonetilmistir:", sz=12, ind=1.25)
for nm, desc in [("README.md","Proje tanitimi, kurulum, calistirma talimatlari"),("docs/TESTING.md","Kapsamli test rehberi"),
    ("docs/CONTRIBUTING.md","PR sureci, code review kontrol listesi"),("docs/TEST_PYRAMID.md","Test piramidi hedefi"),
    ("docs/TEST_METRICS.md","Otomatik guncellenen test metrikleri"),("docs/adr/0001","Test stratejisi mimari karar kaydi"),
    ("docs/adr/0002","Multi-tenant izolasyon karari"),("docs/adr/0003","A11y sifir tolerans politikasi"),
    (".github/PULL_REQUEST_TEMPLATE.md","PR kontrol listesi sablonu")]:
    bul(doc, f"- {nm}: {desc}")

# 3.4 Teknoloji ve Urun Yapilari (Konu #12)
h2(doc, "3.4 Teknoloji ve Urun Yapilari")

# 3.4.1 Frontend Yapisi (Konu #13)
h3(doc, "3.4.1 Frontend Yapisi")
p(doc, "Frontend uygulamasi React.js ve TypeScript ile gelistirilmis olup Vite build araci kullanilmaktadir:", sz=12, ind=1.25)
code(doc, """frontend/src/
+-- app/
|   +-- App.tsx              # Ana yonlendirici
|   +-- ClinicApp.tsx        # Klinik yonetim paneli
|   +-- components/          # 32+ UI bileseni
|   +-- contexts/            # React Context (Auth, Theme)
|   +-- hooks/               # Custom React hooks
|   +-- services/            # API servis katmani (api.ts)
+-- locales/                 # i18n ceviri dosyalari (TR/EN)
+-- mocks/                   # MSW mock handlers
+-- test/                    # Test yardimcilari""")

# 3.4.2 Backend Yapisi (Konu #14)
h3(doc, "3.4.2 Backend, Veritabani ve Guvenlik Yapisi")
code(doc, """backend/
+-- api/                     # Ana uygulama modulu
|   +-- models.py            # 9 veri modeli
|   +-- serializers.py       # 12 serializer
|   +-- views.py             # 15 view/viewset
|   +-- middleware.py        # 3 middleware
|   +-- permissions.py       # 2 RBAC sinifi
|   +-- mixins.py            # AuditLog mixin
+-- core/                    # Django ayarlari
+-- customers/               # Tenant yonetimi""")
p(doc, "Guvenlik Onlemleri:", sz=12, b=True, sb=6)
for item in ["JWT tabanli kimlik dogrulama (Access 60dk + Refresh 7gun)",
    "Rol bazli erisim kontrolu: Admin, Doktor, Asistan", "Schema-based veri izolasyonu",
    "AuditLog ile tum CRUD islemlerinin kayit altina alinmasi (KVKK uyumlulugu)",
    "Hassas veri maskeleme: Sifre, TCKN loglanmaz", "Soft delete: is_active=False"]:
    bul(doc, f"- {item}")

# 3.5 Isimlendirme Standartlari (Konu #15)
h2(doc, "3.5 Isimlendirme Standartlari - PEP 257")
p(doc, "Projede Python tarafinda PEP 8 (stil) ve PEP 257 (docstring) standartlari takip edilmektedir.", sz=12, ind=1.25)
p(doc, "Cizelge 3.2: Isimlendirme Standartlari", sz=10, b=True, al=C, sb=12)
tbl(doc, ["Oge", "Kural", "Ornek"], [
    ["Python Sinif","PascalCase","PatientViewSet, CustomUser"], ["Python Fonksiyon","snake_case","get_queryset, perform_create"],
    ["React Bilesen","PascalCase","AppointmentDialog"], ["React Hook","camelCase (use prefix)","useAuth, useFetch"],
    ["API Endpoint","kebab-case","/api/treatment-types/"], ["Commit Mesaji","Conventional Commits","feat(auth): JWT refresh eklendi"],
], widths=[3.5, 4, 7])

# 3.5.1 API Isimlendirme (Konu #16)
h3(doc, "3.5.1 API Isimlendirme Standartlari")
p(doc, "Cizelge 3.3: API Endpoint Listesi", sz=10, b=True, al=C, sb=12)
tbl(doc, ["HTTP", "Endpoint", "Aciklama", "Yetki"], [
    ["POST","/api/auth/token/","JWT token alma (login)","Acik"],
    ["POST","/api/auth/token/refresh/","Token yenileme","Acik"],
    ["GET","/api/auth/me/","Mevcut kullanici bilgisi","Auth"],
    ["GET/POST","/api/patients/","Hasta listele / olustur","Auth"],
    ["GET/POST","/api/appointments/","Randevu listele / olustur","Auth"],
    ["GET/POST","/api/treatments/","Tedavi listele / olustur","Auth"],
    ["GET/POST","/api/payments/","Odeme listele / olustur","Auth"],
    ["GET","/api/dashboard/today/","Gunluk ozet","Auth"],
    ["GET/PUT","/api/settings/clinic/","Klinik ayarlari","Admin"],
    ["GET","/api/audit-logs/","Islem gecmisi","Admin"],
    ["POST","/api/public/register/","Yeni klinik kaydi","Acik"],
], widths=[2, 4.5, 4, 2.5])

# 3.6 Test Datasi Uretimi (Konu #17)
h2(doc, "3.6 Test Datasi Uretim Sureci")
p(doc, "Projede test verisi uretimi uc katmanda yonetilmektedir:", sz=12, ind=1.25)
p(doc, "Backend (factory-boy + Faker): Her model icin fabrika siniflari. Faker tr_TR locale ile gercekci Turkce isimler uretir.", sz=12, ind=1.25)
p(doc, "Frontend (Factory fonksiyonlari): makePatient(), makeAppointment() gibi fabrika fonksiyonlari.", sz=12, ind=1.25)
p(doc, "Demo Veri Seeding: run-demo.ps1 scripti ile otomatik demo ortami olusturma.", sz=12, ind=1.25)

# 3.7 Versiyon Yonetimi (Konu #27)
h2(doc, "3.7 Versiyon Yonetimi")
p(doc, "Proje Git ile yonetilmekte olup GitHub uzerinde barindirilmaktadir. Branch stratejisi:", sz=12, ind=1.25)
for item in ["main: Production branch", "develop: Gelistirme branch'i", "feat/*: Yeni ozellik branch'leri",
    "fix/*: Hata duzeltme", "test/*: Test eklemeleri"]:
    bul(doc, f"- {item}")

# 3.8 Development Ortami ve IDE (Konu #28)
h2(doc, "3.8 Development Ortami ve IDE")
p(doc, "Cizelge 3.4: Gelistirme Araclari", sz=10, b=True, al=C, sb=12)
tbl(doc, ["Arac", "Kullanim Amaci"], [
    ["VS Code / Cursor AI","Ana IDE - AI destekli kod tamamlama"],
    ["Docker Desktop","PostgreSQL konteynerizasyonu"],
    ["Node.js v18+ / Python 3.12","Calistirma ortamlari"],
    ["Git + GitHub","Versiyon kontrol"],
    ["Postman","API test ve debug"],
    ["Figma","UI/UX tasarim"],
    ["GitHub Copilot","AI destekli kod tamamlama"],
], widths=[4, 10])
doc.add_page_break()

# ====================================================================
# BOLUM 4: BULGULAR  (Sablon: Bolum 4)
# Alt basliklar: Konu #18-26, 29, 30
# ====================================================================
h1(doc, "4. BULGULAR")

# 4.1 Olcumleme (Konu #18)
h2(doc, "4.1 Olcumleme ve Surec Metrikleri")
p(doc, "Cizelge 4.1: Test Istatistikleri", sz=10, b=True, al=C, sb=12)
tbl(doc, ["Metrik", "Baslangic", "Son Durum", "Degisim"], [
    ["Backend Test Sayisi","0 calisan (77 kirik)","169","+169"],
    ["Frontend Unit Test","23","151","+128"],
    ["Frontend E2E Test","67","83","+16"],
    ["A11y Unit + E2E","0","14","+14"],
    ["Tenant Izolasyon Testi","0","21","+21"],
    ["TOPLAM TEST","~90","403","+313"],
    ["Backend Coverage","%55","%78","+23 puan"],
    ["Frontend Lines Coverage","%5","%57","+52 puan"],
    ["WCAG 2.1 AA Ihlali","Bilinmiyor","0","Sifir tolerans"],
    ["CI Required Checks","0","7","+7"],
], widths=[4.5, 3, 3, 3])

# 4.2 Isbirligi (Konu #19)
h2(doc, "4.2 Is Birligine Dayali Gelistirme")
for item in ["GitHub Issues ve Projects: Gorev takibi ve Kanban board",
    "Pull Request Review: Her PR en az 1 onay gerektirir",
    "CODEOWNERS: Kritik dosyalar icin otomatik reviewer",
    "Branch stratejisi: main, develop, feat/*, fix/*, test/*",
    "Pair Programming: Karmasik ozellikler icin esli programlama",
    "AI-Assisted Development: AI araclari ile pair programming"]:
    bul(doc, f"- {item}")

# 4.3 Mimari Cizim (Konu #20)
h2(doc, "4.3 Mimari Cizim ve Katmanli Yapi")
code(doc, """
+--------------------------------------------------------------+
|                    KULLANICI (Browser)                         |
|                   React.js + TypeScript                       |
+--------------------------------------------------------------+
|                    API KATMANI (REST)                          |
|              Django REST Framework + JWT                       |
+--------------------------------------------------------------+
|               TENANT MIDDLEWARE KATMANI                        |
|          HeaderTenantMiddleware (X-Tenant / Host)              |
+--------------------------------------------------------------+
|              VERITABANI KATMANI (PostgreSQL)                   |
|     +----------+  +----------+  +----------+                  |
|     | public   |  | klinik1  |  | klinik2  | <- Schema/tenant |
|     +----------+  +----------+  +----------+                  |
+--------------------------------------------------------------+""")

# UML Diyagramlari
h3(doc, "4.3.1 UML Diyagramlari")
diagram_files = sorted([f for f in os.listdir(DIAG) if f.endswith('.jpeg')])
captions = ["Sekil 4.1: Tedavi ve Odeme Girisi Sekans Diyagrami", "Sekil 4.2: Kullanici Girisi Sekans Diyagrami",
    "Sekil 4.3: Kullanici Girisi (Alternatif)", "Sekil 4.4: Randevu Olusturma ve Cakisma Kontrolu",
    "Sekil 4.5: Dashboard Gunluk Ozet", "Sekil 4.6: Hasta Kayit ve Guncelleme",
    "Sekil 4.7: Klinik Ayarlari Guncelleme", "Sekil 4.8: Genel Mimari Diyagram", "Sekil 4.9: Sistem Bilesen Diyagrami"]
for i, f in enumerate(diagram_files):
    cap = captions[i] if i < len(captions) else f"Sekil 4.{i+1}: {f}"
    img(doc, os.path.join(DIAG, f), w=6.0, cap=cap)

# 4.4 Veri Modeli (Konu #21)
h2(doc, "4.4 Veri Modeli")
p(doc, "Cizelge 4.2: Veritabani Varliklari", sz=10, b=True, al=C, sb=12)
tbl(doc, ["Model", "Tip", "Temel Alanlar", "Iliskiler"], [
    ["Client","Shared","schema_name, name, is_active","1:N Domain"],
    ["Domain","Shared","domain, is_primary","N:1 Client"],
    ["CustomUser","Tenant","username, email, role","1:N Appointment"],
    ["Patient","Tenant","first_name, last_name, phone, tckn","1:1 Anamnesis, 1:N Treatment"],
    ["Anamnesis","Tenant","medical_history, allergies, medications","1:1 Patient"],
    ["Appointment","Tenant","patient, doctor, date, time, status","N:1 Patient, N:1 Doctor"],
    ["Treatment","Tenant","patient, doctor, treatment_type, price","N:1 Patient, N:1 TreatmentType"],
    ["TreatmentType","Tenant","name, category, default_price","1:N Treatment"],
    ["Payment","Tenant","patient, treatment, amount","N:1 Patient, N:1 Treatment"],
    ["Document","Tenant","patient, name, file, uploaded_by","N:1 Patient"],
    ["AuditLog","Tenant","user, action, changes, ip","N:1 User"],
    ["ClinicSettings","Tenant","work_start_time, work_end_time","Singleton"],
], widths=[2.5, 1.5, 5.5, 4.5])

# 4.5 Test ve Deployment (Konu #22)
h2(doc, "4.5 Test ve Deployment Surecleri")

# 4.5.1 Test Sureci (Konu #23)
h3(doc, "4.5.1 Test Sureci ve Uygulamalari")
code(doc, """Test Piramidi:
            /\\
           /E2E\\         %10  - Playwright
          /------\\
         /Integ.  \\      %20  - APIClient/vitest+MSW
        /----------\\
       /   Unit     \\    %70  - pytest / vitest (izole)
      /______________\\""")
p(doc, "Cizelge 4.3: Test Katmanlari", sz=10, b=True, al=C, sb=12)
tbl(doc, ["Katman", "Arac", "Sayi", "Amac"], [
    ["Backend Unit","pytest + factory-boy","148","Model, serializer, view testleri"],
    ["Backend PG-only","FastTenantTestCase","21","Schema izolasyon testleri"],
    ["Frontend Unit","Vitest + Testing Library + MSW","151","Bilesen, hook, servis testleri"],
    ["Frontend A11y","vitest-axe","10","WCAG 2.1 AA taramasi"],
    ["E2E","Playwright","83","Uctan uca akislar"],
    ["E2E A11y","@axe-core/playwright","4","Tarayici erisilebilirlik"],
    ["Mutation","mutmut / Stryker","-","Test kalitesi dogrulama"],
], widths=[3, 4, 1.5, 6])

# 4.5.2 Deployment (Konu #24)
h3(doc, "4.5.2 Deployment Sureci")
p(doc, "Frontend Vercel'e, backend Docker ile Render'a deploy edilmektedir. CI/CD GitHub Actions ile otomatize edilmistir.", sz=12, ind=1.25)
code(doc, """CI/CD Pipeline:
+-----------------------------------------------+
| Backend Matrix (SQLite + PostgreSQL)           |
+-----------------------------------------------+
| Frontend Tests (vitest + coverage esigi)       |
+-----------------------------------------------+
| A11y Tests (WCAG 2.1 AA - sifir tolerans)      |
+-----------------------------------------------+
| E2E Tests (Playwright)                         |
+-----------------------------------------------+
HEPSI YESIL OLMADAN MERGE YOK""")

# 4.6 Uygulama Plani (Konu #25)
h2(doc, "4.6 Uygulama Plani")
p(doc, "Cizelge 4.4: Uygulama Plani", sz=10, b=True, al=C, sb=12)
tbl(doc, ["Faz", "Icerik", "Durum"], [
    ["Faz 1","Proje altyapisi: Django + React + PostgreSQL + Docker","Tamamlandi"],
    ["Faz 2","Multi-tenant mimari: django-tenants, schema izolasyonu","Tamamlandi"],
    ["Faz 3","CRUD API'lari: Hasta, Randevu, Tedavi, Odeme","Tamamlandi"],
    ["Faz 4","Frontend: Dashboard, PatientProfile, AppointmentCalendar","Tamamlandi"],
    ["Faz 5","Auth + RBAC: JWT, rol bazli erisim, login/logout","Tamamlandi"],
    ["Faz 6","Backend test suite: 148 pytest + 21 PG izolasyon","Tamamlandi"],
    ["Faz 7","Frontend test suite: 151 vitest + 83 Playwright E2E","Tamamlandi"],
    ["Faz 8","A11y: WCAG 2.1 AA zero-tolerance","Tamamlandi"],
    ["Faz 9","Deployment: Vercel + Render + CI/CD pipeline","Tamamlandi"],
    ["Faz 10","Dokumantasyon ve rapor","Tamamlandi"],
], widths=[1.5, 9, 3])

# 4.7 Kalite Pratikleri (Konu #26)
h2(doc, "4.7 Kalite Pratikleri")
for item in ["Coverage Threshold: Backend %78, Frontend lines %54 - asla dusurulmez",
    "A11y Zero-Tolerance: Tek WCAG 2.1 AA ihlali CI'i kirar",
    "Pre-commit Hooks (Husky): ESLint, Prettier, TypeScript type-check",
    "Conventional Commits: Standart commit mesaj formati zorunlu",
    "Code Review: Her PR en az 1 onay gerektirir",
    "Flaky Test Politikasi: 1 haftada duzeltilmezse karantina, 2 haftada silinir",
    "Slow Test Budget: Unit <= 200ms, Integration <= 2s, E2E <= 30s",
    "Mutation Testing: mutmut (backend) + Stryker (frontend)"]:
    bul(doc, f"- {item}")

# 4.8 Uygulama Onyzuleri (Konu #29)
h2(doc, "4.8 Uygulama Onyzuleri ve Kullanim Kilavuzu")
p(doc, "Canli uygulama: https://yasca-dental-clinic.vercel.app/", sz=12, ind=1.25)
screenshots = [
    ("saas_landing_page_1780686868207.png", "Sekil 4.10: SaaS Tanitim Sayfasi"),
    ("login_page_1780686889822.png", "Sekil 4.11: Klinik Giris Sayfasi"),
    ("dashboard_page_1780686924336.png", "Sekil 4.12: Dashboard"),
    ("patients_page_1780686936903.png", "Sekil 4.13: Hasta Listesi"),
    ("appointments_page_1780686949382.png", "Sekil 4.14: Randevu Takvimi"),
    ("settings_page_1780686962621.png", "Sekil 4.15: Klinik Ayarlari"),
]
for fname, cap in screenshots:
    img(doc, os.path.join(SS, fname), w=5.8, cap=cap)

# 4.9 Use Cases (Konu #30)
h2(doc, "4.9 Use Case'ler - UML Dokumantasyonu")
p(doc, "Cizelge 4.5: Temel Use Case Listesi", sz=10, b=True, al=C, sb=12)
tbl(doc, ["UC No", "Use Case", "Aktor", "Aciklama"], [
    ["UC-01","Sisteme Giris","Tum Roller","JWT ile kimlik dogrulama, rol bazli yonlendirme"],
    ["UC-02","Hasta Kaydi","Doktor, Asistan","Ad, soyad, telefon (zorunlu) ile kayit"],
    ["UC-03","Randevu Olusturma","Asistan, Admin","Cakisma kontrolu ile randevu"],
    ["UC-04","Tedavi Kaydi","Doktor","Tedavi turu, dis no (FDI), fiyat"],
    ["UC-05","Odeme Kaydi","Asistan, Admin","Tutar, tedaviye baglama"],
    ["UC-06","Dashboard","Tum Roller","Gunun randevulari, istatistikler"],
    ["UC-07","Klinik Ayarlari","Admin","Calisma saatleri, gunler"],
    ["UC-08","Kullanici Yonetimi","Admin","Personel ekleme, rol atama"],
    ["UC-09","Islem Gecmisi","Admin","Audit log kayitlari"],
    ["UC-10","Yeni Klinik Kaydi","Misafir","SaaS kayit formu ile tenant olusturma"],
], widths=[1.5, 3.5, 3, 6])
doc.add_page_break()

# ====================================================================
# BOLUM 5: SONUC VE ONERILER  (Sablon: Bolum 5)
# Alt basliklar: Konu #31-34
# ====================================================================
h1(doc, "5. SONUC VE ONERILER")

# 5.1 Proje Yonetim Pratikleri (Konu #31)
h2(doc, "5.1 Proje Yonetim Pratikleri")
for item in ["Agile/Scrum: Haftalik sprint'ler ile iteratif gelistirme",
    "GitHub Issues: Her is birimi issue olarak takip edilir",
    "Kanban Board: To Do > In Progress > Review > Done",
    "Definition of Done: Kod + Test + PR onayi + CI yesil + Merge",
    "CODEOWNERS: Kritik dosyalar icin otomatik reviewer",
    "Branch Protection: main branch'e dogrudan push YASAK"]:
    bul(doc, f"- {item}")

# 5.2 Neler Ogrendik (Konu #32)
h2(doc, "5.2 Neler Ogrendik - Basarili ve Basarisiz Oneriler")
p(doc, "Basarili Uygulamalar:", sz=12, b=True, sb=6)
for s in ["Dual-mode test stratejisi hem hiz hem dogruluk sagladi",
    "Factory pattern test kodunu ~%70 kisaltti", "A11y'yi bastan entegre etmek cok daha etkili oldu",
    "Pre-commit hook'lari bozuk kodun CI'a ulasmadan yakalanmasini sagladi",
    "AI araclari boilerplate uretiminde buyuk zaman kazandirdi",
    "Schema-based multi-tenancy veri izolasyonunu garanti etti"]:
    bul(doc, f"[+] {s}")
p(doc, "Zorluklar ve Dersler:", sz=12, b=True, sb=12)
for c in ["Coverage rakami yalan soyleyebilir - mutation score onemli",
    "SQLite ile PostgreSQL davranis farklari sorun cikartti",
    "Devraldigimiz 77 kirik test once ayaga kaldirildi",
    "E2E testler flaky olabilir - retry politikasi gerekti",
    "AI urettigi kodda guvenlik aciklari olabilir - insan review sart"]:
    bul(doc, f"[!] {c}")

# 5.3 Takim Degerlendirme (Konu #33)
h2(doc, "5.3 Takim Uyelerinin Proje Degerlendirme Raporu")
for nm, role, ev in [
    ("Yaman Halloum","Backend Gelistirici","Django REST Framework ve django-tenants ile multi-tenant altyapiyi kurmak projedeki en ogretici sureclerden biriydi. Veritabani modellerini ve RESTful API'leri gelistirirken yapay zeka araclarindan hizli prototipleme konusunda buyuk fayda sagladim."),
    ("Ali Ure","Full-Stack Gelistirici & Mimari Kurulum","Projenin hem backend hem de frontend tarafindaki genel mimarisini kurup, Vercel ve Render uzerindeki canli dagitim (deployment) sureclerini yonetmek bana tam yigin (full-stack) gelistirme konusunda ciddi bir deneyim kazandirdi."),
    ("Cihan Kurtbey","Frontend Gelistirici","React ile ozelikle kimlik dogrulama akislari, sifre sifirlama arayuzleri ve e-posta entegrasyonu tarafinda calismak, guvenlik ve kullanici deneyimini bir arada dusunmeyi ogretti. AI araclari bilesen uretimimi cok hizlandirdi."),
    ("Sukru Yesilmen","Test & Kalite Guvence","Projenin test stratejilerini olusturmak ve dokumantasyon sureclerini yurutmek, bir yazilimin sadece calismasinin degil, ayni zamanda surdurulebilir olmasinin ne kadar onemli oldugunu gosterdi. Kalite pratikleri konusunda derinlemesine bilgi edindim.")]:
    p(doc, f"{nm} - {role}", sz=12, b=True, sb=12)
    p(doc, ev, sz=12, ind=1.25)

# 5.4 Yapay Zeka Kullanim Pratikleri (Konu #34)
h2(doc, "5.4 Yapay Zeka Kullanim Pratikleri")
p(doc, "Proje suresince cesitli yapay zeka araclari aktif olarak kullanilmistir:", sz=12, ind=1.25)
p(doc, "Cizelge 5.1: Yapay Zeka Araclari ve Kullanim Alanlari", sz=10, b=True, al=C, sb=12)
tbl(doc, ["AI Araci", "Kullanim Alani", "Somut Katki"], [
    ["GitHub Copilot","Kod tamamlama, boilerplate","Serializer, factory siniflari otomatik uretimi"],
    ["ChatGPT (GPT-4)","Mimari tasarim, problem cozme","Multi-tenant strateji, hata analizi"],
    ["Cursor AI","IDE entegreli AI kodlama","Komponent gelistirme, refactoring"],
    ["Claude Code","Kapsamli kod uretimi","Test suite, dokumantasyon, rapor otomasyonu"],
    ["Gemini","Kod analizi, rapor uretimi","Proje analizi, ekran goruntusu toplama"],
], widths=[3, 3.5, 7.5])
p(doc, "Yapay Zeka Kullaniminda Dikkat Edilen Noktalar:", sz=12, b=True, sb=12)
for note in ["AI urettigi kod her zaman insan review'dan gecirilmistir",
    "Guvenlik kritik bolumlerde AI onerileri ekstra dikkatle incelenmistir",
    "AI araclari test yazimini hizlandirmis, ancak senaryo tasarimi ekip tarafindan yapilmistir",
    "AI ile pair programming yaklasimi benimsenmistir",
    "Prompt muhendisligi: Spesifik ve baglamli prompt'lar daha iyi sonuc vermistir"]:
    bul(doc, f"- {note}")
doc.add_page_break()

# ==================== KAYNAKLAR ====================
h1(doc, "KAYNAKLAR")
for ref in [
    "[1] Django Software Foundation, Django Documentation, https://docs.djangoproject.com/en/5.2/",
    "[2] Django REST Framework, API Guide, https://www.django-rest-framework.org/",
    "[3] django-tenants Documentation, https://django-tenants.readthedocs.io/",
    "[4] React Documentation, https://react.dev/",
    "[5] TypeScript Handbook, https://www.typescriptlang.org/docs/",
    "[6] Vite Build Tool, https://vitejs.dev/",
    "[7] PostgreSQL Documentation, https://www.postgresql.org/docs/15/",
    "[8] pytest Documentation, https://docs.pytest.org/",
    "[9] Vitest Documentation, https://vitest.dev/",
    "[10] Playwright Documentation, https://playwright.dev/",
    "[11] WCAG 2.1 Guidelines, W3C, https://www.w3.org/TR/WCAG21/",
    "[12] OpenAI, ChatGPT, https://chat.openai.com/",
    "[13] GitHub Copilot, https://github.com/features/copilot",
    "[14] Anthropic, Claude, https://claude.ai/",
    "[15] Martin Fowler, Test Pyramid, https://martinfowler.com/articles/practical-test-pyramid.html",
    "[16] PEP 8, https://peps.python.org/pep-0008/",
    "[17] PEP 257, https://peps.python.org/pep-0257/",
]: p(doc, ref, sz=11, sa=6)
doc.add_page_break()

# ==================== EKLER ====================
h1(doc, "EKLER")
h2(doc, "EK-1: Proje Kaynak Kodu Yapisi")
code(doc, """yasca-dental-clinic/
+-- backend/          # Django Backend (api/, customers/, core/)
+-- frontend/         # React Frontend (components/, services/, hooks/)
+-- docs/             # Dokumantasyon (TESTING.md, adr/, diagrams/)
+-- .github/          # CI/CD ve PR template
+-- docker-compose.yml""")
sw = os.path.join(BASE, "swagger_screenshot.png")
if os.path.exists(sw):
    h2(doc, "EK-2: Swagger API Dokumantasyonu")
    img(doc, sw, w=5.5, cap="Sekil EK-2.1: Swagger UI")
doc.add_page_break()

# ==================== OZGECMIS ====================
h1(doc, "OZGECMIS")
for nm in ["Yaman Halloum", "Ali Ure", "Cihan Kurtbey", "Sukru Yesilmen"]:
    p(doc, nm, sz=12, b=True, sb=12)
    p(doc, "Istanbul Saglik ve Teknoloji Universitesi, Yazilim Muhendisligi Bolumu, Lisans Ogrencisi", sz=12)

# ==================== KAYDET ====================
doc.save(OUT)
print(f"[OK] Rapor olusturuldu: {OUT}")
print(f"   Paragraf: {len(doc.paragraphs)}, Tablo: {len(doc.tables)}")
