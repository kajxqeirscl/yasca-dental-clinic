# -*- coding: utf-8 -*-
"""
Grup 4 - Yasca Dis Klinigi Yonetim Sistemi
Final Proje Raporu - DOCX Olusturucu
Times New Roman, 12pt, 1.5 satir araligi
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

FONT = 'Times New Roman'
FONT_SIZE = Pt(12)
LINE_SPACING = 1.5

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Force Times New Roman on run XML (covers ALL character sets) ──
def set_run_font(run, font_name=FONT):
    """Set font name on ALL script types (Latin, EastAsia, Complex) via XML."""
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}" w:cs="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)

def set_style_font(style_obj, font_name=FONT, font_size=None):
    """Set font on a style's rPr element for all character sets."""
    style_obj.font.name = font_name
    if font_size:
        style_obj.font.size = font_size
    # Also set via XML for full coverage
    element = style_obj.element
    rPr = element.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}/>')
        element.append(rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}" w:cs="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)

# ── Set document-level default font ──
doc_defaults = doc.styles.element.find(qn('w:docDefaults'))
if doc_defaults is None:
    doc_defaults = parse_xml(f'<w:docDefaults {nsdecls("w")}/>')
    doc.styles.element.insert(0, doc_defaults)
rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
if rPrDefault is None:
    rPrDefault = parse_xml(f'<w:rPrDefault {nsdecls("w")}/>')
    doc_defaults.append(rPrDefault)
rPr_elem = rPrDefault.find(qn('w:rPr'))
if rPr_elem is None:
    rPr_elem = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:ascii="{FONT}" w:hAnsi="{FONT}" w:eastAsia="{FONT}" w:cs="{FONT}"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr>')
    rPrDefault.append(rPr_elem)
else:
    rFonts = rPr_elem.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT}" w:hAnsi="{FONT}" w:eastAsia="{FONT}" w:cs="{FONT}"/>')
        rPr_elem.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), FONT)
        rFonts.set(qn('w:hAnsi'), FONT)
        rFonts.set(qn('w:eastAsia'), FONT)
        rFonts.set(qn('w:cs'), FONT)
    # Set size (24 half-points = 12pt)
    sz = rPr_elem.find(qn('w:sz'))
    if sz is None:
        sz = parse_xml(f'<w:sz {nsdecls("w")} w:val="24"/>')
        rPr_elem.append(sz)
    else:
        sz.set(qn('w:val'), '24')
    szCs = rPr_elem.find(qn('w:szCs'))
    if szCs is None:
        szCs = parse_xml(f'<w:szCs {nsdecls("w")} w:val="24"/>')
        rPr_elem.append(szCs)
    else:
        szCs.set(qn('w:val'), '24')

# ── Style Setup: Normal ──
style = doc.styles['Normal']
set_style_font(style, FONT, FONT_SIZE)
style.paragraph_format.line_spacing = LINE_SPACING
style.paragraph_format.space_after = Pt(6)

# ── Style Setup: Headings ──
for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0, 0, 0)  # Black headings for academic format
    hs.font.bold = True
    hs.paragraph_format.line_spacing = LINE_SPACING
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    if level == 1:
        set_style_font(hs, FONT, Pt(16))
    elif level == 2:
        set_style_font(hs, FONT, Pt(14))
    else:
        set_style_font(hs, FONT, Pt(13))

# ── Style Setup: List Bullet ──
try:
    lb = doc.styles['List Bullet']
    set_style_font(lb, FONT, FONT_SIZE)
    lb.paragraph_format.line_spacing = LINE_SPACING
except:
    pass

DIAGRAMS = r"c:\Users\Ali\yasca-dental-clinic\docs\diagrams"
SWAGGER = r"c:\Users\Ali\yasca-dental-clinic\swagger_screenshot.png"

# ── Helpers ──
def p(text, bold=False, italic=False, size=None, align=None):
    """Add paragraph with guaranteed Times New Roman."""
    para = doc.add_paragraph()
    if align: para.alignment = align
    para.paragraph_format.line_spacing = LINE_SPACING
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    set_run_font(run, FONT)
    run.font.size = size or FONT_SIZE
    run.bold = bold
    run.italic = italic
    return para

def bullet(text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.line_spacing = LINE_SPACING
    para.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    para.clear()
    run = para.add_run(text)
    set_run_font(run, FONT)
    run.font.size = FONT_SIZE
    return para

def code_block(code):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
    pPr.append(shading)
    run = para.add_run(code)
    set_run_font(run, 'Consolas')
    run.font.size = Pt(9)
    return para

def styled_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        pp = cell.paragraphs[0]
        pp.paragraph_format.line_spacing = LINE_SPACING
        run = pp.add_run(h)
        run.bold = True
        set_run_font(run, FONT)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F497D"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            pp = cell.paragraphs[0]
            pp.paragraph_format.line_spacing = LINE_SPACING
            run = pp.add_run(str(val))
            set_run_font(run, FONT)
            run.font.size = Pt(10)
            if ri % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8EDF3"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def add_image_safe(path, width=None):
    """Safely add an image if it exists."""
    if os.path.exists(path):
        doc.add_picture(path, width=width or Inches(5.5))
        return True
    return False

# diagram files sorted by timestamp name
diagram_files = sorted(os.listdir(DIAGRAMS))

# ═══════════════════════════════════════════════════════════
# KAPAK SAYFASI
# ═══════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()

p('İSTANBUL SAĞLIK VE TEKNOLOJİ ÜNİVERSİTESİ', bold=True, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
p('Mühendislik ve Doğa Bilimleri Fakültesi', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
p('Yazılım Mühendisliği Bölümü', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()

p('YAZ402 – Yazılım Gerçekleme, Test ve Bakım', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

p('FINAL PROJE RAPORU', bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

p('Yaşca Diş Kliniği Yönetim Sistemi', italic=True, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
p('Çok-Kiracılı (Multi-Tenant) SaaS Mimarisi', italic=True, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3): doc.add_paragraph()

p('Grup 4', bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER)
for m in ['Yaman Halloum', 'Ali Üre', 'Cihan Kurtbey', 'Berkay Aydın']:
    p(m, size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
p('2024 – 2025 Bahar Dönemi', size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# İÇİNDEKİLER
# ═══════════════════════════════════════════════════════════
doc.add_heading('İÇİNDEKİLER', level=1)

toc = [
    'BÖLÜM 1: Giriş ve Proje Kapsamı',
    '   1.1. Projenin Amacı',
    '   1.2. Ana Hedefler, Hedef Kullanıcılar ve Kritik Başarı Kriterleri',
    'BÖLÜM 2: Yazılım Gereksinimleri',
    '   2.1. Fonksiyonel İsterler',
    '   2.2. Fonksiyonel Olmayan İsterler',
    'BÖLÜM 3: Detaylı UML Tasarımı',
    '   3.1. Sınıf ve Sekans Analizi',
    '   3.2. Tasarım Detayları',
    'BÖLÜM 4: Teknoloji Karar Analizi (DAR)',
    '   4.1. Teknoloji Seçimlerinin Özeti',
    '   4.2. Karar Matrisleri',
    'BÖLÜM 5: Gerçekleme ve Kalite Analizi (Opsiyon A)',
    '   5.1. İç Arayüzler ve API Kontratları',
    '   5.2. Dış Servis Entegrasyonları',
    '   5.3. Multi-Tenant (SaaS) Mimarisi',
    '   5.4. Gerçek Kod Örnekleri',
    'BÖLÜM 6: Birim Test Sonuçları',
    '   6.1. Birim Test Durum Tablosu',
    '   6.2. Test İyileştirmeleri',
    'BÖLÜM 7: Bakım ve Gelecek Önerileri',
    '   7.1. Sürdürülebilirlik ve İyileştirme',
    'BÖLÜM 8: Sonuç',
    '   8.1. Süreç Değerlendirmesi',
    '   8.2. Temel Çıktıların Kontrolü',
]
for item in toc:
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = LINE_SPACING
    run = para.add_run(item)
    run.font.name = FONT
    run.font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# BÖLÜM 1: GİRİŞ VE PROJE KAPSAMI
# ═══════════════════════════════════════════════════════════
doc.add_heading('BÖLÜM 1: Giriş ve Proje Kapsamı', level=1)

doc.add_heading('1.1. Projenin Amacı', level=2)
p('Bu proje, diş hekimliği kliniklerinde manuel olarak yürütülen hasta kaydı, randevu planlama ve tedavi takibi süreçlerini dijitalleştirmek amacıyla geliştirilmiş çok-kiracılı (multi-tenant) bir SaaS (Software as a Service) web uygulamasıdır.')
p('"Yaşca Diş Kliniği Yönetim Sistemi" adı verilen proje; hasta kaydı ve anamnez (tıbbi özgeçmiş) yönetimi, çakışma kontrollü randevu sistemi, tedavi ve finansal takip, dijital odontogram (interaktif diş şeması) ve klinik yönetim modüllerini tek bir çatı altında sunan kapsamlı bir platformdur.')
p('Sağlık sektöründe hâlâ kağıt tabanlı veya birbirinden kopuk yazılımlarla yürütülen klinik süreçleri, veri kaybı, randevu çakışması ve güvenlik açıkları gibi ciddi sorunlara yol açmaktadır. Bu proje, söz konusu problemleri çözmek üzere tasarlanmış olup, her bir kliniğin verilerini fiziksel olarak izole eden PostgreSQL şema tabanlı multi-tenant mimarisi ile güvenliği en üst düzeyde tutmaktadır.')
p('Sistemin sınırları şu şekilde çizilmiştir: Uygulama yalnızca klinik personeli (Yönetici, Hekim ve Asistan) tarafından kullanılmak üzere tasarlanmıştır. Hastalar sistemin kullanıcısı değildir; sistemde verisi tutulan varlıklardır.')

doc.add_heading('1.2. Ana Hedefler, Hedef Kullanıcılar ve Kritik Başarı Kriterleri', level=2)

p('Ana Hedefler', bold=True)
hedefler = [
    'Hasta kayıt ve anamnez (tıbbi özgeçmiş) yönetimini dijitalleştirmek',
    'Çakışma kontrollü randevu sistemi ile operasyonel verimliliği artırmak',
    'Tedavi ve ödeme takibini entegre bir yapıda sunmak',
    'Dijital diş şeması (Odontogram) ile klinik görselleştirme sağlamak',
    'Güvenli JWT tabanlı kimlik doğrulama ve rol bazlı yetkilendirme (RBAC) uygulamak',
    'Multi-tenant (çok-kiracılı) SaaS mimarisi ile birden fazla kliniğe hizmet vermek',
]
for h in hedefler: bullet(h)

doc.add_paragraph()
p('Hedef Kullanıcılar', bold=True)
p('Sistemde tanımlı üç kullanıcı rolü ve yetkileri aşağıdaki tabloda özetlenmiştir:')
styled_table(
    ['Rol', 'Açıklama', 'Erişim ve Yetkiler'],
    [
        ('Yönetici (Admin)', 'Genellikle Başhekim', 'Klinik ayarları, personel yönetimi, tüm verilere tam erişim, kullanıcı ekleme/şifre sıfırlama'),
        ('Hekim (Doktor)', 'Sistemin ana kullanıcısı', 'Kendi hastalarının randevuları, tedavi girişi, odontogram, dashboard görüntüleme'),
        ('Asistan / Sekreter', 'İdari süreçleri yönetir', 'Hasta kaydı, randevu oluşturma, ödeme girişi, telefon/WhatsApp talep yönetimi'),
    ],
    [3, 4, 10]
)
p('Not: Hasta (Patient) sistemin bir kullanıcısı değildir; sistemde verisi tutulan bir varlıktır.', italic=True)

doc.add_paragraph()
p('Kritik Başarı Kriterleri', bold=True)
kriterler = [
    'Randevu çakışması %0 oranında gerçekleşmelidir (F-008)',
    'Kiracılar arası veri sızıntısı kesinlikle önlenmelidir (PostgreSQL şema izolasyonu)',
    'API yanıt süreleri 2 saniyenin altında olmalıdır (NF-001)',
    'Sistem, en az 54 birim testin tamamını başarıyla geçmelidir',
    'Tüm API endpoint\'leri Swagger (OpenAPI) ile dokümante edilmelidir',
    'Sistem Docker ile tek komutla kurulabilir olmalıdır (NF-010)',
]
for k in kriterler: bullet(k)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# BÖLÜM 2: YAZILIM GEREKSİNİMLERİ
# ═══════════════════════════════════════════════════════════
doc.add_heading('BÖLÜM 2: Yazılım Gereksinimleri', level=1)

doc.add_heading('2.1. Fonksiyonel İsterler', level=2)
p('İster Analizi dokümanında (v01) tanımlanan ve Proje Adımı 1\'de baz alınan fonksiyonel gereksinim listesi aşağıdaki tabloda eksiksiz olarak sunulmaktadır:')

fr_rows = [
    ('F-001', 'Kullanıcı Yönetimi', 'Sistem, kayıtlı kullanıcıların e-posta adresi ve şifreleriyle sisteme güvenli bir şekilde giriş yapmasını sağlamalıdır.', 'Yüksek'),
    ('F-002', 'Kullanıcı Yönetimi', 'Kullanıcı, sistemden güvenli çıkış (logout) yapabilmelidir.', 'Orta'),
    ('F-003', 'Hasta Yönetimi', 'Kullanıcı, sadece Ad, Soyad ve Telefon bilgilerini girerek hızlıca yeni hasta kaydı oluşturabilmelidir. (TC No ve Doğum Tarihi opsiyonel olmalıdır.)', 'Yüksek'),
    ('F-004', 'Hasta Yönetimi', 'Kullanıcı, mevcut hastaları İsim, Soyad veya Telefon numarası ile arama çubuğundan bulabilmelidir.', 'Yüksek'),
    ('F-005', 'Hasta Yönetimi', 'Kullanıcı, kayıtlı bir hastanın profil bilgilerini sonradan güncelleyebilmelidir.', 'Orta'),
    ('F-006', 'Randevu Yönetimi', 'Kullanıcı, randevu takvimini Günlük, Haftalık ve Aylık görünümlerde görüntüleyebilmelidir.', 'Yüksek'),
    ('F-007', 'Randevu Yönetimi', 'Kullanıcı, takvimdeki boş bir saat dilimine tıklayıp hasta seçimi yaparak yeni bir randevu oluşturabilmelidir.', 'Yüksek'),
    ('F-008', 'Randevu Yönetimi', 'Sistem, aynı hekime aynı saatte ikinci bir randevu verilmesini (çakışma) engellemelidir.', 'Yüksek'),
    ('F-009', 'Randevu Yönetimi', 'Kullanıcı, mevcut bir randevunun durumunu "İptal", "Tamamlandı" veya "Gelmedi" olarak değiştirebilmelidir.', 'Orta'),
    ('F-010', 'Tedavi Yönetimi', 'Hekim, hastanın geçmişte yapılan tüm işlemlerini tarih sırasına göre listeleyebilmelidir.', 'Orta'),
    ('F-011', 'Tedavi Yönetimi', 'Hekim, interaktif diş şeması (odontogram) üzerinden diş seçerek yapılan işlemi kaydedebilmelidir.', 'Orta'),
    ('F-012', 'Bildirim Sistemi', 'Sistem, randevu oluşturulduğunda hastanın telefonuna bilgilendirme amaçlı WhatsApp mesajı gönderebilmelidir.', 'Düşük'),
    ('F-013', 'Raporlama', 'Sistem, ana sayfada (Dashboard) o günkü toplam randevu sayısını ve bugünün randevu listesini göstermelidir.', 'Düşük'),
    ('F-014', 'Finansal Yönetim', 'Kullanıcı, bir hastadan alınan ödemeyi (tutar ve açıklama ile) sisteme kaydedebilmelidir.', 'Orta'),
    ('F-015', 'Finansal Yönetim', 'Kullanıcı, bir hastanın profilinde toplam tedavi ücretini, ödenen tutarı ve kalan bakiyeyi görebilmelidir.', 'Orta'),
    ('F-016', 'Doküman Yönetimi', 'Kullanıcı, hastanın profiline röntgen görüntüsü veya fotoğraf (JPEG/PNG formatında) yükleyebilmelidir.', 'Düşük'),
    ('F-017', 'Sistem Yönetimi', 'Yönetici (Admin), sisteme yeni bir Hekim veya Asistan kullanıcısı ekleyebilmelidir.', 'Orta'),
    ('F-018', 'Tedavi Yönetimi', 'Diş şeması üzerinde; "Yapılacak" işlemler farklı renkte, "Tamamlanmış" işlemler farklı renkte gösterilmelidir.', 'Orta'),
    ('F-019', 'Randevu Yönetimi', 'Kullanıcı, randevu oluştururken randevuya özel kısa bir not (Örn: "Diş ağrısı var") ekleyebilmelidir.', 'Orta'),
    ('F-020', 'Tedavi Yönetimi', 'Yönetici, sistemde seçilebilir tedavi türlerini (Örn: Dolgu, Kanal, Çekim) ve bunların varsayılan fiyatlarını tanımlayabilmeli, düzenleyebilmeli ve silebilmelidir.', 'Orta'),
    ('F-021', 'Sistem Yönetimi', 'Yönetici (Admin), diğer kullanıcıların şifrelerini güvenlik nedeniyle talep üzerine sıfırlayabilmelidir.', 'Orta'),
    ('F-022', 'Randevu Yönetimi', 'Yönetici, randevu takviminin başlangıç ve bitiş saatlerini (Örn: 09:00 - 18:00) ve randevu aralıklarını (Örn: 15 dk) ayarlayabilmelidir.', 'Yüksek'),
]

styled_table(
    ['ID', 'Modül', 'Açıklama', 'Öncelik'],
    fr_rows,
    [1.5, 3, 10, 1.5]
)

doc.add_page_break()

doc.add_heading('2.2. Fonksiyonel Olmayan İsterler', level=2)
p('Performans, güvenlik, kullanılabilirlik ve uyumluluk gibi kalite nitelikleri aşağıda listelenmiştir:')

nfr_rows = [
    ('NF-001', 'Performans', 'Sistemdeki listeleme ve arama ekranları, standart bir internet bağlantısında 2 saniyeden kısa sürede yüklenmelidir.', 'Yüksek'),
    ('NF-002', 'Güvenlik', 'Tüm veri iletişimi SSL (HTTPS) sertifikası ile şifrelenmelidir.', 'Yüksek'),
    ('NF-003', 'Güvenlik', 'Hasta verileri (özellikle KVKK kapsamındaki veriler) veritabanında şifreli (AES-256) olarak saklanmalıdır.', 'Yüksek'),
    ('NF-004', 'Güvenlik', 'Kullanıcı oturumları, 30 dakika işlem yapılmadığında otomatik olarak sonlandırılmalıdır (Session Timeout).', 'Orta'),
    ('NF-005', 'Kullanılabilirlik', 'Arayüz, masaüstü/laptop ve mobil cihazlarda sorunsuz ve kullanıcı dostu bir deneyim sunmalıdır (Responsive Tasarım).', 'Yüksek'),
    ('NF-006', 'Kullanılabilirlik', 'Asistan, sisteme giriş yaptıktan sonra en fazla 3 tıklama ile yeni randevu oluşturma ekranına ulaşabilmelidir.', 'Yüksek'),
    ('NF-007', 'Güvenilirlik', 'Sistem, veritabanını günde bir kez otomatik olarak yedeklemelidir.', 'Yüksek'),
    ('NF-008', 'Uyumluluk', 'Uygulama; Chrome, Firefox ve Safari tarayıcılarının güncel sürümlerinde hatasız çalışmalıdır.', 'Orta'),
    ('NF-009', 'Dil Desteği', 'Uygulamanın kullanıcı arayüzü ve tüm sistem mesajları Türkçe olmalıdır.', 'Yüksek'),
    ('NF-010', 'Kurulum', 'Sistem, Docker konteyner yapısı sayesinde tek bir komutla (docker-compose up) kurulabilir olmalıdır.', 'Orta'),
]

styled_table(
    ['ID', 'Kategori', 'Açıklama', 'Öncelik'],
    nfr_rows,
    [1.5, 3, 10, 1.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# BÖLÜM 3: DETAYLI UML TASARIMI
# ═══════════════════════════════════════════════════════════
doc.add_heading('BÖLÜM 3: Detaylı UML Tasarımı', level=1)

doc.add_heading('3.1. Sınıf ve Sekans Analizi', level=2)
p('Proje Adımı 1 kapsamında, sistemin nesne tabanlı tasarımı standart UML notasyonlarına uygun olarak modellenmiştir. Tasar\u0131m modelleri oluşturulurken, İster Analizi dokümanındaki temel fonksiyonel gereksinimler (FR) baz alınmıştır. Sekans diyagramları sistemin dinamik davranışını, sınıf diyagramları ise statik yapısını göstermektedir.')
p('Tasarım süreci, "Sınıf ve Sekans diyagramları arası mutlak tutarlılık" prensibiyle iteratif olarak gerçekleştirilmiştir. Sekans diyagramlarında arayüzden gönderilen tüm mesajlar (fetchDashboardToday(), createAppointment(payload), login(username, password) vb.), Uygulama Katmanı Sınıf Diyagramındaki ApiService sınıfının metodları olarak birebir tanımlanmıştır.')

p('Sekans Diyagramları (Dinamik Model)', bold=True)
p('Projenin 6 temel kullanım durumunu temsil eden dinamik süreçler, Boundary (Arayüz), Control (Servis/API) ve Entity (Varlık) nesneleri kullanılarak modellenmiştir:')

sekans_list = [
    ('SD-01: Kullanıcı Girişi ve Oturum Doğrulaması (FR-01/F-001)', 'Klinik personelinin sisteme giriş yaptığı ve JWT jetonunun localStorage üzerinde saklanarak kullanıcı rolünün doğrulandığı süreçtir. LoginPage → ApiService → AuthController → CustomUser etkileşim akışı modellenmiştir.'),
    ('SD-02: Hasta Kaydı ve Anamnez Güncelleme (FR-02/F-003)', 'Yeni bir hasta profili oluşturulduğu ve tıbbi özgeçmiş (Anamnez) kayıtlarının veritabanına yazıldığı iş akışıdır.'),
    ('SD-03: Randevu Oluşturma ve Çakışma Kontrolü (FR-03/F-008)', 'Randevu talebi oluşturma sürecidir. Sistem, ilgili hekimin uygunluğunu denetler. Çakışma durumunda 400 hatası, uygunluk durumunda 201 Created döndürülür.'),
    ('SD-04: Günlük Dashboard Özet Akışı (FR-04/F-013)', 'Ana panelin veri çekme sürecidir. Bugüne ait randevular, tamamlanan işlem sayıları ve aktif hasta sayısı tek endpoint\'ten döner.'),
    ('SD-05: Tedavi ve Ödeme Girişi (FR-05/F-011)', 'Tedavi türlerinin çekildiği, işlemin kaydedildiği ve finansal ödemenin sisteme yansıtıldığı uçtan uca klinik operasyon sürecidir.'),
    ('SD-06: Klinik Ayarları Güncelleme (FR-06/F-022)', 'Yalnızca Admin yetkisine sahip yöneticinin erişebildiği, yetki kontrolü (403 Forbidden) aşamalarını içeren klinik çalışma saatleri düzenleme sürecidir.'),
]
for title, desc in sekans_list:
    p(title, bold=True, size=Pt(11))
    p(desc)

# Add sequence diagram images
seq_images = {
    'SD-01': 'WhatsApp Image 2026-04-25 at 17.03.46.jpeg',
    'SD-02': 'WhatsApp Image 2026-04-25 at 17.24.32.jpeg',
    'SD-03': 'WhatsApp Image 2026-04-25 at 17.06.03.jpeg',
    'SD-04': 'WhatsApp Image 2026-04-25 at 17.23.21.jpeg',
    'SD-05': 'WhatsApp Image 2026-04-25 at 17.00.59.jpeg',
    'SD-06': 'WhatsApp Image 2026-04-25 at 17.25.28.jpeg',
}

doc.add_page_break()
p('Sekans Diyagramları - Görseller', bold=True)

for sd_name, filename in seq_images.items():
    filepath = os.path.join(DIAGRAMS, filename)
    p(f'Şekil: {sd_name} Sekans Diyagramı', bold=True, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_image_safe(filepath, width=Inches(5.5))
    doc.add_paragraph()

doc.add_page_break()

doc.add_heading('3.2. Tasarım Detayları', level=2)

p('Sınıf Diyagramları (Statik Model)', bold=True)
p('Sekans diyagramlarında kullanılan tüm arayüz, servis ve veri varlıkları, modüler bir mimari yaklaşımıyla iki ana UML Sınıf Diyagramında birleştirilmiştir.')

p('3.2.1. Domain Katmanı (Veri Varlıkları) Sınıf Diyagramı', bold=True)
p('Veritabanı tablolarının nesne tabanlı karşılıklarını ve ilişkilerini gösteren modeldir:')

# Domain class diagram image
domain_img = os.path.join(DIAGRAMS, 'WhatsApp Image 2026-04-25 at 17.27.36.jpeg')
p('Şekil: Domain Katmanı Sınıf Diyagramı', bold=True, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
add_image_safe(domain_img, width=Inches(5))

doc.add_paragraph()
p('Sınıf İlişkileri (Aggregation/Composition):', bold=True)
iliskiler = [
    'Clinic → Patient: 1..* Association — Bir kliniğe birden fazla hasta kayıtlıdır',
    'Clinic → CustomUser: 1..* Association — Bir klinikte birden fazla personel bulunur',
    'Patient → Anamnesis: 1..1 Composition — Her hastanın en fazla bir anamnez kaydı vardır; hasta silinirse anamnez de silinir',
    'Patient → Appointment: 1..* Composition — Hasta silindiğinde randevuları da silinir (CASCADE)',
    'Appointment → Treatment: 1..* Aggregation — Bir randevuda birden fazla tedavi yapılabilir',
    'Treatment → Payment: 1..1 Association — Her tedaviye bir ödeme kaydı bağlanır',
    'TreatmentType → Treatment: 1..* — Bir tedavi türü birden fazla tedavide kullanılabilir',
]
for i in iliskiler: bullet(i)

doc.add_page_break()

p('3.2.2. Uygulama Katmanı (Boundary / Control Sınıfları) Sınıf Diyagramı', bold=True)
app_img = os.path.join(DIAGRAMS, 'WhatsApp Image 2026-04-25 at 17.28.41.jpeg')
p('Şekil: Uygulama Katmanı Sınıf Diyagramı', bold=True, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
add_image_safe(app_img, width=Inches(5.5))

doc.add_paragraph()
p('Boundary (Arayüz Bileşenleri): LoginPage, DashboardUI, PatientDialog, AppointmentDialog gibi React bileşenleri. Control (İş Mantığı): Frontend tarafında ApiService, Backend tarafında ViewSet ve Serializer sınıfları. Tasarım kuralı: Boundary asla doğrudan Entity ile konuşmaz; araya Control katmanı girerek MVC/MVT mimarisine uyum sağlanmıştır.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# BÖLÜM 4: TEKNOLOJİ KARAR ANALİZİ (DAR)
# ═══════════════════════════════════════════════════════════
doc.add_heading('BÖLÜM 4: Teknoloji Karar Analizi (DAR)', level=1)

doc.add_heading('4.1. Teknoloji Seçimlerinin Özeti', level=2)
p('Proje Adımı 2 kapsamında, Adım 1\'deki nesne tabanlı tasarımı en iyi gerçekleyecek teknoloji yığını; performans, ekip yetkinliği, güvenlik ve OOP uyumu kriterleriyle objektif olarak analiz edilmiştir.')

p('Backend Dil Seçimi: Python', bold=True)
p('Sınıf diyagramında tasarlanan nesneler (CustomUser, Patient, Appointment) ve aralarındaki kalıtım/ilişki yapıları nesne yönelimli bir dil gerektirmektedir. Python\'un sunduğu hızlı prototipleme (rapid prototyping) yeteneği ve sağlık verisi şifreleme/hashing kütüphanelerinin zenginliği sebebiyle Python tercih edilmiştir.')

p('Backend Framework: Django REST Framework (DRF)', bold=True)
p('Sekans diyagramlarındaki kontrolcü nesnelerin iş mantığı, Django\'nun MVC mimarisiyle birebir örtüşmektedir. DRF\'in yerleşik JWT kütüphanesi ile AuthController senaryosu doğrudan gerçeklenebilmektedir.')

p('Frontend Framework: React.js (TypeScript)', bold=True)
p('Sınıf diyagramındaki arayüz bileşenleri React\'in Component-Based mimarisiyle doğrudan eşleşir. Dijital Odontogram için Virtual DOM yapısı performans avantajı sağlar.')

p('Veritabanı: PostgreSQL', bold=True)
p('Sınıf diyagramındaki kesin yapısal ilişkiler ve sağlık verilerinin gerektirdiği ACID uyumluluğu nedeniyle PostgreSQL seçilmiştir. Multi-tenant mimaride django-tenants ile şema bazlı izolasyon sağlanmaktadır.')

p('ORM Aracı: Django ORM', bold=True)
p('Sınıf diyagramı nesnelerinin veritabanı tablolarına eşlenmesi Django ORM ile sağlanmakta; "Bire-Çok" ilişkiler ForeignKey, "Bire-Bir" ilişkiler OneToOneField ile koda dökülmektedir.')

doc.add_heading('4.2. Karar Matrisleri', level=2)
p('Alternatifler 1 (en düşük) – 5 (en yüksek) arasında puanlanmıştır.')

p('4.2.1. Backend Framework Karar Matrisi', bold=True)
styled_table(
    ['Kriter', 'Django REST (Python)', 'Express.js (Node)', 'Spring Boot (Java)'],
    [
        ('Geliştirme Hızı / Öğrenme Eğrisi', '5', '4', '2'),
        ('Dahili Güvenlik (Auth/RBAC)', '5', '2', '4'),
        ('Ekip Yetkinliği', '5', '3', '2'),
        ('Sınıf Diyagramı (OOP) Uyumluluğu', '4', '3', '5'),
        ('TOPLAM PUAN', '19', '12', '13'),
    ], [5.5, 3.5, 3.5, 3.5]
)
doc.add_paragraph()
p('Sonuç: Güvenlik altyapısı ve ekibin hızı göz önüne alınarak Django REST Framework seçilmiştir.', bold=True, italic=True)

doc.add_paragraph()
p('4.2.2. Veritabanı Karar Matrisi', bold=True)
styled_table(
    ['Kriter', 'PostgreSQL', 'MongoDB', 'SQLite'],
    [
        ('Veri Bütünlüğü (ACID Uyumu)', '5', '2', '4'),
        ('Sınıf/İlişki Modelleme Kolaylığı', '5', '3', '5'),
        ('Production Ölçeklenebilirlik', '5', '5', '1'),
        ('Medikal Veri Güvenliği Standartları', '5', '3', '2'),
        ('TOPLAM', '20', '13', '12'),
    ], [5.5, 3.5, 3.5, 3.5]
)
doc.add_paragraph()
p('Sonuç: İlişkisel yapının zorunluluğu ve ölçeklenebilirlik sebebiyle PostgreSQL seçilmiştir.', bold=True, italic=True)

doc.add_paragraph()
p('4.2.3. Frontend Framework Karar Matrisi', bold=True)
styled_table(
    ['Kriter', 'React.js', 'Vue.js', 'Angular'],
    [
        ('Topluluk Desteği ve Kaynak', '5', '4', '3'),
        ('Odontogram için DOM Performansı', '5', '4', '3'),
        ('Ekip Yetkinliği', '5', '2', '1'),
        ('Sınır Sınıfları (Component) Uyumu', '5', '5', '5'),
        ('TOPLAM', '20', '15', '12'),
    ], [5.5, 3.5, 3.5, 3.5]
)
doc.add_paragraph()
p('Sonuç: DOM manipülasyon gücü ve ekip tecrübesi nedeniyle React.js seçilmiştir.', bold=True, italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# BÖLÜM 5: GERÇEKLEME VE KALİTE ANALİZİ (OPSİYON A)
# ═══════════════════════════════════════════════════════════
doc.add_heading('BÖLÜM 5: Gerçekleme ve Kalite Analizi', level=1)
p('[Opsiyon A: İç ve Dış Arayüzlerin Gerçeklenmesi ve Dokümantasyonu]', bold=True, italic=True)
p('Bu bölüm, Proje Adımı 3 kapsamında gerçeklenen çalışmayı içermektedir. Yaşca Diş Kliniği Yönetim Sistemi, React.js (Frontend) ve Django REST Framework (Backend) kullanılarak geliştirilen API odaklı (API-first) bir SaaS projesidir. Sistemde yoğun bir arayüz/entegrasyon mimarisi bulunduğu için Opsiyon A seçilmiştir.')

doc.add_heading('5.1. İç Arayüzler ve API Kontratları', level=2)
p('Projedeki katmanlar (Model, ViewSet, Serializer) birbirleriyle JSON paketleri üzerinden haberleşmektedir. Endüstri standardı olan Swagger (drf-spectacular) kullanılarak tüm API uç noktaları canlı olarak dokümante edilmiştir.')

# Swagger screenshot
if os.path.exists(SWAGGER):
    p('Şekil: Swagger UI – Canlı API Dokümantasyonu', bold=True, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_image_safe(SWAGGER, width=Inches(5.5))
    doc.add_paragraph()

p('API Endpoint Kontratları ve Hata Kodları', bold=True)

p('A) Kimlik Doğrulama: POST /api/auth/token/ (F-001)', bold=True)
styled_table(['Parametre','Değer'],[
    ('Açıklama', 'JWT ile sisteme giriş'),
    ('Request Body', '{"username": "admin_user", "password": "securepassword123"}'),
    ('Response (200 OK)', '{"access": "eyJhbGciOi...", "refresh": "eyJhbGciOi..."}'),
    ('Hata Kodları', '401 Unauthorized (Hatalı giriş), 400 Bad Request (Eksik parametre)'),
],[4,13])

doc.add_paragraph()
p('B) Hasta Yönetimi: GET & POST /api/patients/ (F-003, F-004, F-005)', bold=True)
p('POST Request Body:')
code_block('{\n  "first_name": "Ahmet",\n  "last_name": "Yılmaz",\n  "phone": "05551234567"\n}')
p('GET Response (200 OK):')
code_block('[\n  {\n    "id": 1,\n    "first_name": "Ahmet",\n    "last_name": "Yılmaz",\n    "full_name": "Ahmet Yılmaz",\n    "phone": "05551234567",\n    "tckn": null,\n    "last_visit": "2026-05-20"\n  }\n]')
p('Hata Kodları: 400 Bad Request (Zorunlu alan eksik).')

doc.add_paragraph()
p('C) Randevu Yönetimi: POST /api/appointments/ (F-007, F-008)', bold=True)
styled_table(['Parametre','Değer'],[
    ('Açıklama', 'Çakışma kontrollü randevu oluşturma'),
    ('Request Body', '{"patient": 15, "doctor": 3, "date": "2026-06-15", "time": "14:30:00"}'),
    ('Response (201 Created)', 'Oluşturulan randevu nesnesi (id: 42)'),
    ('Hata Kodları', '400 Bad Request (Çakışma-Conflict), 403 Forbidden (Yetkisiz)'),
],[4,13])

doc.add_paragraph()
p('D) Dashboard: GET /api/dashboard/today/ (F-013)', bold=True)
p('Response (200 OK):')
code_block('{\n  "today_appointments": [...],\n  "today_total": 12,\n  "today_completed": 4,\n  "total_patients": 150\n}')

doc.add_page_break()

doc.add_heading('5.2. Dış Servis Entegrasyonları', level=2)
p('Harici SMS Sağlayıcı Entegrasyonu (F-012)', bold=True)
p('Randevu oluşturulduğunda harici SMS API\'sine bildirim isteği atılır. Dış API timeout olursa işlem çökmez; hata log sistemine kaydedilir ve randevu asenkron olarak tamamlanır.')
p('Payload: {"api_key": "YOUR_SECRET_KEY", "phone_number": "+905551234567", "message": "Randevunuz onaylanmıştır."}')

p('Frontend (React.js) Entegrasyonu', bold=True)
p('Frontend, sunucu ile axios kullanarak haberleşmektedir. JWT (Bearer Token) otomatik olarak başlığa eklenir:')
code_block('// frontend/src/services/api.js\nimport axios from \'axios\';\n\nconst api = axios.create({ baseURL: \'http://127.0.0.1:8000/api/\' });\n\napi.interceptors.request.use((config) => {\n    const token = localStorage.getItem(\'accessToken\');\n    if (token) config.headers.Authorization = `Bearer ${token}`;\n    return config;\n});')

doc.add_heading('5.3. Multi-Tenant (SaaS) Mimarisi', level=2)
p('Uygulama django-tenants kütüphanesi kullanılarak çoklu kiracı (schema-per-tenant) yapısında kodlanmıştır. HeaderTenantMiddleware, her gelen HTTP isteğinde domain/tenant kimliğini okur ve ilgili kliniğin PostgreSQL şemasına yönlendirir. Böylece veri güvenliği fiziksel izolasyonla sağlanır.')
p('Middleware Çalışma Mantığı:', bold=True)
bullet('X-Tenant header varsa → o tenant\'a bağlan (canlı ortam)')
bullet('Host header\'da bilinen subdomain varsa → o tenant\'a bağlan (lokal geliştirme)')
bullet('Hiçbiri değilse → public tenant\'a bağlan (ana site)')

code_block('class HeaderTenantMiddleware:\n    def __call__(self, request):\n        TenantModel = get_tenant_model()\n        DomainModel = get_tenant_domain_model()\n        tenant = None\n        \n        # 1. X-Tenant header\n        tenant_header = request.META.get(\'HTTP_X_TENANT\', \'\').strip()\n        if tenant_header:\n            try:\n                tenant = TenantModel.objects.get(schema_name=tenant_header)\n            except TenantModel.DoesNotExist:\n                pass\n        \n        # 2. Host header fallback\n        if not tenant:\n            hostname = request.get_host().split(\':\')[0]\n            try:\n                domain_obj = DomainModel.objects.select_related(\'tenant\').get(domain=hostname)\n                tenant = domain_obj.tenant\n            except DomainModel.DoesNotExist:\n                pass\n        \n        # 3. Public fallback\n        if not tenant:\n            tenant = TenantModel.objects.get(schema_name=\'public\')\n        \n        request.tenant = tenant\n        connection.set_tenant(tenant)\n        return self.get_response(request)')

doc.add_heading('5.4. Gerçek Kod Örnekleri', level=2)

p('Model Katmanı (Entity) – Appointment Modeli:', bold=True)
code_block('class Appointment(models.Model):\n    class Status(models.TextChoices):\n        SCHEDULED = "scheduled", "Planlandı"\n        COMPLETED = "completed", "Tamamlandı"\n        CANCELLED = "cancelled", "İptal"\n        NO_SHOW = "no_show", "Gelmedi"\n\n    patient = models.ForeignKey(Patient, on_delete=models.CASCADE,\n                                related_name="appointments")\n    doctor = models.ForeignKey(CustomUser, on_delete=models.PROTECT)\n    date = models.DateField("Tarih")\n    time = models.TimeField("Saat")\n    status = models.CharField(max_length=20, choices=Status.choices,\n                              default=Status.SCHEDULED)')

p('Serializer Katmanı – Çakışma Algoritması (F-008):', bold=True)
code_block('class AppointmentCreateSerializer(serializers.ModelSerializer):\n    def validate(self, data):\n        doctor = data.get("doctor")\n        date = data.get("date")\n        time = data.get("time")\n        if doctor and date and time:\n            existing = Appointment.objects.filter(\n                doctor=doctor, date=date, time=time,\n                status=Appointment.Status.SCHEDULED, is_active=True\n            )\n            if self.instance:\n                existing = existing.exclude(pk=self.instance.pk)\n            if existing.exists():\n                raise serializers.ValidationError(\n                    "Bu hekime bu saatte zaten randevu kayıtlı."\n                )\n        return data')

p('ViewSet Katmanı (Controller) – RBAC:', bold=True)
code_block('class AppointmentViewSet(AuditLogMixin, viewsets.ModelViewSet):\n    permission_classes = [IsAuthenticated]\n\n    def get_queryset(self):\n        user = self.request.user\n        qs = Appointment.objects.filter(is_active=True)\\\n                                .select_related("patient", "doctor")\n        # RBAC: Doktor sadece kendi randevularını görür\n        if user.role == CustomUser.Role.DOCTOR and not user.is_superuser:\n            qs = qs.filter(doctor=user)\n        return qs.order_by("date", "time")')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# BÖLÜM 6: BİRİM TEST SONUÇLARI
# ═══════════════════════════════════════════════════════════
doc.add_heading('BÖLÜM 6: Birim Test Sonuçları', level=1)

p('Bu bölüm, Proje Adımı 4 kapsamında gerçekleştirilen birim test çalışmalarının sonuçlarını sunmaktadır. Test ortamı olarak pytest + pytest-django kullanılmış, sahte test verisi factory-boy + Faker ile üretilmiş ve bağımlılıklar unittest.mock / monkeypatch ile taklit edilmiştir.')

p('Test edilen 5 kritik sınıf, sınıf diyagramından zengin iş mantığı ve hata yönetimi içeren yapılar olarak seçilmiştir.')

doc.add_heading('6.1. Birim Test Durum Tablosu', level=2)
p('Koşum: pytest api/tests/test_serializers.py api/tests/test_middleware.py customers/tests/test_register.py → 54 passed')

test_rows = [
    ('UT-001', 'AppointmentCreateSerializer', 'validate()', 'Negatif', 'Aynı hekime aynı tarih/saatte ikinci randevu', 'PASSED'),
    ('UT-002', 'AppointmentSerializer', 'validate()', 'Negatif', 'Güncelleme serializer\'ında da çakışma reddi', 'PASSED'),
    ('UT-003', 'AppointmentCreateSerializer', 'validate()', 'Pozitif', 'Aynı hekime farklı saatte randevu', 'PASSED'),
    ('UT-004', 'AppointmentCreateSerializer', 'validate()', 'Pozitif', 'İPTAL edilmiş randevunun slotu bloke etmemesi', 'PASSED'),
    ('UT-005', 'AppointmentSerializer', 'validate()', 'Pozitif', 'Güncelleme: kendi slotunu çakışma saymaması', 'PASSED'),
    ('UT-006', 'HeaderTenantMiddleware', '__call__()', 'Pozitif', 'X-Tenant header ile doğru kiracı çözümleme', 'PASSED'),
    ('UT-007', 'HeaderTenantMiddleware', '__call__()', 'Pozitif', 'Header yokken Host adından kiracı çözümleme', 'PASSED'),
    ('UT-008', 'HeaderTenantMiddleware', '__call__()', 'Pozitif', 'Header ve Host birlikte: Header önceliği', 'PASSED'),
    ('UT-009', 'HeaderTenantMiddleware', '__call__()', 'Negatif', 'Geçersiz X-Tenant → Host\'a fallback', 'PASSED'),
    ('UT-010', 'HeaderTenantMiddleware', '__call__()', 'Negatif', 'Boş X-Tenant → public tenant\'a düşme', 'PASSED'),
    ('UT-011', 'HeaderTenantMiddleware', '__call__()', 'Negatif', 'Public tenant tanımsız → HTTP 500', 'PASSED'),
    ('UT-012', 'RegisterClinicView', 'post()', 'Pozitif', 'Geçerli verilerle yeni klinik kaydı (mock)', 'PASSED'),
    ('UT-013', 'RegisterClinicView', 'post()', 'Pozitif', 'Kayıt sonrası admin superuser oluşması', 'PASSED'),
    ('UT-014', 'RegisterClinicView', 'post()', 'Negatif', 'Zorunlu alan eksik → HTTP 400', 'PASSED'),
    ('UT-015', 'RegisterClinicView', 'post()', 'Negatif', 'Kayıtlı subdomain ile tekrar kayıt → 400', 'PASSED'),
    ('UT-016', 'RegisterClinicView', 'post()', 'Negatif', 'Beklenmedik istisna → HTTP 500', 'PASSED'),
    ('UT-017', 'CheckDomainView', 'get()', 'Pozitif', 'Var olan subdomain sorgusu → exists=True', 'PASSED'),
    ('UT-018', 'CheckDomainView', 'get()', 'Negatif', 'Boş subdomain parametresi → 400', 'PASSED'),
    ('UT-019', 'PatientSerializer', 'create()', 'Pozitif', 'Anamnez olmadan hasta kaydı', 'PASSED'),
    ('UT-020', 'PatientSerializer', 'create()', 'Pozitif', 'Nested anamnez verisiyle hasta oluşturma', 'PASSED'),
    ('UT-021', 'PatientSerializer', 'update()', 'Pozitif', 'Anamnez yoksa güncelleme ile oluşturma', 'PASSED'),
    ('UT-022', 'PatientSerializer', 'create()', 'Negatif', 'Telefon olmadan kayıt → is_valid=False', 'PASSED'),
    ('UT-023', 'PatientSerializer', 'create()', 'Negatif', 'Boş ad (first_name) ile kayıt → reddedildi', 'PASSED'),
    ('UT-024', 'TreatmentSerializer', 'validate()', 'Pozitif', 'Aynı gün farklı dişe tedavi → geçerli', 'PASSED'),
    ('UT-025', 'TreatmentSerializer', 'validate()', 'Negatif', 'Aynı gün/diş/tür mükerrer → ValidationError', 'PASSED'),
]

styled_table(
    ['Test ID', 'Hedef Sınıf', 'Metot', 'Tür', 'Senaryo Açıklaması', 'Durum'],
    test_rows,
    [1.3, 3.2, 1.5, 1.3, 7, 1.3]
)

doc.add_paragraph()
p('Özet: 25 temsili senaryo (13 Pozitif + 12 Negatif) — TÜMÜ PASSED. İlgili 3 test dosyasının tam koşumu: 54/54 PASSED.', bold=True)

doc.add_heading('6.2. Test İyileştirmeleri', level=2)
p('Tüm 54 birim test başarıyla (PASSED) sonuçlanmıştır; başarısız (FAILED) test bulunmamaktadır. Ancak test geliştirme sürecinde şu iyileştirmeler yapılmıştır:')
bullet('Çakışma Kontrolü İyileştirmesi: İlk sürümde iptal edilmiş randevular da çakışma olarak algılanıyordu. validate() metoduna status=SCHEDULED filtresi eklenerek yalnızca planlanmış randevuların çakışma kontrolüne dahil edilmesi sağlandı.')
bullet('Güncelleme Çakışma Bypass: Mevcut bir randevu güncellenirken kendi slotunun çakışma olarak sayılmaması için exclude(pk=self.instance.pk) kontrolü eklendi.')
bullet('Middleware Fallback Güvenliği: Public tenant bulunamadığında sistemin sessizce hata vermesi yerine açık bir HTTP 500 yanıtı döndürmesi sağlandı.')
bullet('Mocking Stratejisi: HeaderTenantMiddleware ve RegisterClinicView testlerinde gerçek PostgreSQL/django-tenants bağımlılığı olmadan, monkeypatch ile izole test ortamı kuruldu.')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# BÖLÜM 7: BAKIM VE GELECEK ÖNERİLERİ
# ═══════════════════════════════════════════════════════════
doc.add_heading('BÖLÜM 7: Bakım ve Gelecek Önerileri', level=1)

doc.add_heading('7.1. Sürdürülebilirlik ve İyileştirme', level=2)

p('Refactoring İhtiyaçları', bold=True)
bullet('Servis Katmanı Ayrımı: Mevcut durumda iş mantığının bir kısmı ViewSet ve Serializer katmanlarında yer almaktadır. Karmaşık iş kurallarını (randevu çakışması, ödeme doğrulaması) bağımsız bir Service Layer\'a taşımak bakım maliyetini düşürecektir.')
bullet('Test Kapsamının Genişletilmesi: Mevcut birim testlerin yanına API seviyesinde entegrasyon testleri ve Playwright ile uçtan uca (E2E) testler eklenmelidir.')
bullet('Hata Yönetimi Merkezi Yapısı: API hatalarının RFC 7807 Problem Details standardında döndürülmesi için merkezi bir exception handler yapısına geçiş planlanmaktadır.')

p('Yeni Özellik Önerileri', bold=True)
bullet('Gerçek SMS/WhatsApp entegrasyonu ile otomatik randevu hatırlatmaları (F-012 tam gerçekleme)')
bullet('Randevu takviminde sürükle-bırak (drag & drop) desteği')
bullet('Dijital röntgen ve görüntüleme dosyalarının hasta profiline eklenmesi (F-016 geliştirilmesi)')
bullet('Raporlama modülü: Aylık/yıllık gelir, hasta istatistikleri, tedavi dağılımları')
bullet('Mobil uygulama (React Native) ile hekimlerin uzaktan erişimi')
bullet('Çok dilli destek (i18n) ile farklı ülkelere açılma imkânı')

p('Versiyon Güncelleme Planı', bold=True)
styled_table(
    ['Versiyon', 'Planlanan Özellikler', 'Tahmini Tarih'],
    [
        ('v1.1', 'SMS/WhatsApp entegrasyonu, takvim drag & drop', '2025 Q3'),
        ('v1.2', 'Raporlama modülü, PDF çıktı desteği', '2025 Q4'),
        ('v2.0', 'Mobil uygulama (React Native), çok dilli destek (i18n)', '2026 Q1'),
        ('v2.1', 'Röntgen entegrasyonu, gelişmiş analitik dashboard', '2026 Q2'),
    ],
    [3, 10, 4]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# BÖLÜM 8: SONUÇ
# ═══════════════════════════════════════════════════════════
doc.add_heading('BÖLÜM 8: Sonuç', level=1)

doc.add_heading('8.1. Süreç Değerlendirmesi', level=2)
p('"Yaşca Diş Kliniği Yönetim Sistemi", yazılım mühendisliğinin temel disiplinleri olan analiz, tasarım, gerçekleme, test ve bakım aşamalarının tamamını kapsayacak şekilde geliştirilmiştir.')
p('Proje Adımı 1\'de gerçekleştirilen sistem analizi ve UML modelleme çalışmasıyla, İster Analizi dokümanındaki 22 fonksiyonel ve 10 fonksiyonel olmayan gereksinim detaylı sekans ve sınıf diyagramları ile modellenmiştir. Boundary-Control-Entity ayrımına dayanan bu modelleme, kodlama aşamasının temelini oluşturmuştur.')
p('Proje Adımı 2\'de hazırlanan Teknoloji Karar Analizi (DAR) raporu ile Django REST Framework, React.js ve PostgreSQL\'den oluşan teknoloji yığını, karar matrisleri kullanılarak objektif kriterlerle seçilmiştir.')
p('Proje Adımı 3\'te iç ve dış arayüzlerin gerçeklenmesi kapsamında, RESTful API kontratları Swagger ile canlı olarak dokümante edilmiş; frontend-backend entegrasyonu ve django-tenants ile multi-tenant altyapı kodlanmıştır.')
p('Proje Adımı 4\'te gerçekleştirilen birim test çalışmasında, sınıf diyagramından seçilen 5 kritik sınıfın çekirdek metotları hem pozitif hem negatif senaryolarla test edilmiştir. Bağımlılıklar monkeypatch/MagicMock ile taklit edilerek metotlar gerçekten izole edilmiştir. İlgili test dosyalarının tamamı 54/54 PASSED ile sonuçlanmıştır.')

p('Karşılaşılan Zorluklar:', bold=True)
bullet('Multi-tenant mimari kurulumu ve django-tenants ile şema yönetimi ilk aşamada karmaşık bulunmuş, ancak HeaderTenantMiddleware ile çözülmüştür.')
bullet('Çakışma kontrolü algoritmasında iptal edilmiş randevuların da çakışma olarak algılanması sorunu tespit edilmiş ve status filtresi ile giderilmiştir.')
bullet('Birim testlerde gerçek veritabanı bağımlılığının kaldırılması için mocking stratejisinin doğru kurgulanması iteratif çalışma gerektirmiştir.')

doc.add_heading('8.2. Temel Çıktıların Kontrolü', level=2)
p('Proje sürecinde elde edilen temel çıktılar aşağıda doğrulanmıştır:')

ciktilar = [
    ('Uçtan uca test edilmiş modüller', '54/54 birim test PASSED; 5 kritik sınıf pozitif ve negatif senaryolarla doğrulandı.'),
    ('Standartlara uygun UML dokümantasyonu', '6 sekans diyagramı + 2 sınıf diyagramı standart UML notasyonlarına uygun olarak hazırlandı.'),
    ('Teknik karar verme yetkinliği', '3 ayrı karar matrisi (Backend, DB, Frontend) ile teknoloji yığını objektif olarak seçildi.'),
    ('Ölçeklenebilir multi-tenant SaaS mimarisi', 'django-tenants ile PostgreSQL şema bazlı kiracı izolasyonu gerçeklendi.'),
    ('Canlı API dokümantasyonu', 'Swagger (drf-spectacular) ile tüm endpointler interaktif olarak dokümante edildi.'),
    ('Bütünsel tutarlılık', 'İsterler → UML → DAR → Gerçekleme → Test zincirinde tam izlenebilirlik sağlandı.'),
]

styled_table(
    ['Çıktı', 'Doğrulama'],
    ciktilar,
    [5, 12]
)

doc.add_paragraph()
p('Bu proje, yazılım geliştirme yaşam döngüsünün tüm aşamalarını kapsayan kapsamlı bir mühendislik çalışması olarak tamamlanmıştır. Edinilen deneyimler ve oluşturulan altyapı, sistemin gelecekte genişletilmesi için sağlam bir temel oluşturmaktadır.')

# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════
out = r"c:\Users\Ali\yasca-dental-clinic\Grup_4_Final_Proje_Raporu.docx"
doc.save(out)
print(f"Rapor olusturuldu: {out}")
